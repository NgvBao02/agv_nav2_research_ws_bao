#!/usr/bin/env python3
"""Generate the Vietnamese–English terminology and writing guide for PSTMO."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PSTMO_thuat_ngu_Anh_Viet.docx"
FONT = "Times New Roman"
NAVY = "17365D"
TEAL = "137A7F"
LIGHT = "EAF2F8"
PALE = "EAF6F2"
WHITE = "FFFFFF"


def set_cell_shading(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 9.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[tuple | list], widths: list[float] | None = None, size: float = 9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, item in enumerate(headers):
        set_cell_text(header.cells[i], item, bold=True, color=WHITE, size=size)
        set_cell_shading(header.cells[i], NAVY)
        if widths:
            header.cells[i].width = Cm(widths[i])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, item in enumerate(row):
            set_cell_text(cells[i], str(item), size=size)
            if widths:
                cells[i].width = Cm(widths[i])
            if ridx % 2:
                set_cell_shading(cells[i], "F4F7F9")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def add_para(doc: Document, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.15
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Trang ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    for name, size, color in [("Title", 24, NAVY), ("Heading 1", 16, NAVY), ("Heading 2", 13, TEAL), ("Heading 3", 11.5, NAVY)]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    props = doc.core_properties
    props.title = "Sổ tay thuật ngữ và cách diễn đạt PSTMO"
    props.subject = "Thuật ngữ Anh–Việt dùng trong slide và báo cáo PSTMO"
    props.author = "NGUYỄN TIẾN CƯƠNG"
    props.keywords = "PSTMO, ROS 2, Nav2, terminology, path smoothing, RViz2, Gazebo"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("SỔ TAY THUẬT NGỮ\nVÀ CÁCH DIỄN ĐẠT PSTMO")
    r.bold = True
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Bản chuẩn hóa Anh–Việt dùng cho slide, thuyết trình và báo cáo")
    r.italic = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    p.add_run("NGUYỄN TIẾN CƯƠNG\n2026").bold = True
    doc.add_page_break()

    add_heading(doc, "1. Nguyên tắc dùng thuật ngữ", 1)
    add_bullets(doc, [
        "Lần xuất hiện đầu tiên: viết thuật ngữ tiếng Việt trước, thuật ngữ tiếng Anh trong ngoặc; các lần sau ưu tiên tiếng Việt.",
        "Tên lớp, chủ đề ROS, thông báo và mã lỗi trong mã nguồn được giữ nguyên trong kiểu chữ mã, sau đó giải thích bằng tiếng Việt.",
        "Không dùng một từ cho hai khái niệm khác nhau. Đặc biệt phải tách đường hình học (path) khỏi quỹ đạo theo thời gian (trajectory).",
        "Mọi đại lượng phải đi kèm ký hiệu và đơn vị. Góc phải ghi rõ độ hoặc radian; độ cong và tích phân bình phương độ cong dùng m⁻¹.",
        "Không dùng các từ đánh giá mơ hồ như “tốt hơn”, “đẹp hơn”, “an toàn hơn” nếu chưa nói rõ chỉ số, phương án đối chứng và phạm vi dữ liệu.",
        "Raw, Simple, Savitzky–Golay, Constrained và PSTMO là tên phương án; giữ cách viết hoa nhất quán.",
    ])
    add_para(doc, "Mẫu viết lần đầu: “Bộ làm mượt đường đi (path smoother) PSTMO nhận nav_msgs/Path và trả về một đường hình học đã xử lý.” Mẫu viết về sau: “bộ làm mượt PSTMO”.")

    add_heading(doc, "2. Các phân biệt bắt buộc", 1)
    distinctions = [
        ("Đường đi và quỹ đạo", "Đường đi/path là chuỗi tư thế hình học, chưa gắn thời gian. Quỹ đạo/trajectory gắn trạng thái với thời gian và có thể kèm v, ω, gia tốc."),
        ("Điểm và tư thế", "Điểm chỉ có tọa độ; tư thế/pose gồm vị trí và hướng. Không gọi (x, y, ψ) là một điểm đơn thuần."),
        ("Lập kế hoạch, làm mượt, điều khiển", "Bộ lập kế hoạch chọn đường đi qua đâu; bộ làm mượt chỉnh hình học đường; bộ điều khiển phát lệnh v, ω để bám đường."),
        ("Bản đồ và bản đồ chi phí", "Bản đồ mô tả môi trường; bản đồ chi phí/costmap là lưới giá trị phục vụ lập kế hoạch và kiểm tra an toàn."),
        ("Hình bao và vùng quét", "Hình bao robot/footprint là đa giác tại một tư thế; vùng quét hình bao/swept footprint là hợp các hình bao dọc chuyển động nội suy."),
        ("Đường kế hoạch và quỹ đạo thực mô phỏng", "Đường kế hoạch được gửi cho FollowPath; quỹ đạo thực mô phỏng lấy từ trạng thái vật lý Gazebo. Hai đường không được gọi chung là kết quả RViz2."),
        ("Giới hạn vận tốc và biểu đồ vận tốc", "Giới hạn vận tốc là ngưỡng cục bộ theo độ cong; biểu đồ vận tốc là dãy v(s) sau quét tiến/lùi để thỏa tăng tốc và giảm tốc."),
        ("κmax và Eκ", "κmax phản ánh đoạn uốn gắt nhất. Eκ=∫κ²ds phản ánh mức uốn tích lũy; Eκ là chỉ số hình học, không phải điện năng."),
        ("G¹ và G²", "G¹ bảo đảm hướng tiếp tuyến liên tục; G² còn bảo đảm độ cong liên tục. G² không đồng nghĩa với jerk theo thời gian bằng 0."),
        ("Góc chuyển hướng θ và góc hướng ψ", "θ là độ đổi hướng giữa cạnh vào và cạnh ra; ψ là hướng của một tư thế trong hệ tọa độ."),
        ("Khoảng cách bánh mô hình và hiệu dụng", "b=0,2548 m là khoảng cách bánh dùng trong mô hình động học; 0,2834 m là giá trị hiệu dụng trong Gazebo sau hệ số nhân. Phải ghi rõ loại giá trị."),
        ("RPP và Collision Monitor", "Cảnh báo dự báo va chạm phía trước của RPP không đồng nghĩa với một lần Collision Monitor can thiệp."),
        ("Hình 3D và ảnh mô phỏng", "Hình 3D từ STL/SDF dùng để giải thích cấu tạo; ảnh Gazebo xác nhận trạng thái vật lý và môi trường; ảnh RViz2 xác nhận dữ liệu điều hướng. Không dùng hình 3D minh họa thay bằng chứng thực nghiệm."),
    ]
    add_table(doc, ["Cặp khái niệm", "Cách phân biệt chuẩn"], distinctions, [5.0, 12.0], 9.3)

    add_heading(doc, "3. Bảng sửa từ ngữ", 1)
    replacements = [
        ("planner", "bộ lập kế hoạch; bộ lập kế hoạch toàn cục", "Khối tạo đường từ điểm đầu tới điểm đích"),
        ("đường planner", "đường do bộ lập kế hoạch tạo", "Tránh ghép Việt–Anh không tự nhiên"),
        ("smoother", "bộ làm mượt đường đi", "Khối chỉnh hình học đường"),
        ("controller", "bộ điều khiển bám đường", "Khối phát lệnh vận tốc"),
        ("ứng viên; candidate dùng riêng lẻ", "phương án xử lý tại góc", "Cách gọi tự nhiên hơn mà vẫn đúng vai trò"),
        ("transition", "đoạn chuyển tiếp", "Trong bài là đoạn Bézier bậc năm G²"),
        ("pivot", "quay tại chỗ", "Tâm robot không tịnh tiến khi đổi hướng"),
        ("pass-through / đi xuyên qua", "giữ nguyên góc", "Góc nhỏ được giữ, không chèn đoạn chuyển tiếp"),
        ("waypoint", "điểm đường", "Một điểm lấy mẫu trên đường"),
        ("pose", "tư thế", "Gồm vị trí và hướng"),
        ("yaw", "góc hướng; góc quay quanh trục z", "Ghi rõ đơn vị độ hoặc radian"),
        ("costmap", "bản đồ chi phí", "Lưới chi phí dùng trong Nav2"),
        ("cost", "giá trị chi phí; chi phí bản đồ", "Không dùng “rủi ro cost”"),
        ("footprint", "hình bao chiếm chỗ của robot", "Đa giác thân xe chiếm chỗ"),
        ("swept-footprint", "vùng quét hình bao robot", "Hợp các hình bao dọc chuyển động"),
        ("clearance", "khoảng hở hình bao", "Khoảng cách nhỏ nhất tới vật cản"),
        ("speed cap", "giới hạn vận tốc cục bộ", "Chỉ một ngưỡng; không đại diện cho cả biểu đồ"),
        ("speed profile", "biểu đồ vận tốc dọc đường", "Dãy v(s) sau các ràng buộc gia tốc"),
        ("time gate", "điều kiện ưu thế thời gian", "So sánh đoạn chuyển tiếp với quay tại chỗ"),
        ("hard gate / cổng cứng", "điều kiện loại bắt buộc", "Không đạt thì loại phương án"),
        ("soft objective", "hàm mục tiêu lựa chọn", "Chỉ áp dụng sau các điều kiện bắt buộc"),
        ("pipeline", "quy trình xử lý", "Chuỗi bước theo thứ tự"),
        ("runtime", "thời gian xử lý thuật toán", "Ghi ms hoặc s"),
        ("diagnostics", "dữ liệu chẩn đoán", "Dữ liệu nội bộ giúp truy vết quyết định"),
        ("baseline", "phương án đối chứng", "Mốc dùng để so sánh"),
        ("benchmark", "phép thử đánh giá; thử nghiệm đối sánh", "Nêu rõ thiết kế và số lần lặp"),
        ("ground truth", "giá trị thực mô phỏng", "Trạng thái vật lý lấy trực tiếp từ Gazebo"),
        ("live image", "ảnh chụp trực tiếp từ hệ thống", "Không gọi chung là ảnh minh họa"),
        ("exact path", "đường ROS Path được lưu trực tiếp", "Nêu nguồn dữ liệu thay vì từ cảm tính"),
        ("start / goal", "điểm đầu / điểm đích", "Dùng tiếng Việt trong câu văn"),
        ("success", "thành công; hoàn tất tác vụ", "Phải nêu điều kiện xác định thành công"),
        ("collision count", "số mẫu hình bao bị va chạm", "Phân biệt với sự kiện Collision Monitor"),
        ("win / lose", "nhanh hơn / chậm hơn; giá trị nhỏ hơn / lớn hơn", "Nêu đúng chỉ số so sánh"),
        ("energy", "chỉ số uốn chuẩn hóa eκ", "Không làm người đọc hiểu nhầm là điện năng"),
        ("curvature energy", "tích phân bình phương độ cong Eκ", "Tên toán học chính xác hơn"),
        ("risk", "thành phần chi phí bản đồ chuẩn hóa r_cost", "Không đồng nhất với an toàn tuyệt đối"),
        ("angular", "thành phần vận tốc góc chuẩn hóa rω", "Nêu rõ đại lượng được chuẩn hóa"),
        ("profile độ cong", "phân bố độ cong", "Cách viết tiếng Việt tự nhiên"),
        ("light smoother", "bộ làm mượt có chi phí xử lý thấp", "Không dùng “nhẹ” nếu chưa nêu tiêu chí"),
        ("Clr min", "khoảng hở nhỏ nhất (m)", "Tiêu đề bảng phải đọc được độc lập"),
        ("T chạy", "thời gian chạy (s)", "Ghi rõ đại lượng và đơn vị"),
        ("TB / Min / Max", "trung bình / nhỏ nhất / lớn nhất", "Tránh viết tắt không cần thiết"),
        ("đường đẹp hơn", "đường có κmax hoặc Eκ nhỏ hơn", "Gắn nhận định với chỉ số đo"),
        ("an toàn tuyệt đối", "không phát hiện va chạm trên cấu hình thử", "Giới hạn kết luận theo dữ liệu"),
        ("chứng minh an toàn", "cho thấy đạt kiểm tra an toàn trong mô phỏng", "Không suy rộng sang robot thật"),
        ("hết thời gian kiên nhẫn", "vượt thời gian chờ của bộ điều khiển", "Diễn giải đúng PATIENCE_EXCEEDED"),
        ("lượt tới goal", "lượt tới đích", "Tránh trộn Việt–Anh"),
    ]
    add_table(doc, ["Không nên dùng", "Dùng thống nhất", "Lý do / phạm vi"], replacements, [4.4, 5.5, 7.2], 8.8)

    add_heading(doc, "4. Từ điển thuật ngữ Anh–Việt", 1)
    glossary = [
        ("acceleration limit", "giới hạn gia tốc", "Giới hạn tốc độ thay đổi vận tốc"),
        ("AMCL", "định vị Monte Carlo thích nghi", "Ước lượng tư thế robot trên bản đồ"),
        ("corner-handling option (candidate solution)", "phương án xử lý tại góc", "Một trạng thái với d và α cụ thể; tiếng Việt không dùng từ “ứng viên”"),
        ("clearance", "khoảng hở hình bao", "Khoảng cách nhỏ nhất từ hình bao tới vật cản"),
        ("collision-free", "không phát hiện va chạm", "Kết quả nhị phân trên bản đồ chi phí hiện tại"),
        ("collision monitor", "bộ giám sát va chạm", "Mô-đun an toàn giám sát lệnh chuyển động"),
        ("conditioned polyline", "đường gấp khúc đã điều kiện hóa", "Đường sau bỏ trùng, RDP và triệt zíc-zắc"),
        ("controller", "bộ điều khiển bám đường", "Phát v và ω"),
        ("costmap", "bản đồ chi phí", "Lưới chi phí của Nav2"),
        ("curvature", "độ cong", "Mức đổi hướng trên một đơn vị chiều dài"),
        ("curvature energy", "tích phân bình phương độ cong", "Eκ=∫κ²ds, đơn vị m⁻¹"),
        ("differential drive", "truyền động vi sai", "Robot điều khiển bằng hai bánh chủ động"),
        ("dynamic programming", "quy hoạch động", "Chọn chuỗi trạng thái tương thích toàn đường"),
        ("effective wheel separation", "khoảng cách bánh hiệu dụng", "Khoảng cách có kể hệ số hiệu chỉnh mô phỏng"),
        ("endpoint invariant", "bất biến điểm đầu–đích", "Điểm đầu/đích và hướng đích không đổi"),
        ("feasible", "khả thi", "Đạt các điều kiện áp dụng"),
        ("footprint", "hình bao robot", "Đa giác thân xe chiếm chỗ"),
        ("G¹ continuity", "liên tục hình học bậc một", "Hướng tiếp tuyến liên tục"),
        ("G² continuity", "liên tục hình học bậc hai", "Hướng tiếp tuyến và độ cong liên tục"),
        ("global planner", "bộ lập kế hoạch toàn cục", "Tạo đường từ đầu tới đích"),
        ("ground truth", "giá trị thực mô phỏng", "Tư thế vật lý do Gazebo cung cấp"),
        ("hard constraint", "điều kiện bắt buộc", "Không đạt thì loại"),
        ("heading", "hướng chuyển động", "Hướng tiếp tuyến của đường"),
        ("kinematic feasibility", "tính khả thi động học", "Khả thi theo quan hệ v, ω và vận tốc bánh"),
        ("lateral acceleration", "gia tốc ngang", "a_y=v²|κ|"),
        ("lookahead", "khoảng nhìn trước", "Khoảng RPP dùng để chọn điểm bám"),
        ("motion primitive", "mẫu chuyển động nguyên thủy", "Đoạn chuyển động rời rạc của bộ lập kế hoạch"),
        ("odometry", "dữ liệu đo chuyển động", "Ước lượng chuyển động tương đối của robot"),
        ("pass-through", "giữ nguyên góc", "Không chèn chuyển tiếp khi góc nhỏ"),
        ("path", "đường đi; đường hình học", "Chuỗi tư thế chưa gắn thời gian"),
        ("path smoother", "bộ làm mượt đường đi", "Chỉnh hình học đường"),
        ("PATIENCE_EXCEEDED", "vượt thời gian chờ của bộ điều khiển", "Mã kết quả FollowPath 104 trong ca C30"),
        ("peak cost", "chi phí bản đồ lớn nhất", "Giá trị lớn nhất dọc phương án đang xét"),
        ("pivot", "quay tại chỗ", "Đổi hướng khi tâm robot không tịnh tiến"),
        ("pose", "tư thế", "Vị trí và hướng"),
        ("quintic Bézier curve", "đường cong Bézier bậc năm", "Đường cong có sáu điểm điều khiển"),
        ("Regulated Pure Pursuit", "điều khiển bám đường Pure Pursuit có điều chỉnh", "Bộ điều khiển RPP của Nav2"),
        ("RDP simplification", "rút gọn Ramer–Douglas–Peucker", "Giảm số điểm trong sai số cho phép"),
        ("shape ratio", "tỷ lệ hình dạng", "α=q/d"),
        ("smoothing time", "thời gian làm mượt", "Thời gian xử lý của thuật toán làm mượt"),
        ("speed limit", "giới hạn vận tốc", "Ngưỡng vận tốc cục bộ"),
        ("speed profile", "biểu đồ vận tốc", "Dãy v(s) dọc đường"),
        ("swept footprint", "vùng quét hình bao robot", "Hợp hình bao giữa các tư thế"),
        ("time gate", "điều kiện ưu thế thời gian", "Tfastest+0,15 s<Tpivot"),
        ("time parameterization", "tham số hóa theo thời gian", "Gán vận tốc và thời gian để kiểm tra khả thi"),
        ("trajectory", "quỹ đạo theo thời gian", "Trạng thái là hàm của thời gian"),
        ("transition", "đoạn chuyển tiếp", "Đoạn Bézier thay góc gãy"),
        ("trim distance", "khoảng cắt", "d: chiều dài cắt trên cạnh vào/ra"),
        ("velocity smoother", "bộ làm mượt vận tốc", "Làm mượt lệnh vận tốc, không làm mượt đường"),
        ("wheel speed", "vận tốc tuyến tính tại bánh", "vL hoặc vR"),
        ("yaw", "góc hướng", "Góc quay quanh trục z"),
    ]
    add_table(doc, ["Thuật ngữ tiếng Anh", "Tiếng Việt dùng trong bài", "Định nghĩa ngắn"], glossary, [4.6, 5.6, 6.9], 8.7)

    add_heading(doc, "5. Ký hiệu, công thức và đơn vị", 1)
    symbols = [
        ("x, y", "Tọa độ phẳng", "m"), ("ψ", "Góc hướng của tư thế", "rad hoặc °"),
        ("θ", "Góc chuyển hướng giữa cạnh vào và cạnh ra", "rad hoặc °"), ("L", "Chiều dài đường/đoạn", "m"),
        ("κ", "Độ cong", "m⁻¹"), ("κmax", "Trị tuyệt đối độ cong lớn nhất", "m⁻¹"),
        ("Eκ=∫κ²ds", "Tích phân bình phương độ cong", "m⁻¹"), ("Eref", "Thang chuẩn hóa Eκ", "1 m⁻¹"),
        ("d", "Khoảng cắt", "m"), ("q", "Khoảng điều khiển q=αd", "m"),
        ("α", "Tỷ lệ hình dạng", "không đơn vị"), ("b", "Khoảng cách hai bánh chủ động", "m"),
        ("v", "Vận tốc tịnh tiến", "m/s"), ("ω", "Vận tốc góc", "rad/s"),
        ("vL, vR", "Vận tốc tuyến tính tại bánh trái/phải", "m/s"), ("aacc, adec", "Giới hạn tăng tốc/giảm tốc", "m/s²"),
        ("ay", "Gia tốc ngang", "m/s²"), ("aω", "Gia tốc góc", "rad/s²"),
        ("T", "Thời gian", "s"), ("m", "Biên không chồng lấn", "m; cấu hình hiệu dụng 0,05 m"),
        ("rcost", "Chi phí bản đồ chuẩn hóa", "không đơn vị"), ("rω", "Vận tốc góc chuẩn hóa", "không đơn vị"),
        ("eκ", "Chỉ số uốn chuẩn hóa", "không đơn vị"), ("J", "Giá trị hàm mục tiêu", "không đơn vị"),
    ]
    add_table(doc, ["Ký hiệu", "Ý nghĩa", "Đơn vị"], symbols, [3.2, 9.9, 4.0], 9.0)
    add_para(doc, "Hàm mục tiêu dùng trong mã nguồn: r_cost=min(1, peak_cost/252); r_ω=min(1, |ω|max/ωmax); eκ=(Eκ/Eref)/(Eκ/Eref+1), với Eref=1 m⁻¹; J=0,15r_cost+0,10r_ω+0,75eκ. Thuật toán ưu tiên J nhỏ hơn.")
    add_para(doc, "Điều kiện ưu thế thời gian: T_fastest+0,15 s<T_pivot. Cửa sổ cạnh tranh tối đa là 10 s; đây là cửa sổ lọc các phương án được so sánh, không phải thời gian chạy của robot.")

    add_heading(doc, "6. Cách mô tả năm phương án", 1)
    methods = [
        ("Raw – đường gốc", "Đường do bộ lập kế hoạch tạo, chưa đi qua bộ làm mượt; dùng làm phương án đối chứng."),
        ("Simple", "Bộ làm mượt lặp cục bộ trên các điểm đường, cân bằng bám đường gốc và độ trơn cục bộ."),
        ("Savitzky–Golay", "Bộ lọc đa thức theo cửa sổ, có tác dụng giảm dao động tần số cao của chuỗi điểm."),
        ("Constrained", "Bộ làm mượt tối ưu có các thành phần độ trơn, độ cong, sai lệch và chi phí bản đồ; kết quả phụ thuộc cấu hình trọng số."),
        ("PSTMO", "Bộ làm mượt tạo đoạn chuyển tiếp Bézier bậc năm G², kiểm tra động học bánh xe, vùng quét hình bao, thời gian và chọn trạng thái toàn đường bằng quy hoạch động."),
    ]
    add_table(doc, ["Tên phương án", "Cách giới thiệu chuẩn"], methods, [4.2, 12.9], 9.3)

    add_heading(doc, "7. Câu mẫu có thể dùng trong báo cáo", 1)
    templates = [
        ("Nêu vấn đề", "Đường do bộ lập kế hoạch tạo có thể hợp lệ trên bản đồ chi phí nhưng vẫn chứa góc gãy khiến hướng tiếp tuyến thay đổi đột ngột và gây khó khăn cho bộ điều khiển bám đường."),
        ("Nêu mục tiêu", "PSTMO thay góc gãy bằng đoạn chuyển tiếp Bézier bậc năm G², đồng thời kiểm tra tính khả thi động học, vùng quét hình bao robot và ưu thế thời gian."),
        ("Giải thích G²", "G² được chọn thay vì chỉ G¹ vì G² bảo đảm cả hướng tiếp tuyến và độ cong liên tục tại mối nối; nhờ đó độ cong không nhảy từ đoạn thẳng sang đường cong."),
        ("Giải thích bậc năm", "Cấu trúc Bézier bậc năm cho phép áp các điều kiện biên vị trí, tiếp tuyến và độ cong bằng 0 ở hai đầu, đồng thời giữ tỷ lệ α làm tham số điều chỉnh hình dạng."),
        ("Định nghĩa phương án", "Một phương án xử lý tại góc là một cách xử lý cụ thể, gồm giữ nguyên góc, quay tại chỗ hoặc tạo đoạn chuyển tiếp xác định bởi d và α."),
        ("Nêu hàm mục tiêu", "Sau khi đạt các điều kiện bắt buộc, các phương án được so sánh theo chi phí bản đồ chuẩn hóa, vận tốc góc chuẩn hóa và chỉ số uốn chuẩn hóa; phương án có J nhỏ hơn được ưu tiên."),
        ("Nêu kết quả hình học", "Trên 34 nhóm ghép cặp đầy đủ, PSTMO giảm κmax trung bình 90,3% và giảm Eκ trung bình 98,5% so với đường gốc; các tỷ lệ này mô tả độ mượt hình học, không phải mức tiết kiệm điện năng."),
        ("Nêu kết quả thời gian", "Trong các cặp thực thi thành công, PSTMO có thời gian tới đích trung bình ngắn hơn đường gốc 3,97 s; kết luận này chỉ áp dụng cho cấu hình mô phỏng và số lần lặp hiện tại."),
        ("Nêu khoảng hở", "Khoảng hở hình bao của PSTMO thấp hơn một số phương án đối chứng nhưng toàn bộ mẫu hình bao vẫn đạt kiểm tra va chạm trên bản đồ chi phí dùng trong thử nghiệm."),
        ("Nêu ca C30", "Ca C30 gồm hai cơ chế thất bại khác nhau: Simple bị máy chủ làm mượt từ chối do va chạm hình bao; bốn phương án còn lại bị RPP dự báo va chạm và FollowPath kết thúc với PATIENCE_EXCEEDED, mã 104."),
        ("Nêu giới hạn", "Kết quả mô phỏng chưa đủ để khẳng định an toàn tuyệt đối trên robot thật, vì chưa mô hình hóa đầy đủ vật cản động, trượt bánh, tải, nhiễu và độ trễ cảm biến."),
    ]
    add_table(doc, ["Mục đích", "Câu viết đề xuất"], templates, [3.5, 13.6], 9.2)

    add_heading(doc, "8. Cách phân tích một thí nghiệm", 1)
    add_bullets(doc, [
        "Xác định ca thử: môi trường, bộ lập kế hoạch, điểm đầu, điểm đích và mã băm của đường gốc.",
        "Mô tả hình học đầu vào: số điểm, chiều dài, số góc, dải góc và góc có trị tuyệt đối lớn nhất.",
        "Mô tả quyết định của PSTMO: số góc giữ nguyên, số đoạn chuyển tiếp, số lần quay tại chỗ, d, α và lý do loại từng phương án nếu có.",
        "So sánh hình học theo cùng đường gốc: chiều dài, κmax, Eκ, khoảng hở nhỏ nhất và số mẫu hình bao bị va chạm.",
        "Tách thời gian xử lý thuật toán khỏi thời gian robot di chuyển tới đích.",
        "Đối chiếu đường kế hoạch trên RViz2 với quỹ đạo thực mô phỏng của Gazebo; không suy ra chất lượng bám đường chỉ từ hình ảnh tĩnh.",
        "Nếu thất bại, ghi đúng khối phát lỗi, thông báo, mã lỗi và thời điểm; không dùng thời gian dừng sớm như thời gian hoàn thành.",
        "Kết luận bằng chỉ số cụ thể và phạm vi dữ liệu, tránh câu tuyệt đối hóa.",
    ])
    add_para(doc, "Mẫu đọc góc: “Đường đã điều kiện hóa có 10 góc được xét, từ −45,77° tới +56,30°. Dấu âm/dương cho biết chiều quay; độ khó hình học không chỉ phụ thuộc trị tuyệt đối của góc mà còn phụ thuộc chiều dài hai cạnh kề, hành lang và vùng quét hình bao.”")

    add_heading(doc, "9. Các lỗi kỹ thuật đã sửa trong slide và báo cáo", 1)
    corrections = [
        ("Hàm mục tiêu cũ viết Eκ/(Eκ+1)", "Đã bổ sung Eref=1 m⁻¹ để phép cộng không sai thứ nguyên và đổi tên thành eκ, chỉ số uốn chuẩn hóa."),
        ("Dùng từ “energy”", "Đã đổi thành tích phân bình phương độ cong Eκ hoặc chỉ số uốn chuẩn hóa eκ; ghi rõ không phải điện năng."),
        ("Dùng dmax cho giảm tốc", "Đã dùng a_dec để không nhầm với khoảng cắt lớn nhất d_max."),
        ("Thiếu cửa sổ cạnh tranh", "Đã bổ sung giá trị tối đa 10 s và tách khỏi biên ưu thế thời gian 0,15 s."),
        ("segment_margin=0 bị hiểu là không có biên", "Đã nêu biên hiệu dụng 0,05 m do thuật toán tự chọn theo khoảng lấy mẫu và độ phân giải bản đồ chi phí."),
        ("pass-through dịch là “đi xuyên qua”", "Đã đổi thành “giữ nguyên góc”, đúng với hành vi không chèn đoạn chuyển tiếp."),
        ("PATIENCE_EXCEEDED dịch nghĩa đen", "Đã diễn giải là “vượt thời gian chờ của bộ điều khiển”, giữ nguyên mã 104."),
        ("Gộp cảnh báo RPP với Collision Monitor", "Đã tách dự báo va chạm nội bộ của RPP khỏi số lần Collision Monitor can thiệp."),
        ("Gộp khoảng cách bánh", "Đã tách b mô hình 0,2548 m và giá trị hiệu dụng Gazebo 0,2834 m."),
        ("Nhãn bảng viết tắt", "Đã đổi Clr min, T chạy, TB/Min/Max thành nhãn tiếng Việt đầy đủ và kèm đơn vị."),
        ("Lẫn path và trajectory", "Đã tách đường kế hoạch khỏi quỹ đạo thực mô phỏng trong phần RViz2/Gazebo."),
        ("Dùng hình 3D như ảnh kết quả", "Đã ghi rõ hình 3D chỉ giải thích cấu tạo; ảnh Gazebo/RViz2 và dữ liệu ROS mới là bằng chứng thực nghiệm."),
        ("Mô tả C30 quá chung", "Đã tách lỗi hình học Simple tại tọa độ cụ thể khỏi lỗi thực thi RPP/PATIENCE_EXCEEDED của bốn phương án còn lại."),
    ]
    add_table(doc, ["Lỗi / điểm mơ hồ", "Cách sửa đã áp dụng"], corrections, [6.0, 11.1], 9.1)

    add_heading(doc, "10. Danh sách kiểm tra trước khi nộp", 1)
    add_bullets(doc, [
        "Tên phương án, ký hiệu và đơn vị có nhất quán giữa hình, bảng, slide và báo cáo không?",
        "Mọi từ tiếng Anh lần đầu xuất hiện đã có thuật ngữ tiếng Việt tương ứng chưa?",
        "Đã phân biệt đường đi, quỹ đạo, tư thế và điểm chưa?",
        "Đã phân biệt thời gian làm mượt với thời gian tới đích chưa?",
        "Khi nói “giảm” hoặc “nhanh hơn”, đã nêu phương án đối chứng, số cặp và đơn vị chưa?",
        "Khi nói an toàn, đã giới hạn kết luận trong bản đồ chi phí và cấu hình thử chưa?",
        "Eκ có được gọi đúng là tích phân bình phương độ cong, không phải điện năng chưa?",
        "Góc đã ghi dấu, trị số và đơn vị; kết luận có xét thêm cạnh kề và hành lang chưa?",
        "Mã lỗi và khối phát lỗi có được ghi nguyên văn, sau đó mới diễn giải tiếng Việt chưa?",
        "Ảnh RViz2/Gazebo có chú thích nguồn dữ liệu và phân biệt đường kế hoạch với quỹ đạo thực thi chưa?",
    ])
    add_para(doc, "Quy ước cuối cùng: nếu một câu có thể khiến người đọc hiểu thành một khối hệ thống khác, một đại lượng khác hoặc một mức bảo đảm rộng hơn dữ liệu, câu đó phải được viết lại trước khi nộp.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
