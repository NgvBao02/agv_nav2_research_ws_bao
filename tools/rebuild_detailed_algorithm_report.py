#!/usr/bin/env python3

"""Rebuild the detailed Vietnamese tutorial report from Linh's 53-page master.

The supplied DOCX is the source of truth for prose, layout, figures, tables and
writing voice.  This script does not recreate or shorten that document.  It
copies the master and then inserts the verified 2026-07-26 controller, hardware
and Gazebo-audit material into the relevant chapters.

The script intentionally depends only on python-docx and the Python standard
library so that the report can be regenerated without a browser renderer.
"""

from __future__ import annotations

import argparse
import gzip
import hashlib
import json
import math
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_MASTER = Path("/home/linh-pham/Downloads/Phong cách Phạm Hải Linh.docx")
DEFAULT_OUTPUT = ROOT / "docs" / "BAO_CAO_TOAN_DIEN_ADAPTIVE_HYBRID_PIVOT_G2.docx"
EXPECTED_MASTER_SHA256 = (
    "ce75e70a7ea961b771bed83081223ce07fed1ac865816c5dab73d1afb8b165f0"
)
AUDIT_DIR = ROOT / "results" / "current_full_audit_20260726"
ASSETS_DIR = ROOT / "docs" / "bao_cao_toan_dien_assets"
CURRENT_FIGURE_MANIFEST = ASSETS_DIR / "current_figure_manifest.json"

ENVIRONMENTS = (
    "research_warehouse",
    "narrow_aisles",
    "office_maze",
    "open_arena",
    "warehouse_cross_aisles",
    "warehouse_dispatch",
    "warehouse_long_aisles",
)

FINAL_TRACES = {
    "research_warehouse": AUDIT_DIR / "lower_left_diagonal_pivot_g2_final.json.gz",
    "narrow_aisles": AUDIT_DIR / "narrow_aisles_pivot_g2_final.json.gz",
    "office_maze": AUDIT_DIR / "office_maze_pivot_g2_final.json.gz",
    "open_arena": AUDIT_DIR / "open_arena_pivot_g2_final.json.gz",
    "warehouse_cross_aisles": AUDIT_DIR / "warehouse_cross_aisles_pivot_g2_final.json.gz",
    "warehouse_dispatch": AUDIT_DIR / "warehouse_dispatch_pivot_g2_final.json.gz",
    "warehouse_long_aisles": AUDIT_DIR / "warehouse_long_aisles_pivot_g2_final.json.gz",
}

BEFORE_AFTER_TRACES = {
    "lower_before": AUDIT_DIR / "lower_left_diagonal_pivot_g2_baseline.json.gz",
    "lower_after": FINAL_TRACES["research_warehouse"],
    "rack_before": AUDIT_DIR / "right_rack_detour_pivot_g2_baseline.json.gz",
    "rack_after": AUDIT_DIR / "right_rack_detour_pivot_g2_final.json.gz",
    "narrow_before_braking": AUDIT_DIR / "narrow_aisles_pivot_g2_optimized.json.gz",
    "narrow_after_braking": FINAL_TRACES["narrow_aisles"],
}

ENV_LABEL = {
    "research_warehouse": "Kho nghiên cứu",
    "narrow_aisles": "Lối đi hẹp",
    "office_maze": "Mê cung văn phòng",
    "open_arena": "Không gian mở",
    "warehouse_cross_aisles": "Kho có lối giao cắt",
    "warehouse_dispatch": "Kho điều phối–xuất hàng",
    "warehouse_long_aisles": "Kho có lối đi dài",
}

MUTATED_MASTER_PARAGRAPH_PREFIXES = (
    "7.7 Thực thi Pivot và terminal servo",
    "Khoảng cách hai tâm vệt lăn vật lý là 0,2548 m",
    "9.2 Cảm biến",
    "9.3 Bảng tham số chính",
    "Scenario đại diện lower_left_diagonal chạy",
    "Scenario đại diện southwest_northeast_weave chạy",
    "Scenario đại diện office_long_diagonal chạy",
    "Scenario đại diện southwest_northeast chạy",
    "Scenario đại diện cross_aisle_transfer chạy",
    "Scenario đại diện full_replenishment chạy",
    "Scenario đại diện diagonal_replenishment chạy",
    "11.2 Định nghĩa metric",
    "12.2 Những lỗi đã sửa được bằng dữ liệu",
    "Bảng này cho thấy việc sửa projection",
    "12.3 So sánh hình học trên 7.200 dòng",
    "12.4 So sánh chạy kín tám phương pháp ở tốc độ thích nghi",
    "12.5 Tương thích với năm planner",
    "12.6 Ca phản ví dụ phải được giữ lại trong báo cáo",
    "Dữ liệu chạy kín: quỹ đạo tham chiếu",
    "Cùng một hình học môi trường trong occupancy grid RViz2",
    "Sai số chạy kín của Pivot–G² thích nghi trên bảy môi trường.",
    "Pure Pivot–G² thích nghi giảm 92,2% Eκ so với Raw.",
    "So sánh hình học: năng lượng độ cong, clearance và tỷ lệ sinh đường thành công.",
    "Mức giảm năng lượng độ cong theo từng môi trường và từng phương pháp.",
    "Ảnh hưởng của smoother và chế độ tốc độ đến thời gian",
    "Điểm riêng mạnh nhất của dự án không phải việc sử dụng Bézier bậc năm",
    "Mô phỏng không thay thế robot thật.",
    "Adaptive Hybrid Pivot–G² giải quyết khoảng trống đó theo một chuỗi kín.",
    "Nguồn dữ liệu, hình ảnh và các kết quả định lượng",
)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def load_json(path: Path) -> dict:
    if not path.exists():
        raise FileNotFoundError(f"Thiếu dữ liệu kiểm chứng: {path}")
    if path.suffix == ".gz":
        with gzip.open(path, "rt", encoding="utf-8") as stream:
            return json.load(stream)
    return json.loads(path.read_text(encoding="utf-8"))


def cm(value_m: float) -> str:
    return f"{100.0 * float(value_m):.3f}".replace(".", ",")


def dec(value: float, digits: int = 3) -> str:
    return f"{float(value):.{digits}f}".replace(".", ",")


def percent_reduction(after: float, before: float) -> float:
    return 100.0 * (float(before) - float(after)) / float(before)


def normalized_text(text: str) -> str:
    return " ".join(text.split())


def find_paragraph(document: Document, exact_text: str):
    expected = normalized_text(exact_text)
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if normalized_text(paragraph.text) == expected
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected exactly one paragraph {exact_text!r}, found {len(matches)}"
        )
    return matches[0]


def find_paragraph_prefix(document: Document, prefix: str):
    expected = normalized_text(prefix)
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if normalized_text(paragraph.text).startswith(expected)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected exactly one paragraph beginning {prefix!r}, found {len(matches)}"
        )
    return matches[0]


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_paragraph_after(document: Document, cursor, text: str, style: str = "Normal"):
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(text)
    cursor_element = cursor._p if hasattr(cursor, "_p") else cursor._tbl
    cursor_element.addnext(paragraph._p)
    return paragraph


def insert_picture_after(
    document: Document,
    cursor,
    image_path: Path,
    *,
    width= Cm(16.0),
):
    if not image_path.exists():
        raise FileNotFoundError(f"Thiếu hình báo cáo: {image_path}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(image_path), width=width)
    cursor_element = cursor._p if hasattr(cursor, "_p") else cursor._tbl
    cursor_element.addnext(paragraph._p)
    return paragraph


def _drawing_paragraph_after(paragraph):
    element = paragraph._p.getnext()
    while element is not None:
        if element.tag == qn("w:p"):
            if element.xpath(".//a:blip"):
                return element
            text = normalized_text("".join(element.itertext()))
            if text:
                style = element.pPr.pStyle if element.pPr is not None else None
                style_id = style.get(qn("w:val")) if style is not None else ""
                if style_id and style_id.startswith("Heading"):
                    break
        element = element.getnext()
    raise RuntimeError(f"Không tìm thấy hình sau đoạn {paragraph.text!r}")


def replace_first_picture_after_heading(
    document: Document,
    heading_text: str,
    image_path: Path,
) -> None:
    if not image_path.exists():
        raise FileNotFoundError(f"Thiếu hình thay thế: {image_path}")
    heading = find_paragraph(document, heading_text)
    drawing_p = _drawing_paragraph_after(heading)
    inline = drawing_p.xpath(".//wp:inline")
    if len(inline) != 1:
        raise RuntimeError(f"Hình sau {heading_text!r} không có đúng một inline")
    extent = inline[0].find(qn("wp:extent"))
    width = Emu(int(extent.get("cx")))
    # Reuse the existing paragraph so all surrounding pagination remains stable.
    for child in list(drawing_p):
        drawing_p.remove(child)
    drawing_p.getparent()  # keep lxml ownership explicit before python-docx wraps it
    from docx.text.paragraph import Paragraph

    paragraph = Paragraph(drawing_p, document._body)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=width)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_table(table, font_size: float = 9.0) -> None:
    table.style = "Table Grid"
    table.autofit = True
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, "D9E2F3")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def insert_table_after(
    document: Document,
    cursor,
    headers: Sequence[str],
    rows: Iterable[Sequence[str]],
    *,
    font_size: float = 9.0,
):
    materialized = [tuple(str(value) for value in row) for row in rows]
    table = document.add_table(rows=1, cols=len(headers))
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = str(value)
    for values in materialized:
        if len(values) != len(headers):
            raise ValueError("Table row length does not match header length")
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    format_table(table, font_size)
    cursor_element = cursor._p if hasattr(cursor, "_p") else cursor._tbl
    cursor_element.addnext(table._tbl)
    return table


def add_sequence(document: Document, cursor, blocks):
    """Insert a sequence of (style, text) or table dictionaries after cursor."""
    for block in blocks:
        if isinstance(block, tuple):
            style, text = block
            cursor = insert_paragraph_after(document, cursor, text, style)
        elif isinstance(block, dict):
            cursor = insert_table_after(
                document,
                cursor,
                block["headers"],
                block["rows"],
                font_size=block.get("font_size", 9.0),
            )
        else:
            raise TypeError(f"Unsupported report block: {type(block)!r}")
    return cursor


def append_table_rows(table, rows: Iterable[Sequence[str]], font_size: float = 9.0) -> None:
    column_count = len(table.columns)
    for values in rows:
        values = tuple(str(value) for value in values)
        if len(values) != column_count:
            raise ValueError("Appended row length does not match table")
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    format_table(table, font_size)


def table_after_heading(document: Document, heading_text: str):
    heading = find_paragraph(document, heading_text)
    element = heading._p.getnext()
    while element is not None:
        if element.tag == qn("w:tbl"):
            for table in document.tables:
                if table._tbl is element:
                    return table
            raise RuntimeError(f"Table object not found after heading {heading_text!r}")
        if element.tag == qn("w:p"):
            style_id = element.pPr.pStyle.get(qn("w:val")) if (
                element.pPr is not None and element.pPr.pStyle is not None
            ) else ""
            if style_id and style_id.startswith("Heading"):
                break
        element = element.getnext()
    raise RuntimeError(f"No table found after heading {heading_text!r}")


def validate_master(document: Document, master_path: Path) -> None:
    actual_hash = sha256(master_path)
    if actual_hash != EXPECTED_MASTER_SHA256:
        raise RuntimeError(
            "File gốc không đúng phiên bản 53 trang đã kiểm kê. "
            f"SHA-256 nhận được: {actual_hash}"
        )
    if len(document.paragraphs) != 401:
        raise RuntimeError(f"Expected 401 master paragraphs, got {len(document.paragraphs)}")
    if len(document.tables) != 25:
        raise RuntimeError(f"Expected 25 master tables, got {len(document.tables)}")
    if len(document.inline_shapes) != 34:
        raise RuntimeError(
            f"Expected 34 master inline figures, got {len(document.inline_shapes)}"
        )
    required = (
        "VII, Profile vận tốc và điều khiển vòng kín",
        "IX, Mô hình robot, cảm biến và các giới hạn đang dùng",
        "XI, Benchmark đo gì và làm thế nào để tránh kết luận sai?",
        "XII, Kết quả hình học và chạy kín",
        "PHỤ LỤC B, Vị trí source chính và tài liệu tham khảo",
    )
    for text in required:
        find_paragraph(document, text)


def master_preservation_snapshot(document: Document) -> dict:
    preserved_paragraphs = []
    for paragraph in document.paragraphs:
        text = normalized_text(paragraph.text)
        if not text:
            continue
        if any(text.startswith(prefix) for prefix in MUTATED_MASTER_PARAGRAPH_PREFIXES):
            continue
        preserved_paragraphs.append(text)
    preserved_tables = []
    # Tables 7, 15, 18, 19, 24 and 25 are intentionally extended, refreshed,
    # or have their metric headers made more explicit.
    mutable_table_indices = {6, 14, 17, 18, 23, 24}
    for index, table in enumerate(document.tables):
        if index in mutable_table_indices:
            continue
        preserved_tables.append(
            tuple(tuple(normalized_text(cell.text) for cell in row.cells) for row in table.rows)
        )
    return {
        "paragraphs": tuple(preserved_paragraphs),
        "tables": tuple(preserved_tables),
    }


def validate_master_preservation(document: Document, snapshot: dict) -> None:
    output_paragraphs = [normalized_text(p.text) for p in document.paragraphs]
    missing_paragraphs = [
        text for text in snapshot["paragraphs"] if text not in output_paragraphs
    ]
    if missing_paragraphs:
        raise RuntimeError(
            "Nội dung gốc bị mất ngoài danh sách cập nhật có chủ đích: "
            f"{missing_paragraphs[:3]}"
        )
    output_tables = {
        tuple(tuple(normalized_text(cell.text) for cell in row.cells) for row in table.rows)
        for table in document.tables
    }
    missing_tables = [table for table in snapshot["tables"] if table not in output_tables]
    if missing_tables:
        raise RuntimeError(
            f"Có {len(missing_tables)} bảng gốc không còn nguyên vẹn"
        )


def validate_audit(final_data: dict[str, dict], audit_data: dict[str, dict]) -> None:
    if set(final_data) != set(ENVIRONMENTS):
        raise RuntimeError("Đợt kiểm chứng hiện tại không đủ bảy môi trường")
    for environment, row in final_data.items():
        if not row.get("success") or not row.get("physically_settled"):
            raise RuntimeError(f"Trace cuối chưa thành công và dừng vật lý: {environment}")
        if row.get("adaptive_speed_nominal_max_abs_jerk_mps3", math.inf) > 0.900001:
            raise RuntimeError(f"Nominal jerk vượt 0,9 m/s³: {environment}")
        for key in (
            "planner_start_anchor_adjustment_m",
            "planner_goal_anchor_adjustment_m",
            "selected_start_anchor_adjustment_m",
            "selected_goal_anchor_adjustment_m",
        ):
            if float(row.get(key, 0.0)) > 0.080001:
                raise RuntimeError(f"Neo path vượt hợp đồng 0,08 m: {environment}/{key}")
    for label, row in audit_data.items():
        if not row.get("success") or not row.get("physically_settled"):
            raise RuntimeError(f"Trace before/after không hợp lệ: {label}")


def validate_current_figure_manifest() -> None:
    if not CURRENT_FIGURE_MANIFEST.exists():
        raise FileNotFoundError(
            "Thiếu manifest hình hiện tại. Chạy "
            "tools/generate_full_algorithm_tutorial_report.py trước."
        )
    manifest = json.loads(CURRENT_FIGURE_MANIFEST.read_text(encoding="utf-8"))
    if manifest.get("version") != "current_full_audit_20260726":
        raise RuntimeError("Manifest hình không đúng phiên bản kiểm chứng 26/07/2026")
    expected_validation = {
        name: path for name, path in FINAL_TRACES.items()
    }
    expected_audit = {
        "lower_before": BEFORE_AFTER_TRACES["lower_before"],
        "lower_after": BEFORE_AFTER_TRACES["lower_after"],
        "rack_before": BEFORE_AFTER_TRACES["rack_before"],
        "rack_after": BEFORE_AFTER_TRACES["rack_after"],
        "narrow_before_angular_braking": BEFORE_AFTER_TRACES[
            "narrow_before_braking"
        ],
        "narrow_after": BEFORE_AFTER_TRACES["narrow_after_braking"],
    }
    for group_name, expected in (
        ("validation_sources", expected_validation),
        ("before_after_sources", expected_audit),
    ):
        recorded = manifest.get(group_name, {})
        if set(recorded) != set(expected):
            raise RuntimeError(f"Manifest hình thiếu nguồn trong {group_name}")
        for name, path in expected.items():
            if recorded[name].get("sha256") != sha256(path):
                raise RuntimeError(
                    f"Hình đã cũ so với dữ liệu nguồn: {group_name}/{name}"
                )
    generated = tuple(manifest.get("generated_figures", ()))
    if len(generated) != 10:
        raise RuntimeError("Manifest phải chứa 10 hình phụ thuộc dữ liệu hiện tại")
    figure_hashes = manifest.get("figure_sha256", {})
    if set(figure_hashes) != set(generated):
        raise RuntimeError("Manifest chưa ghi đủ SHA-256 của hình được sinh")
    for filename in generated:
        path = ASSETS_DIR / filename
        if not path.exists() or path.stat().st_size < 10_000:
            raise RuntimeError(f"Hình hiện tại bị thiếu hoặc rỗng: {path}")
        if figure_hashes[filename] != sha256(path):
            raise RuntimeError(f"Hình không khớp SHA-256 trong manifest: {path}")


def validate_source_contract() -> None:
    """Refuse to publish prose that no longer matches the current source tree."""
    expected_snippets = {
        ROOT / "src/vacuum_robot_gazebo/config/real_robot_profile.yaml": (
            "effective_wheel_separation_m: 0.2834",
            "wheel_separation_multiplier: 1.1122448980",
            "motor_model: GA25_encoder_130rpm",
            "gearbox_ratio: 45.0",
            "nominal_output_rpm: 130.0",
            "rated_load_output_rpm: 100.0",
            "stall_current_per_motor_a: 1.3",
            "topology: 4S4P",
            "total_cell_count: 16",
            "capacity_ah: 10.4",
            "full_voltage_assumption_v: 16.8",
            "regulated_voltage_v: 12.0",
        ),
        ROOT / "src/vacuum_robot_gazebo/config/nav2_params.yaml": (
            "pivot_effective_angular_deceleration: 0.18",
            "initial_alignment_preview_distance: 0.30",
            "initial_alignment_enter_angle: 0.15",
            "initial_alignment_exit_angle: 0.035",
            "adaptive_max_linear_speed: 0.30",
            "adaptive_max_wheel_linear_speed: 0.36",
            "adaptive_max_linear_jerk: 0.90",
        ),
        ROOT / "src/adaptive_pivot_g2_benchmark/adaptive_pivot_g2_benchmark/path_contract.py": (
            "maximum_adjustment: float = 0.08",
            "def anchor_path_start(",
            "def anchor_path_goal(",
        ),
        ROOT / "src/adaptive_pivot_g2_controller/src/maneuver_path.cpp": (
            "double angular_braking_speed_limit(",
            "std::sqrt(2.0 * effective_angular_deceleration * available_angle)",
        ),
        ROOT / "src/adaptive_pivot_g2_controller/src/adaptive_speed_profile.cpp": (
            "const bool zero_speed_override = unconstrained_speed < -kEpsilon;",
            "upper_cap_override || zero_speed_override",
            "result.speed <= kEpsilon ?",
        ),
        ROOT / "src/vacuum_robot_gazebo/urdf/vacuum_robot.urdf": (
            '<mechanicalReduction>45</mechanicalReduction>',
            '<limit effort="0.3530394" velocity="13.6135681656"/>',
        ),
        ROOT / "src/vacuum_robot_gazebo/models/vacuum_robot/model.sdf": (
            "<wheel_separation>0.2834</wheel_separation>",
            "<wheel_radius>0.0425</wheel_radius>",
        ),
        ROOT / "docs/REV_ECIT_2026_ADAPTIVE_HYBRID_PIVOT_G2_SUPPLEMENT.html": (
            "<td>53.094</td><td>38.265</td><td>-27.9%</td>",
            "<td>1.883</td><td>1.412</td><td>-25.0%</td>",
            "<td>2.309</td><td>1.071</td><td>-53.6%</td>",
            "<td>0.0138</td><td>0.0061</td><td>-55.6%</td>",
        ),
    }
    for path, snippets in expected_snippets.items():
        if not path.exists():
            raise FileNotFoundError(f"Thiếu source để đối chiếu báo cáo: {path}")
        source = path.read_text(encoding="utf-8")
        missing = [snippet for snippet in snippets if snippet not in source]
        if missing:
            raise RuntimeError(f"Source contract changed in {path}: missing {missing}")


def insert_revision_note(document: Document) -> None:
    cursor = find_paragraph_prefix(
        document,
        "Tài liệu này được viết theo đúng mạch phát triển của dự án",
    )
    add_sequence(
        document,
        cursor,
        [
            (
                "Normal",
                "Bản đang đọc là bản rà soát toàn diện ngày 26/07/2026. "
                "Toàn bộ nội dung của bản gốc được giữ nguyên để không làm mất "
                "mạch suy luận đã hình thành trong quá trình nghiên cứu. Những "
                "phần được bổ sung tập trung vào bốn vấn đề đã bộc lộ khi quan "
                "sát đồng thời RViz2, Gazebo và dữ liệu trace: neo chính xác "
                "start–goal của path, xác định hướng xuất phát, phanh khi quay, "
                "và hợp đồng phần cứng của hai động cơ GA25 cùng bộ nguồn 4S4P. "
                "Kết quả cũ ngày 25/07 vẫn được giữ như dấu vết phát triển; kết "
                "quả hiện tại được ghi riêng để người đọc không nhầm hai phiên bản."
            ),
        ],
    )


def insert_controller_update(document: Document) -> None:
    old_heading = find_paragraph(document, "7.7 Thực thi Pivot và terminal servo")
    set_paragraph_text(old_heading, "7.11 Thực thi Pivot và terminal servo")
    anchor = find_paragraph_prefix(
        document,
        "Nếu một safety cap đột ngột thấp hơn vận tốc jerk-limited",
    )
    blocks = [
        (
            "Heading 2",
            "7.7 Neo start và goal liên tục trước khi đánh giá hoặc làm mượt",
        ),
        (
            "Normal",
            "Một lỗi nhỏ về biểu diễn có thể làm sai toàn bộ quá trình đánh giá. "
            "Global planner làm việc trên occupancy grid nên pose đầu hoặc cuối "
            "của nav_msgs/Path có thể bị đưa về tâm ô lưới gần nhất, trong khi "
            "pose thực của robot và goal người dùng chọn là tọa độ liên tục trong "
            "frame map. Nếu giữ nguyên tâm ô như thể đó là start thật, controller "
            "sẽ nhìn thấy một sai số ngang giả ngay từ lúc bắt đầu. Khi robot vừa "
            "ra khỏi đường cong hoặc đang căn goal, sai số giả này còn có thể làm "
            "profile vận tốc phanh hoặc tăng tốc sai thời điểm."
        ),
        ("Equation", "δgrid,max = √[(rmap/2)² + (rmap/2)²] = rmap/√2"),
        (
            "Normal",
            "Trong đó: δgrid,max là độ lệch lớn nhất giữa một điểm liên tục và "
            "tâm ô chứa điểm đó; rmap là độ phân giải map, hiện bằng 0,05 m/ô."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: với rmap = 0,05 m, độ lệch cực đại chỉ "
            "do lượng tử hóa đã bằng 0,035355 m, tức khoảng 3,54 cm. Con số này "
            "lớn hơn sai số bám trung bình mà hệ thống đang cố giảm, vì vậy không "
            "thể xem nó là nhiễu không đáng kể."
        ),
        (
            "Normal",
            "Hợp đồng path mới thực hiện hai bước. Trước hết kiểm tra khoảng "
            "cách từ pose đầu của planner đến start liên tục và từ pose cuối đến "
            "goal liên tục. Nếu khoảng điều chỉnh không vượt 0,08 m, vị trí đầu "
            "và cuối được phục hồi đúng theo yêu cầu gốc. Orientation ở hai pose "
            "biên được lấy từ start/goal yêu cầu; controller không dùng riêng "
            "orientation đầu để suy ra tiếp tuyến mà tính lại bằng preview của "
            "selected path. Nếu phải điều chỉnh "
            "lớn hơn 0,08 m, trial bị loại thay vì âm thầm kéo path qua một vùng "
            "mà planner chưa chứng minh là an toàn."
        ),
        ("Equation", "p₀ ← pstart,  pN ← pgoal,  nếu ‖p₀ - pstart‖ ≤ 0,08 và ‖pN - pgoal‖ ≤ 0,08"),
        (
            "Normal",
            "Trong đó: p₀ và pN là hai pose biên của path; pstart là pose liên "
            "tục tại thời điểm lập đường; pgoal là pose đích do scenario hoặc "
            "RViz2 cung cấp; ‖·‖ là khoảng cách Euclidean trên mặt phẳng."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: giới hạn 0,08 m đủ lớn để sửa lượng tử hóa "
            "của map 0,05 m nhưng vẫn đủ nhỏ để phát hiện một path bị lệch frame, "
            "bị dùng nhầm start, hoặc bị ghép sai sau smoother."
        ),
        (
            "Normal",
            "Bước neo được áp dụng cho cả raw path do planner trả về và selected "
            "path sau khi smoother lựa chọn. Bốn đại lượng "
            "planner_start_anchor_adjustment_m, planner_goal_anchor_adjustment_m, "
            "selected_start_anchor_adjustment_m và "
            "selected_goal_anchor_adjustment_m được ghi vào kết quả. Cách làm này "
            "bảo đảm tám smoother được chấm trên cùng một miền start–goal, đồng "
            "thời không để chính bước làm mượt âm thầm di chuyển hai điều kiện biên."
        ),
        (
            "Heading 2",
            "7.8 Xác định hướng xuất phát theo hai tầng",
        ),
        (
            "Normal",
            "Hướng đặt robot ban đầu không nên được lấy đơn giản bằng đường thẳng "
            "nối start đến goal. Trong map có kệ hoặc tường, bearing trực tiếp có "
            "thể chỉ vào vật cản dù đường hợp lệ phải rẽ sang một hành lang khác. "
            "Nếu Gazebo spawn robot theo bearing sai, controller sẽ mất nhiều thời "
            "gian quay ngược; nếu vừa quay vừa tiến, bánh dễ kéo xe lệch khỏi ray "
            "tham chiếu ngay trước góc đầu tiên."
        ),
        ("Equation", "ψdirect = atan2(ygoal - ystart, xgoal - xstart)"),
        (
            "Normal",
            "Trong đó: ψdirect là hướng trực tiếp từ start đến goal; xstart, "
            "ystart là tọa độ start; xgoal, ygoal là tọa độ goal; atan2 giữ đúng "
            "góc phần tư và trả yaw trong radian."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: ψdirect chỉ là ứng viên đầu tiên, không "
            "phải đáp án bắt buộc. Nó được chấp nhận khi hướng đó có khoảng thăm "
            "dò footprint an toàn hoặc khi line-of-sight trên occupancy grid không "
            "đi qua ô chiếm dụng."
        ),
        (
            "Normal",
            "Tầng thứ nhất diễn ra trước khi spawn hoặc trước khi gửi trial. File "
            "YAML của map cung cấp resolution, origin và đường dẫn PGM. Tọa độ "
            "world được đổi về tọa độ cục bộ của map, sau đó lấy chỉ số cột và "
            "hàng bằng phép floor. Thuật toán không chỉ kiểm tra một tia từ tâm; "
            "nó thăm dò theo các hướng ứng viên với bề rộng footprint để tránh "
            "chọn một hướng mà tâm đi qua được nhưng góc thân xe chạm kệ."
        ),
        ("Equation", "c = floor(xlocal/rmap),  r = H - 1 - floor(ylocal/rmap)"),
        (
            "Normal",
            "Trong đó: c và r là chỉ số cột, hàng của ảnh PGM; H là chiều cao "
            "ảnh; xlocal, ylocal là tọa độ sau khi trừ origin và quay về hệ map; "
            "phép H - 1 - · xuất hiện vì trục hàng của ảnh tăng từ trên xuống, "
            "trong khi trục y của bản đồ tăng từ dưới lên."
        ),
        (
            "Normal",
            "Nếu bearing trực tiếp không an toàn, tầng này ước lượng một tuyến "
            "grid route có ưu tiên clearance và lấy hướng của đoạn đầu đủ dài. "
            "Tên nguồn quyết định, ví dụ map_aware_grid_route hoặc direct_bearing, "
            "được ghi vào initial_heading_source. Khi scenario đã cung cấp yaw "
            "tường minh, yaw đó được giữ vì nó là một phần của điều kiện thử, "
            "không bị thuật toán tự ý thay thế."
        ),
        (
            "Normal",
            "Tầng thứ hai diễn ra sau khi planner và smoother đã tạo selected "
            "path. Controller không tin tuyệt đối orientation ở pose đầu, vì một "
            "planner 2D có thể để quaternion mặc định hoặc orientation chưa phản "
            "ánh đúng tiếp tuyến. Nó lấy một điểm preview cách đầu đường 0,30 m "
            "theo chiều dài cung rồi tính hướng dây cung từ pose đầu đến điểm đó."
        ),
        ("Equation", "ψpreview = atan2(ypreview - y₀, xpreview - x₀),  spreview = 0,30 m"),
        (
            "Normal",
            "Trong đó: ψpreview là hướng controller cần căn; p₀ = (x₀,y₀) là "
            "đầu selected path; ppreview là điểm nội suy tại tiến độ 0,30 m. "
            "Khoảng preview đủ dài để không bị chi phối bởi hai waypoint gần như "
            "trùng nhau nhưng vẫn ngắn hơn phần lớn đoạn vào hành lang."
        ),
        (
            "Normal",
            "Nếu |wrap(ψpreview - ψrobot)| nhỏ hơn 0,15 rad, robot được phép vào "
            "bám đường ngay. Nếu lớn hơn hoặc bằng ngưỡng này, trạng thái "
            "initial_alignment được kích hoạt: vận tốc tuyến tính bằng 0 và robot "
            "quay tại chỗ. Trạng thái chỉ được nhả khi sai số nhỏ hơn 0,035 rad, "
            "đồng thời vận tốc tuyến tính và góc đều đã nằm trong ngưỡng dừng. "
            "Hai ngưỡng khác nhau tạo hysteresis, tránh trạng thái bật–tắt liên "
            "tục quanh một giá trị duy nhất."
        ),
        (
            "Heading 2",
            "7.9 Bao phanh góc theo khả năng giảm tốc thực",
        ),
        (
            "Normal",
            "Giới hạn gia tốc lệnh 1,2 rad/s² chỉ quy định mỗi chu kỳ được thay "
            "đổi lệnh bao nhiêu; nó không chứng minh thân xe thực sự giảm vận tốc "
            "góc với đúng giá trị đó. Trace Gazebo cho thấy khi phải đảo chiều "
            "quay trong initial alignment, mức giảm tốc góc hiệu dụng chỉ khoảng "
            "0,18 rad/s². Nếu tiếp tục dùng một lệnh tốc độ góc gần hằng số đến "
            "sát target rồi mới hạ, robot sẽ đi qua hướng cần đạt và quay ngược "
            "lại. Quá trình này lặp thành counter-rotation, làm mất thời gian và "
            "có thể kéo robot lệch đường khi chuyển sang vận tốc tuyến tính."
        ),
        ("Equation", "ωbrake = min[ωmax, √(2αeff·max(|eψ| - ψtol, 0))]"),
        (
            "Normal",
            "Trong đó: ωbrake là trần vận tốc góc còn được phép phát; ωmax = "
            "0,70 rad/s là trần quay; αeff = 0,18 rad/s² là giảm tốc góc hiệu "
            "dụng đo từ Gazebo; eψ là sai số heading có dấu; ψtol = 0,015 rad là "
            "deadband của Pivot."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: công thức được suy ra từ quan hệ động học "
            "ω² = 2αθ. Phần |eψ| - ψtol là góc thật sự còn được dùng để phanh; "
            "deadband được trừ trước để robot không cố dừng tại một điểm toán học "
            "không có bề rộng."
        ),
        ("Equation", "ωdes = sign(eψ)·min(Kψ|eψ|, ωbrake)"),
        (
            "Normal",
            "Trong đó: ωdes là vận tốc góc mong muốn; Kψ là hệ số phản hồi theo "
            "sai số; sign(eψ) chọn chiều quay. Sau đó lệnh thực vẫn bị giới hạn "
            "bởi |ωcmd - ωmeasured| ≤ αcmdΔt với αcmd = 1,2 rad/s²."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: luật tách ba vai trò. Kψ tạo phản hồi gần "
            "target; bao căn bậc hai quyết định lúc phải phanh theo động học; "
            "αcmd giới hạn độ dốc của lệnh theo chu kỳ. Không nên gộp ba đại "
            "lượng này thành một tham số vì chúng mô tả ba hiện tượng khác nhau."
        ),
        (
            "Normal",
            "Ví dụ với |eψ| = 0,961 rad, ψtol = 0,015 rad và αeff = 0,18 "
            "rad/s², trần phanh bằng khoảng 0,584 rad/s, thấp hơn ωmax = 0,70 "
            "rad/s. Robot bắt đầu giảm tốc sớm thay vì giữ 0,70 rad/s đến gần "
            "đích. Trên narrow_aisles, số mẫu initial_alignment giảm từ 209 xuống "
            "93 và thời gian hoàn thành giảm từ 74,721 s xuống 67,104 s."
        ),
        (
            "Heading 2",
            "7.10 Biên vận tốc bằng không và cách báo jerk trung thực",
        ),
        (
            "Normal",
            "Bộ tạo S-curve hoạt động tốt trong miền v > 0, nhưng tại thời điểm "
            "dừng có một biên khả thi đặc biệt. Nếu gia tốc trước đang âm và tốc "
            "độ còn lại quá nhỏ, phép tích phân vprev + acmdΔt có thể cho kết quả "
            "âm trước khi jerk kịp kéo gia tốc về 0. Robot không được phép chạy "
            "với độ lớn vận tốc âm trong nhánh chỉ chuyển động tiến, vì vậy phải "
            "clip tại 0. Mẫu clip này là một safety override thật, không phải một "
            "mẫu nominal jerk."
        ),
        ("Equation", "vunconstrained = vprev + acmdΔt"),
        ("Equation", "vshape = max(0, vunconstrained)"),
        (
            "Normal",
            "Trong đó: vunconstrained là kết quả tích phân trước ràng buộc; "
            "vshape là vận tốc không âm; acmd là gia tốc sau giới hạn jerk."
        ),
        (
            "Normal",
            "Nếu vunconstrained < 0, cờ zero_speed_override được bật. Nếu một "
            "trần an toàn mới làm vtarget nhỏ hơn vshape, cờ upper_cap_override "
            "được bật. Hai trường hợp đều được gộp vào safety_override. Jerk đo "
            "từ chênh lệch gia tốc vẫn được lưu nguyên giá trị để dữ liệu trung "
            "thực, nhưng không được tính vào adaptive_speed_nominal_max_abs_jerk."
        ),
        ("Equation", "safety_override = upper_cap_override ∨ zero_speed_override"),
        (
            "Normal",
            "Trong đó: ký hiệu ∨ nghĩa là phép OR logic. Chỉ cần một điều kiện "
            "an toàn xảy ra thì mẫu đó không còn thuộc miền S-curve danh nghĩa."
        ),
        (
            "Normal",
            "Khi vshape đã về 0, trạng thái gia tốc nội bộ được reset về 0. Nếu "
            "giữ lại một gia tốc âm ngoài miền hợp lệ, chu kỳ sau sẽ tưởng robot "
            "vẫn đang phanh và có thể tạo một xung jerk giả khi tăng tốc lại. "
            "Nhờ phân loại này, giới hạn nominal jerk 0,9 m/s³ được kiểm tra đúng "
            "ngữ nghĩa, còn các clip vì an toàn vẫn hiện rõ trong trace thay vì "
            "bị che đi bởi một thống kê đã lọc."
        ),
    ]
    add_sequence(document, anchor, blocks)


def insert_hardware_update(document: Document) -> None:
    old_sensor = find_paragraph(document, "9.2 Cảm biến")
    old_parameters = find_paragraph(document, "9.3 Bảng tham số chính")
    set_paragraph_text(old_sensor, "9.7 Cảm biến")
    set_paragraph_text(old_parameters, "9.8 Bảng tham số chính")

    calibration = find_paragraph_prefix(
        document,
        "Khoảng cách hai tâm vệt lăn vật lý là 0,2548 m",
    )
    set_paragraph_text(
        calibration,
        "Khoảng cách hai tâm vệt lăn vật lý theo CAD là 0,2548 m và được dùng "
        "trong động học, giới hạn vận tốc bánh và profile vận tốc. Plugin "
        "DiffDrive trong Gazebo dùng khoảng cách hiệu dụng 0,2834 m, tương ứng "
        "multiplier 1,1122448980, để odometry khớp ground truth của contact "
        "model. Giá trị 0,2834 m là kết quả làm tròn của fit bình phương tối "
        "thiểu có trọng số 0,283385 m từ hai trial độc lập ngày 26/07/2026. Hai "
        "giá trị không được trộn lẫn: 0,2548 m mô tả hình học thật, còn 0,2834 m "
        "chỉ hiệu chỉnh quan hệ bánh–thân xe trong mô phỏng."
    )

    motor_rows = [
        ("Số lượng và loại", "2 × GA25 encoder 130 rpm", "Hai bánh chủ động độc lập"),
        ("Điện áp định mức", "12 V DC", "Điện áp rail motor sau bộ hạ áp"),
        ("Chiều quay", "CW/CCW", "Driver phải điều khiển được hai chiều"),
        ("Tỷ số truyền", "45:1", "45 vòng armature xấp xỉ 1 vòng trục ra"),
        ("Tốc độ armature", "6000 rpm", "Tốc độ phần motor trước hộp số"),
        ("Không tải", "130 ±10% rpm; 60 mA/motor", "Giới hạn khi gần như không mang tải"),
        ("Tải định mức", "100 ±10% rpm; 300 mA/motor", "Điểm làm việc danh nghĩa"),
        ("Mô-men định mức", "1 kgf·cm = 0,0980665 N·m", "Mức tải dùng liên tục theo dữ liệu cung cấp"),
        ("Stall", "1,3 A; 3,6 kgf·cm = 0,3530394 N·m", "Giới hạn tuyệt đối, không được giữ lâu"),
        ("Kích thước", "dài 68 mm; đường kính 25 mm", "Kích thước thân motor trong mô hình"),
    ]
    battery_rows = [
        ("Tổng số cell", "N = S·P = 4·4 = 16", "Đúng cấu hình 4S4P"),
        ("Dung lượng một cell", "2600 mAh = 2,6 Ah", "Dung lượng điện tích danh nghĩa"),
        ("Dòng xả ghi trên cell", "5C = 13 A", "13 A = 5 × 2,6 Ah"),
        ("Điện áp pack danh nghĩa", "4 × 3,7 = 14,8 V", "Giả định Li-ion chuẩn, cần đối chiếu cell thật"),
        ("Điện áp pack đầy", "4 × 4,2 = 16,8 V", "Không được cấp thẳng vào GA25 12 V"),
        ("Dung lượng pack", "4 × 2,6 = 10,4 Ah", "Các nhánh song song cộng dung lượng"),
        ("Năng lượng danh nghĩa", "14,8 × 10,4 = 153,92 Wh", "Ước lượng năng lượng lý thuyết"),
        ("Dòng xả lý thuyết của cell-array", "4 × 13 = 52 A", "Không phải dòng an toàn mặc định của toàn xe"),
        ("Rail motor", "12 V đã điều áp", "Bắt buộc dùng buck phù hợp trước driver"),
    ]
    blocks = [
        (
            "Heading 2",
            "9.2 Đọc đúng thông số của động cơ GA25",
        ),
        (
            "Normal",
            "Mỗi bên xe dùng một động cơ giảm tốc GA25 có encoder, điện áp định "
            "mức 12 V DC và tỷ số truyền 45:1. Cụm từ 130 rpm trong tên thường "
            "chỉ tốc độ trục ra ở chế độ không tải danh nghĩa, không có nghĩa "
            "bánh xe luôn chạy 130 vòng/phút khi mang toàn bộ khối lượng robot. "
            "Để đưa thông số thương mại vào URDF, SDF và controller mà không làm "
            "sai bản chất, cần tách rõ điện áp, dòng, tốc độ, tỷ số truyền và "
            "mô-men."
        ),
        {
            "headers": ("Thông số GA25", "Giá trị", "Ý nghĩa và lý do sử dụng"),
            "rows": motor_rows,
            "font_size": 8.5,
        },
        (
            "Normal",
            "Voltage, ký hiệu U và đơn vị volt (V), là hiệu điện thế cấp cho "
            "motor. Rated voltage nghĩa là điện áp định mức mà nhà sản xuất dùng "
            "để công bố các điểm làm việc. Current, ký hiệu I và đơn vị ampere "
            "(A), là dòng điện motor lấy từ nguồn. No-load current là dòng khi "
            "gần như không có tải cơ; rated-load current là dòng tại tải danh "
            "nghĩa; stall current là dòng khi rotor bị khóa. Ba giá trị không "
            "được thay thế cho nhau trong thiết kế driver hoặc dây nguồn."
        ),
        (
            "Normal",
            "Speed có đơn vị rpm, viết đầy đủ là revolutions per minute, tức số "
            "vòng trong một phút. ROS và Gazebo thường dùng radian trên giây nên "
            "phải đổi đơn vị. Một vòng bằng 2π radian và một phút bằng 60 giây."
        ),
        ("Equation", "ω = 2πn/60"),
        (
            "Normal",
            "Trong đó: n là tốc độ trục ra theo rpm; ω là vận tốc góc theo "
            "rad/s; π xấp xỉ 3,14159265."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: 130 rpm tương ứng 13,613568 rad/s; "
            "100 rpm tương ứng 10,471976 rad/s. Giá trị 13,613568 rad/s được "
            "giữ trong URDF/SDF như giới hạn vận tốc tuyệt đối của joint, còn "
            "100 rpm phản ánh vùng tải định mức thực tế hơn."
        ),
        (
            "Normal",
            "Torque là mô-men xoắn, đo khả năng motor tạo lực quay. Dữ liệu đầu "
            "vào dùng kgf·cm, trong khi URDF dùng N·m. Một kgf là lực do khối "
            "lượng 1 kg chịu gia tốc trọng trường tiêu chuẩn 9,80665 m/s²; cánh "
            "tay đòn 1 cm bằng 0,01 m."
        ),
        ("Equation", "τ[N·m] = τ[kgf·cm] × 9,80665 × 0,01 = τ[kgf·cm] × 0,0980665"),
        (
            "Normal",
            "Trong đó: τ là mô-men xoắn. Do đó 1 kgf·cm bằng 0,0980665 N·m và "
            "3,6 kgf·cm bằng 0,3530394 N·m."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: mô-men 0,0980665 N·m là điểm định mức; "
            "0,3530394 N·m chỉ là stall torque. Không được dùng stall torque như "
            "mô-men làm việc liên tục, vì tại đó mỗi motor hút 1,3 A nhưng trục "
            "không quay, công suất điện chủ yếu biến thành nhiệt trong cuộn dây."
        ),
        (
            "Heading 2",
            "9.3 Từ tốc độ trục motor đến tốc độ tuyến tính của robot",
        ),
        (
            "Normal",
            "Bánh xe có đường kính 0,085 m nên bán kính r = 0,0425 m. Nếu bỏ "
            "qua biến dạng lốp và trượt tiếp xúc, vận tốc tuyến tính tại vành "
            "bánh bằng tích của bán kính và vận tốc góc trục ra."
        ),
        ("Equation", "vwheel = rwheel·ωoutput = rwheel·2πnoutput/60"),
        (
            "Normal",
            "Trong đó: vwheel là vận tốc tiếp tuyến bánh; rwheel = 0,0425 m; "
            "ωoutput là vận tốc góc trục ra; noutput là tốc độ rpm sau hộp số."
        ),
        (
            "Normal",
            "Với 130 rpm, tốc độ tuyến tính lý thuyết không tải bằng 0,578577 "
            "m/s. Với 100 rpm, tốc độ lý thuyết tại tải định mức bằng 0,445059 "
            "m/s. Controller chỉ cho thân xe tối đa 0,30 m/s và profile giới hạn "
            "tốc độ bánh ở 0,36 m/s. Như vậy lệnh điều khiển nằm dưới tốc độ tải "
            "định mức lý thuyết, tạo khoảng dự trữ cho quay vi sai, sai số điện "
            "áp, tải hàng và tổn hao cơ khí."
        ),
        (
            "Normal",
            "Khi robot quay, một bánh có thể nhanh hơn vận tốc tâm thân xe. Với "
            "vận tốc thân v và độ cong κ, vận tốc hai bánh được xấp xỉ bởi "
            "vR = v(1 + Lκ/2) và vL = v(1 - Lκ/2). Vì vậy kiểm tra v ≤ 0,30 "
            "m/s chưa đủ; profile phải tiếp tục kiểm tra max(|vR|,|vL|) ≤ 0,36 "
            "m/s ở từng mẫu."
        ),
        (
            "Heading 2",
            "9.4 Cách biểu diễn GA25 trong URDF, SDF và transmission",
        ),
        (
            "Normal",
            "URDF mô tả cấu trúc link–joint, hình học, khối lượng và giới hạn "
            "joint để ROS biết robot gồm những phần nào. SDF mô tả mô hình được "
            "Gazebo tích phân trong physics. Hai file cùng chứa thân motor GA25 "
            "dài 68 mm, đường kính 25 mm dưới dạng visual. Các motor visual "
            "không có collision và inertial riêng vì khối lượng motor đã nằm "
            "trong tham số chassis; nếu cộng thêm lần nữa sẽ làm sai tổng khối "
            "lượng và mô-men quán tính."
        ),
        (
            "Normal",
            "Hai transmission loại SimpleTransmission dùng mechanical reduction "
            "45:1 để ghi rõ quan hệ armature–trục bánh. Joint bánh giữ giới hạn "
            "vận tốc 13,613568 rad/s và effort 0,3530394 N·m như hard absolute "
            "limits. Đây là giới hạn để chặn cấu hình phi vật lý, không phải lệnh "
            "mặc định mà controller cố đạt."
        ),
        (
            "Normal",
            "SDF DiffDrive nhận lệnh vận tốc thân xe và tạo chuyển động bánh theo "
            "contact model. Vì plugin mô phỏng không phải mô hình điện chi tiết "
            "của GA25, các số dòng 60 mA, 300 mA và 1,3 A hiện được lưu trong "
            "hardware profile chứ chưa được dùng để tính sụt áp hoặc nhiệt. Điều "
            "này phải được ghi rõ để người đọc không nhầm việc đã thêm thông số "
            "vào URDF với việc đã mô phỏng đầy đủ mạch điện và động cơ DC."
        ),
        (
            "Heading 2",
            "9.5 Encoder: những gì đã biết và những gì tuyệt đối không được đoán",
        ),
        (
            "Normal",
            "Tên sản phẩm có chữ encoder nhưng dữ liệu được cung cấp chưa nêu "
            "encoder nằm ở trục armature hay trục ra, số xung mỗi vòng là bao "
            "nhiêu, và driver đếm cạnh x1, x2 hay x4. Vì vậy các trường "
            "pulses_per_revolution, quadrature_decode, encoder_ticks_per_rev, "
            "radians_per_tick và metres_per_tick đang được để null trong "
            "real_robot_profile.yaml. Giá trị null ở đây là một ràng buộc an "
            "toàn dữ liệu, không phải thiếu sót cần lấp bằng một con số lấy từ "
            "motor GA25 khác."
        ),
        ("Equation", "Ntick,out = PPR·q·G  (nếu encoder đặt trước hộp số)"),
        ("Equation", "Δφwheel = 2π/Ntick,out,  Δswheel = rwheel·Δφwheel"),
        (
            "Normal",
            "Trong đó: PPR là số pulse mỗi vòng theo cách nhà sản xuất định "
            "nghĩa; q là hệ số giải mã quadrature, có thể là 1, 2 hoặc 4; G = "
            "45 là tỷ số truyền; Ntick,out là số count cho một vòng trục ra; "
            "Δφwheel là góc bánh trên mỗi count; Δswheel là quãng đường lý "
            "thuyết trên mỗi count."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: chỉ được nhân G khi encoder nằm ở phía "
            "armature. Nếu encoder đã nằm ở trục ra thì nhân thêm 45 sẽ làm sai "
            "odometry đúng 45 lần. Trước khi viết driver robot thật, cần đánh "
            "dấu một cạnh bánh, quay đúng một vòng trục ra, đọc count thô trên "
            "cả hai kênh và ghi rõ chế độ giải mã. Sau đó mới tính metres_per_tick "
            "và hiệu chuẩn bán kính bánh bằng quãng đường đo thực."
        ),
        (
            "Heading 2",
            "9.6 Bộ nguồn 16 cell 18650 mắc 4S4P",
        ),
        (
            "Normal",
            "Ký hiệu 4S4P có nghĩa bốn cell nối tiếp trong mỗi nhánh và bốn "
            "nhánh như vậy nối song song. Nối tiếp làm tăng điện áp nhưng không "
            "tăng dung lượng Ah của một nhánh; song song làm tăng dung lượng và "
            "khả năng cấp dòng nhưng không tăng điện áp. Tổng số cell bằng tích "
            "của số cell nối tiếp và song song."
        ),
        {
            "headers": ("Đại lượng nguồn", "Phép tính", "Ý nghĩa và lý do sử dụng"),
            "rows": battery_rows,
            "font_size": 8.5,
        },
        ("Equation", "Ncell = S·P = 4·4 = 16"),
        ("Equation", "Upack = S·Ucell,  Qpack = P·Qcell"),
        ("Equation", "Enom ≈ Unom·Qpack = 14,8·10,4 = 153,92 Wh"),
        (
            "Normal",
            "Trong đó: S = 4 là số cell nối tiếp; P = 4 là số nhánh song song; "
            "Ucell là điện áp một cell; Qcell là dung lượng một cell; Upack và "
            "Qpack là điện áp, dung lượng của pack; Enom là năng lượng danh nghĩa."
        ),
        (
            "Normal",
            "Các giá trị 3,7 V danh nghĩa và 4,2 V khi đầy là giả định hóa học "
            "Li-ion tiêu chuẩn, không phải thông số đã được xác minh từ mã cell "
            "cụ thể. Với giả định đó, pack là 14,8 V danh nghĩa và 16,8 V khi "
            "đầy. Vì GA25 định mức 12 V, pack 4S đầy tuyệt đối không được cấp "
            "trực tiếp vào motor. Cần một buck converter tạo rail 12 V ổn định "
            "trước motor driver."
        ),
        ("Equation", "Icell = C·Qcell = 5·2,6 = 13 A,  Iarray,ideal = P·Icell = 52 A"),
        (
            "Normal",
            "Trong đó: C = 5 là C-rate; Qcell = 2,6 Ah; Icell = 13 A là dòng xả "
            "được ghi cho một cell; Iarray,ideal = 52 A là tổng lý thuyết của "
            "bốn nhánh song song."
        ),
        (
            "Normal",
            "Ý nghĩa và lý do sử dụng: 52 A không phải dòng mà xe sẽ luôn lấy và "
            "cũng không tự động là dòng an toàn của hệ thống. Giới hạn an toàn "
            "thực bằng phần tử có rating thấp nhất trong chuỗi cell, mối hàn, "
            "BMS 4S, cầu chì, dây, connector, buck, PCB, motor driver và khả năng "
            "tản nhiệt. Model BMS, fuse, buck và driver hiện chưa được cung cấp "
            "nên các rating tương ứng phải tiếp tục để null."
        ),
        (
            "Normal",
            "Hai motor lấy tổng khoảng 0,12 A ở không tải, 0,60 A tại tải định "
            "mức và có thể đạt 2,60 A nếu cả hai cùng stall. Stall vẫn phải bị "
            "ngăn bằng giới hạn dòng, timeout lệnh 0,25 s, phát hiện bánh không "
            "quay và bảo vệ nhiệt; việc pack có thể cấp dòng lớn hơn không làm "
            "stall trở thành một chế độ vận hành hợp lệ."
        ),
    ]
    add_sequence(document, calibration, blocks)

    parameter_table = table_after_heading(document, "9.8 Bảng tham số chính")
    append_table_rows(
        parameter_table,
        [
            ("Khoảng vệt lăn hiệu dụng Gazebo", "0,2834 m", "Hiệu chuẩn contact/odometry, không thay CAD"),
            ("GA25", "2 × 12 V, 45:1, 130 rpm", "Nguồn động lực hai bánh"),
            ("Giới hạn joint tuyệt đối", "13,613568 rad/s; 0,3530394 N·m", "No-load speed và stall torque"),
            ("Tốc độ tải định mức lý thuyết", "0,445059 m/s", "Mốc vật lý cao hơn controller 0,30 m/s"),
            ("Nguồn", "16 × 18650, 4S4P, 10,4 Ah", "Pack 14,8 V danh nghĩa; 16,8 V đầy"),
            ("Rail motor", "12 V qua buck", "Không cấp thẳng 4S đầy vào GA25"),
            ("Encoder PPR", "Chưa biết, để null", "Không suy đoán sai odometry robot thật"),
        ],
        9.0,
    )


def update_map_results(document: Document, final_data: dict[str, dict]) -> None:
    prefix_by_environment = {
        "research_warehouse": "Scenario đại diện lower_left_diagonal chạy",
        "narrow_aisles": "Scenario đại diện southwest_northeast_weave chạy",
        "office_maze": "Scenario đại diện office_long_diagonal chạy",
        "open_arena": "Scenario đại diện southwest_northeast chạy",
        "warehouse_cross_aisles": "Scenario đại diện cross_aisle_transfer chạy",
        "warehouse_dispatch": "Scenario đại diện full_replenishment chạy",
        "warehouse_long_aisles": "Scenario đại diện diagonal_replenishment chạy",
    }
    for environment, prefix in prefix_by_environment.items():
        row = final_data[environment]
        paragraph = find_paragraph_prefix(document, prefix)
        set_paragraph_text(
            paragraph,
            f"Scenario đại diện {row['scenario']} chạy với ThetaStar, "
            "Pivot–G² thích nghi và profile vận tốc hiện tại. "
            f"Thời gian hoàn thành là {dec(row['execution_time_s'], 3)} s, "
            f"ground-truth RMSE là {cm(row['tracking_rmse_m'])} cm, "
            f"sai số cực đại là {cm(row['tracking_max_error_m'])} cm, "
            f"exit RMSE là {cm(row['curve_exit_tracking_rmse_m'])} cm và "
            f"clearance footprint kế hoạch là "
            f"{cm(row['planned_footprint_clearance_min_m'])} cm. "
            "Đây là trace cuối ngày 26/07/2026, đã đạt action success, đạt "
            "ground-truth goal và dừng vật lý; số liệu không chỉ phản ánh "
            "smoother mà còn chứa ảnh hưởng của định vị, controller và physics Gazebo."
        )

    result_table = table_after_heading(
        document, "12.1 Kết quả đại diện trên bảy môi trường"
    )
    for table_row, environment in zip(result_table.rows[1:], ENVIRONMENTS):
        row = final_data[environment]
        values = (
            ENV_LABEL[environment],
            row["scenario"],
            dec(row["execution_time_s"], 3),
            cm(row["tracking_rmse_m"]),
            cm(row["tracking_max_error_m"]),
            cm(row["estimated_tracking_rmse_m"]),
            cm(row["localization_position_error_p95_m"]),
            cm(row["planned_footprint_clearance_min_m"]),
        )
        for cell, value in zip(table_row.cells, values):
            cell.text = value
    format_table(result_table, 7.5)


def replace_data_dependent_figures(document: Document) -> None:
    replacements = {
        "7.11 Thực thi Pivot và terminal servo": ASSETS_DIR / "figure_10_speed_trace.png",
        "10.2 Kho nghiên cứu - research_warehouse": ASSETS_DIR / "map_research_warehouse.png",
        "10.3 Lối đi hẹp - narrow_aisles": ASSETS_DIR / "map_narrow_aisles.png",
        "10.4 Mê cung văn phòng - office_maze": ASSETS_DIR / "map_office_maze.png",
        "10.5 Không gian mở - open_arena": ASSETS_DIR / "map_open_arena.png",
        "10.6 Kho có lối giao cắt - warehouse_cross_aisles": ASSETS_DIR / "map_warehouse_cross_aisles.png",
        "10.7 Kho điều phối–xuất hàng - warehouse_dispatch": ASSETS_DIR / "map_warehouse_dispatch.png",
        "10.8 Kho có lối đi dài - warehouse_long_aisles": ASSETS_DIR / "map_warehouse_long_aisles.png",
        "12.1 Kết quả đại diện trên bảy môi trường": ASSETS_DIR / "figure_14_all_map_error.png",
    }
    for heading, path in replacements.items():
        replace_first_picture_after_heading(document, heading, path)

    set_paragraph_text(
        find_paragraph_prefix(document, "Dữ liệu chạy kín: quỹ đạo tham chiếu"),
        "Trace cuối ngày 26/07/2026 trên "
        "research_warehouse/lower_left_diagonal: selected path, ground truth "
        "Gazebo, profile vận tốc và sai số controller. Các số trong tiêu đề "
        "hình được đọc trực tiếp từ file final JSON.",
    )
    map_caption = (
        "Cùng một hình học môi trường trong occupancy grid RViz2 và mô hình SDF Gazebo."
    )
    map_captions = [
        paragraph
        for paragraph in document.paragraphs
        if normalized_text(paragraph.text) == map_caption
    ]
    if len(map_captions) != 7:
        raise RuntimeError(f"Expected seven map captions, got {len(map_captions)}")
    for paragraph, environment in zip(map_captions, ENVIRONMENTS):
        set_paragraph_text(
            paragraph,
            f"{ENV_LABEL[environment]}: occupancy grid RViz2, hình học SDF "
            "Gazebo, selected path và ground truth lấy từ trace cuối ngày "
            "26/07/2026. Các điểm start–goal mờ là toàn bộ scenario có trong map.",
        )
    set_paragraph_text(
        find_paragraph_prefix(
            document,
            "Sai số chạy kín của Pivot–G² thích nghi trên bảy môi trường.",
        ),
        "Ba nhóm sai số của bảy trace cuối ngày 26/07/2026. Chiều cao từng cột "
        "được sinh trực tiếp từ cùng JSON dùng để điền bảng ngay phía trên.",
    )


def prune_unused_main_document_images(document: Document) -> None:
    """Remove superseded image parts after replacing old data figures."""
    used_relationships = {
        blip.get(qn("r:embed"))
        for blip in document.element.xpath(".//a:blip")
        if blip.get(qn("r:embed"))
    }
    stale = [
        relationship_id
        for relationship_id, relationship in document.part.rels.items()
        if "relationships/image" in relationship.reltype
        and relationship_id not in used_relationships
    ]
    for relationship_id in stale:
        document.part.drop_rel(relationship_id)


def insert_benchmark_contract(document: Document) -> None:
    old_heading = find_paragraph(document, "11.2 Định nghĩa metric")
    set_paragraph_text(old_heading, "11.3 Định nghĩa metric")
    anchor = find_paragraph_prefix(
        document,
        "Tầng thứ hai là chạy kín.",
    )
    blocks = [
        (
            "Heading 2",
            "11.2 Hợp đồng dữ liệu của đợt kiểm chứng ngày 26/07/2026",
        ),
        (
            "Normal",
            "Đợt rà soát mới không thay thế ma trận 7.200 dòng và 42 trial ngày "
            "25/07. Nó là một lớp kiểm chứng hồi quy tập trung vào những thay đổi "
            "có nguy cơ làm xe sai hướng hoặc lệch khỏi đường sau cong. Bảy trace "
            "đại diện bao phủ đủ bảy environment; thêm hai cặp before–after ở "
            "lower_left_diagonal và right_rack_detour, cùng một cặp trước–sau "
            "luật phanh góc ở narrow_aisles."
        ),
        (
            "Normal",
            "Một trace chỉ được nhận khi action trả success, ground-truth pose "
            "nằm trong tolerance, ground-truth yaw đạt yêu cầu và robot đã dừng "
            "vật lý. selected_path_xy và ground_truth_state_trace phải tồn tại "
            "để kết quả có thể tính lại. Bốn độ dịch neo start–goal phải không "
            "vượt 0,08 m. Nominal jerk phải không vượt 0,9 m/s³ sau khi loại "
            "đúng những mẫu safety override."
        ),
        (
            "Normal",
            "Báo cáo phân biệt nominal jerk và actual jerk. Nominal jerk đo "
            "những chu kỳ S-curve còn đầy đủ miền khả thi. Actual jerk cũng chứa "
            "các lần cap an toàn hoặc biên v = 0 buộc lệnh thay đổi nhanh hơn. "
            "Nếu chỉ báo actual jerk mà không kèm cờ override, người đọc có thể "
            "kết luận sai rằng thuật toán vi phạm giới hạn ở mọi thời điểm. Nếu "
            "chỉ báo nominal jerk, người đọc lại không nhìn thấy các can thiệp "
            "an toàn. Vì vậy hai đại lượng phải tồn tại song song."
        ),
        (
            "Normal",
            "Ground-truth tracking error đo khoảng cách từ pose vật lý Gazebo "
            "đến selected path. Estimated tracking error dùng pose trong TF mà "
            "controller nhìn thấy. Odometry error so odom với chuyển động vật "
            "lý sau căn hệ. Localization error so map pose ước lượng với ground "
            "truth. Việc tách bốn nguồn giúp xác định xe lệch vì path, vì "
            "controller, vì odometry hay vì AMCL thay vì gộp mọi lỗi vào một RMSE."
        ),
    ]
    add_sequence(document, anchor, blocks)


def insert_current_audit_results(
    document: Document,
    final_data: dict[str, dict],
    audit_data: dict[str, dict],
) -> None:
    # Preserve the old published table, but do not present it as a current raw trace.
    old_heading = find_paragraph(document, "12.2 Những lỗi đã sửa được bằng dữ liệu")
    set_paragraph_text(
        old_heading,
        "12.2 Bảng legacy đã công bố trong supplement ngày 25/07/2026",
    )
    old_explanation = find_paragraph_prefix(
        document,
        "Bảng này cho thấy việc sửa projection",
    )
    set_paragraph_text(
        old_explanation,
        "Bảng phía trên được giữ nguyên theo "
        "docs/REV_ECIT_2026_ADAPTIVE_HYBRID_PIVOT_G2_SUPPLEMENT.html để lưu "
        "dấu vết kết quả đã công bố trước đợt audit hiện tại. Repository không "
        "còn giữ một cặp raw JSON độc lập đủ rõ để tái tính riêng bảng này, nên "
        "nó được ghi nhãn legacy và không được dùng làm bằng chứng cho phiên bản "
        "26/07/2026. Về mặt lịch sử, bảng cho thấy projection, profile vận tốc "
        "và terminal servo đã ảnh hưởng trực tiếp đến thời gian, tracking error "
        "và sai số cuối như thế nào. Không dùng các số này để thay cho trace "
        "hiện tại; đợt kiểm chứng sau đây có dữ liệu, mã điều khiển và hiệu chuẩn "
        "Gazebo mới hơn."
    )

    renumber = {
        "12.3 So sánh hình học trên 7.200 dòng": "12.4 So sánh hình học trên 7.200 dòng, ma trận ngày 25/07/2026",
        "12.4 So sánh chạy kín tám phương pháp ở tốc độ thích nghi": "12.5 Tám phương pháp trên lower_left_diagonal/ThetaStar, ma trận ngày 25/07/2026",
        "12.5 Tương thích với năm planner": "12.6 Năm planner trên lower_left_diagonal, ma trận ngày 25/07/2026",
        "12.6 Ca phản ví dụ phải được giữ lại trong báo cáo": "12.7 Ca phản ví dụ phải được giữ lại trong báo cáo",
    }
    for old, new in renumber.items():
        set_paragraph_text(find_paragraph(document, old), new)

    geometry_explanation = find_paragraph_prefix(
        document, "Pure Pivot–G² thích nghi giảm 92,2% Eκ so với Raw."
    )
    set_paragraph_text(
        geometry_explanation,
        geometry_explanation.text
        + " Các giá trị Eκ, clearance, chiều dài, độ lệch và runtime trong bảng "
        "là trung bình trên những hàng sinh đường thành công; số thành công trên "
        "900 được trình bày riêng để phương pháp có nhiều ca fail không được che "
        "bởi một giá trị trung bình đẹp.",
    )
    set_paragraph_text(
        find_paragraph_prefix(
            document,
            "So sánh hình học: năng lượng độ cong, clearance và tỷ lệ sinh đường thành công.",
        ),
        "So sánh hình học từ 7.200 hàng của ma trận ngày 25/07/2026: năng lượng "
        "độ cong, clearance và tỷ lệ sinh đường thành công.",
    )
    set_paragraph_text(
        find_paragraph_prefix(
            document,
            "Mức giảm năng lượng độ cong theo từng môi trường và từng phương pháp.",
        ),
        "Mức giảm năng lượng độ cong theo từng môi trường và từng phương pháp, "
        "tính từ ma trận hình học ngày 25/07/2026.",
    )
    set_paragraph_text(
        find_paragraph_prefix(
            document,
            "Ảnh hưởng của smoother và chế độ tốc độ đến thời gian",
        ),
        "Ảnh hưởng của smoother và chế độ tốc độ trong ma trận chạy kín ngày "
        "25/07/2026. Bảng ngay phía trên chỉ lấy tám method có tốc độ thích nghi "
        "trên cùng research_warehouse/lower_left_diagonal/ThetaStar.",
    )

    execution_table = table_after_heading(
        document,
        "12.5 Tám phương pháp trên lower_left_diagonal/ThetaStar, ma trận ngày 25/07/2026",
    )
    execution_table.rows[0].cells[6].text = "v thực max (m/s)"
    execution_table.rows[0].cells[7].text = "Jerk nominal P95 (m/s³)"
    format_table(execution_table, 7.5)
    planner_table = table_after_heading(
        document,
        "12.6 Năm planner trên lower_left_diagonal, ma trận ngày 25/07/2026",
    )
    planner_table.rows[0].cells[5].text = "Exit RMSE (cm)"
    format_table(planner_table, 8.5)

    lower_before = audit_data["lower_before"]
    lower_after = audit_data["lower_after"]
    rack_before = audit_data["rack_before"]
    rack_after = audit_data["rack_after"]
    narrow_before = audit_data["narrow_before_braking"]
    narrow_after = audit_data["narrow_after_braking"]

    comparison_rows = []
    for scenario, before, after in (
        ("lower_left_diagonal", lower_before, lower_after),
        ("right_rack_detour", rack_before, rack_after),
    ):
        comparison_rows.extend(
            [
                (
                    scenario,
                    "Thời gian (s)",
                    dec(before["execution_time_s"], 3),
                    dec(after["execution_time_s"], 3),
                    f"{dec(percent_reduction(after['execution_time_s'], before['execution_time_s']), 1)}%",
                ),
                (
                    scenario,
                    "GT RMSE (cm)",
                    cm(before["tracking_rmse_m"]),
                    cm(after["tracking_rmse_m"]),
                    f"{dec(percent_reduction(after['tracking_rmse_m'], before['tracking_rmse_m']), 1)}%",
                ),
                (
                    scenario,
                    "Odom vị trí P95 (cm)",
                    cm(before["odometry_position_error_p95_m"]),
                    cm(after["odometry_position_error_p95_m"]),
                    f"{dec(percent_reduction(after['odometry_position_error_p95_m'], before['odometry_position_error_p95_m']), 1)}%",
                ),
                (
                    scenario,
                    "Odom yaw P95 (rad)",
                    dec(before["odometry_yaw_error_p95_rad"], 6),
                    dec(after["odometry_yaw_error_p95_rad"], 6),
                    f"{dec(percent_reduction(after['odometry_yaw_error_p95_rad'], before['odometry_yaw_error_p95_rad']), 1)}%",
                ),
            ]
        )

    anchor = old_explanation
    blocks = [
        (
            "Heading 2",
            "12.3 Rà soát kín ngày 26/07/2026 sau khi sửa hướng và vận tốc",
        ),
        (
            "Normal",
            "Đợt rà soát này được thực hiện từ dữ liệu Gazebo ground truth, "
            "estimated pose, odometry, lệnh vận tốc và telemetry controller. "
            "Mục tiêu không phải chọn riêng một metric đẹp nhất, mà kiểm tra xem "
            "các phần mới có cùng dẫn đến một chuyển động logic hay không: robot "
            "spawn đúng hướng có thể đi, selected path bắt đầu đúng vị trí, lệnh "
            "quay phanh trước target, xe không tăng tốc khi còn lệch heading, "
            "và sau cong có thể trở về đường thay vì trôi ra ngoài."
        ),
        {
            "headers": ("Scenario", "Metric", "Trước", "Sau", "Mức giảm"),
            "rows": comparison_rows,
            "font_size": 8.0,
        },
        (
            "Normal",
            "Ở right_rack_detour, thời gian giảm "
            f"{dec(percent_reduction(rack_after['execution_time_s'], rack_before['execution_time_s']), 1)}%, "
            "GT RMSE giảm "
            f"{dec(percent_reduction(rack_after['tracking_rmse_m'], rack_before['tracking_rmse_m']), 1)}%, "
            "exit max giảm "
            f"{dec(percent_reduction(rack_after['curve_exit_tracking_max_error_m'], rack_before['curve_exit_tracking_max_error_m']), 1)}% "
            "và odometry yaw P95 giảm "
            f"{dec(percent_reduction(rack_after['odometry_yaw_error_p95_rad'], rack_before['odometry_yaw_error_p95_rad']), 1)}%. Đây là ca thể "
            "hiện rõ nhất lợi ích của việc đồng bộ hiệu chuẩn khoảng vệt lăn, "
            "neo path và profile quay. Ở lower_left_diagonal, odometry position "
            "P95 giảm 62,7%, odometry yaw P95 giảm 83,2% và estimated RMSE giảm "
            "27,5%. GT RMSE của riêng lần chạy cuối tăng nhẹ khoảng 1,7%, từ "
            "2,112 cm lên 2,148 cm; sai số cuối cũng biến động nhưng vẫn trong "
            "tolerance. Kết quả này được ghi thẳng thay vì chỉ giữ metric có lợi, "
            "vì hai lần mô phỏng có nhiễu AMCL và không thể dùng một mẫu để tuyên "
            "bố mọi đại lượng đều cải thiện tuyệt đối."
        ),
        {
            "headers": ("Metric narrow_aisles", "Trước luật phanh", "Sau luật phanh", "Mức giảm"),
            "rows": [
                (
                    "Thời gian (s)",
                    dec(narrow_before["execution_time_s"], 3),
                    dec(narrow_after["execution_time_s"], 3),
                    f"{dec(percent_reduction(narrow_after['execution_time_s'], narrow_before['execution_time_s']), 1)}%",
                ),
                (
                    "Mẫu initial_alignment",
                    str(narrow_before["adaptive_speed_mode_sample_counts"]["initial_alignment"]),
                    str(narrow_after["adaptive_speed_mode_sample_counts"]["initial_alignment"]),
                    f"{dec(percent_reduction(narrow_after['adaptive_speed_mode_sample_counts']['initial_alignment'], narrow_before['adaptive_speed_mode_sample_counts']['initial_alignment']), 1)}%",
                ),
                (
                    "Thời gian alignment xấp xỉ (s)",
                    dec(0.05 * narrow_before["adaptive_speed_mode_sample_counts"]["initial_alignment"], 2),
                    dec(0.05 * narrow_after["adaptive_speed_mode_sample_counts"]["initial_alignment"], 2),
                    f"{dec(percent_reduction(narrow_after['adaptive_speed_mode_sample_counts']['initial_alignment'], narrow_before['adaptive_speed_mode_sample_counts']['initial_alignment']), 1)}%",
                ),
                (
                    "GT RMSE (cm)",
                    cm(narrow_before["tracking_rmse_m"]),
                    cm(narrow_after["tracking_rmse_m"]),
                    f"{dec(percent_reduction(narrow_after['tracking_rmse_m'], narrow_before['tracking_rmse_m']), 1)}%",
                ),
            ],
            "font_size": 8.5,
        },
        (
            "Normal",
            "Chu kỳ telemetry là 0,05 s nên 209 mẫu initial_alignment tương "
            "đương khoảng 10,45 s, còn 93 mẫu tương đương khoảng 4,65 s. Mức "
            "giảm 55,5% không đến từ việc tăng bừa ωmax; ωmax vẫn là 0,70 rad/s. "
            "Nó đến từ việc giảm counter-rotation bằng bao phanh căn bậc hai theo "
            "αeff = 0,18 rad/s²."
        ),
        (
            "Normal",
            "Cần đọc kết quả narrow_aisles một cách thận trọng: GT RMSE tăng "
            f"{dec(-percent_reduction(narrow_after['tracking_rmse_m'], narrow_before['tracking_rmse_m']), 1)}% "
            "và exit max tăng "
            f"{dec(-percent_reduction(narrow_after['curve_exit_tracking_max_error_m'], narrow_before['curve_exit_tracking_max_error_m']), 1)}%. "
            "Vì vậy luật phanh mới chứng minh rõ việc giảm thời gian căn hướng "
            "và counter-rotation, nhưng chưa chứng minh mọi metric bám đường trên "
            "narrow_aisles đều tốt hơn. Các cột tăng được tô đỏ trong hình dưới."
        ),
        (
            "Normal",
            "Bảy trace cuối đều thành công, đạt ground-truth goal, dừng vật lý "
            "và giữ nominal jerk không quá 0,9 m/s³. Ground-truth RMSE nằm từ "
            f"{cm(min(row['tracking_rmse_m'] for row in final_data.values()))} "
            "đến "
            f"{cm(max(row['tracking_rmse_m'] for row in final_data.values()))} "
            "cm. Map open_arena có RMSE thấp nhất; warehouse_dispatch cao nhất, "
            "phù hợp với mật độ vật cản và số pha chuyển động phức tạp hơn."
        ),
    ]
    cursor = add_sequence(document, anchor, blocks)
    cursor = insert_picture_after(
        document,
        cursor,
        ASSETS_DIR / "figure_15_current_audit_comparison.png",
    )
    insert_paragraph_after(
        document,
        cursor,
        "So sánh chuẩn hóa ba cặp before–after ngày 26/07/2026. Mốc 100% là "
        "trace trước; cột xanh là giảm, cột đỏ là tăng. Hình được sinh từ cùng "
        "JSON dùng để điền hai bảng phía trên.",
        "Caption VN",
    )


def update_conclusion_and_limits(document: Document) -> None:
    contribution = find_paragraph_prefix(
        document,
        "Điểm riêng mạnh nhất của dự án không phải việc sử dụng Bézier bậc năm",
    )
    set_paragraph_text(
        contribution,
        "Điểm riêng mạnh nhất của dự án không phải việc sử dụng Bézier bậc năm, "
        "vì Bézier và tính liên tục G² đã tồn tại trong nghiên cứu trước. Đóng "
        "góp nằm ở cách toàn bộ hệ thống được ghép thành một chuỗi logic phù hợp "
        "với Nav2 và robot vi sai: neo lại start–goal liên tục để loại sai số tâm "
        "ô; xác định heading ban đầu theo occupancy map rồi xác nhận lại bằng "
        "selected path; tìm trim thích nghi ở từng góc; giữ Pivot như một trạng "
        "thái thực; dùng DP chống overlap; kiểm tra swept footprint; fallback "
        "bằng Hybrid gate; tạo speed envelope hai chiều; sửa projection theo "
        "hướng; phanh góc theo khả năng giảm tốc đo từ Gazebo; phân loại đúng "
        "zero-speed safety override; thực thi terminal servo; hiệu chuẩn odometry "
        "tách khỏi hình học CAD; và đánh giá bằng ground truth phân tầng."
    )

    hardware_limit = find_paragraph_prefix(
        document,
        "Mô phỏng không thay thế robot thật.",
    )
    set_paragraph_text(
        hardware_limit,
        "Mô phỏng không thay thế robot thật. Ma sát sàn, backlash hộp số GA25, "
        "tải hàng, sụt áp của pack 4S4P, độ trễ driver, nhiễu lidar, sai số IMU "
        "và biến dạng bánh xe ngoài thực tế khác Gazebo. Trước khi chạy thật cần "
        "đo lại footprint, L, bán kính bánh, vận tốc bánh tối đa và hệ số "
        "acceleration. Đặc biệt, PPR, vị trí encoder và chế độ quadrature hiện "
        "chưa biết nên chưa được phép tính metres_per_tick. Model BMS, fuse, "
        "buck 12 V và motor driver cũng phải được chốt bằng datasheet trước khi "
        "cấp nguồn; giá trị 52 A của cell-array không phải rating mặc định của xe."
    )

    conclusion = find_paragraph_prefix(
        document,
        "Adaptive Hybrid Pivot–G² giải quyết khoảng trống đó theo một chuỗi kín.",
    )
    set_paragraph_text(
        conclusion,
        "Adaptive Hybrid Pivot–G² giải quyết khoảng trống đó theo một chuỗi kín. "
        "Raw path được neo đúng start–goal liên tục rồi điều kiện hóa. Hướng ban "
        "đầu được suy ra từ map và được controller xác nhận lại theo preview của "
        "selected path. Mỗi góc có trạng thái Pivot và nhiều ứng viên G². Tìm "
        "kiếm thích nghi chọn trim, DP giải xung đột toàn đường, swept footprint "
        "loại va chạm, Hybrid gate giữ Simple hoặc Raw khi cần. Sau đó speed "
        "envelope hai chiều, luật phanh góc và maneuver-aware RPP biến path thành "
        "lệnh vận tốc có thể thực hiện. Benchmark tách hình học, controller, "
        "odometry, định vị và ground truth để tránh kết luận dựa trên hình ảnh RViz2."
    )

    final_source = find_paragraph_prefix(
        document,
        "Nguồn dữ liệu, hình ảnh và các kết quả định lượng",
    )
    set_paragraph_text(
        final_source,
        "Nguồn dữ liệu, hình ảnh và các kết quả định lượng trong báo cáo này "
        "được tổng hợp từ ma trận nghiên cứu ngày 25/07/2026, đợt kiểm chứng "
        "current_full_audit_20260726 và mã nguồn hiện tại của workspace "
        "agv_nav2_research_ws. Các bảng lịch sử được ghi rõ phiên bản; bảng bảy "
        "map hiện tại lấy trực tiếp từ các trace *_pivot_g2_final.json.gz."
    )


def update_appendices(document: Document) -> None:
    glossary = table_after_heading(
        document, "PHỤ LỤC A, Thuật ngữ dùng xuyên suốt dự án"
    )
    append_table_rows(
        glossary,
        [
            ("4S4P", "Bốn cell nối tiếp trong mỗi nhánh và bốn nhánh song song; tổng 16 cell."),
            ("Buck converter", "Bộ hạ áp DC–DC; trong xe dùng để tạo rail 12 V từ pack 4S tối đa 16,8 V."),
            ("C-rate", "Tỷ lệ dòng xả theo dung lượng Ah; 5C của cell 2,6 Ah tương ứng 13 A."),
            ("CW/CCW", "Clockwise/counter-clockwise, quay thuận hoặc ngược chiều kim đồng hồ."),
            ("Encoder PPR", "Số xung trên một vòng theo định nghĩa phần cứng; hiện chưa biết và không được suy đoán."),
            ("Effective wheel separation", "Khoảng vệt lăn hiệu dụng dùng để hiệu chuẩn contact/odometry Gazebo, không thay kích thước CAD."),
            ("Initial alignment", "Pha dừng và quay để hướng thân xe khớp tiếp tuyến đầu selected path trước khi chạy tiến."),
            ("Locked rotor / stall", "Trạng thái trục motor bị khóa; dòng và nhiệt tăng cao, không phải chế độ vận hành."),
            ("No-load speed", "Tốc độ khi motor gần như không mang tải; không phải tốc độ chạy liên tục của xe."),
            ("Rated-load", "Điểm làm việc có tải danh nghĩa dùng để đánh giá vận hành liên tục."),
            ("Safety override", "Can thiệp bắt buộc khi cap an toàn hoặc biên v = 0 phải thắng giới hạn jerk danh nghĩa."),
            ("Start/goal anchoring", "Phục hồi hai đầu path về đúng pose liên tục trong một ngưỡng kiểm tra."),
        ],
        8.5,
    )

    sources = table_after_heading(
        document, "PHỤ LỤC B, Vị trí source chính và tài liệu tham khảo"
    )
    append_table_rows(
        sources,
        [
            ("Neo start/goal liên tục", "src/adaptive_pivot_g2_benchmark/adaptive_pivot_g2_benchmark/path_contract.py"),
            ("Định hướng ban đầu map-aware", "src/adaptive_pivot_g2_benchmark/adaptive_pivot_g2_benchmark/initial_heading.py"),
            ("Phanh góc và preview heading", "src/adaptive_pivot_g2_controller/src/maneuver_path.cpp"),
            ("State machine initial alignment", "src/adaptive_pivot_g2_controller/src/maneuver_aware_rpp_controller.cpp"),
            ("Zero-speed override", "src/adaptive_pivot_g2_controller/src/adaptive_speed_profile.cpp"),
            ("Hợp đồng GA25 và nguồn 4S4P", "src/vacuum_robot_gazebo/config/real_robot_profile.yaml"),
            ("Mô hình URDF", "src/vacuum_robot_gazebo/urdf/vacuum_robot.urdf"),
            ("Mô hình SDF và DiffDrive", "src/vacuum_robot_gazebo/models/vacuum_robot/model.sdf"),
            ("Dữ liệu kiểm chứng hiện tại", "results/current_full_audit_20260726/*.json.gz"),
            ("Manifest hình theo dữ liệu hiện tại", "docs/bao_cao_toan_dien_assets/current_figure_manifest.json"),
            ("Ma trận hình học lịch sử", "results/conference_geometry_20260725/*.csv"),
            ("Ma trận chạy kín lịch sử", "results/conference_execution_20260725/conference_execution_compact.csv"),
            ("Bảng legacy before–after", "docs/REV_ECIT_2026_ADAPTIVE_HYBRID_PIVOT_G2_SUPPLEMENT.html"),
            ("Dựng lại báo cáo chi tiết", "tools/rebuild_detailed_algorithm_report.py"),
            ("Xuất PDF và HTML tự chứa hình", "tools/export_detailed_algorithm_report.py"),
        ],
        8.5,
    )


def set_document_metadata(document: Document) -> None:
    props = document.core_properties
    props.title = "Báo cáo toàn diện Adaptive Hybrid Pivot–G²"
    props.subject = (
        "Giải thích từ nền tảng đến thuật toán, ROS 2/Nav2, Gazebo, GA25, "
        "nguồn 4S4P và kết quả kiểm chứng ngày 26/07/2026"
    )
    props.author = "Phạm Hải Linh"
    props.keywords = (
        "Adaptive Hybrid Pivot-G2, ROS 2, Nav2, Gazebo, differential drive, "
        "GA25, 4S4P, path smoothing"
    )
    props.comments = (
        "Bản mở rộng trực tiếp từ tài liệu gốc 53 trang; không phải conference paper "
        "và không phải supplement."
    )


def validate_output(document: Document, final_data: dict[str, dict]) -> dict:
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required_phrases = (
        "7.7 Neo start và goal liên tục",
        "7.8 Xác định hướng xuất phát theo hai tầng",
        "7.9 Bao phanh góc theo khả năng giảm tốc thực",
        "7.10 Biên vận tốc bằng không",
        "9.2 Đọc đúng thông số của động cơ GA25",
        "9.6 Bộ nguồn 16 cell 18650 mắc 4S4P",
        "12.3 Rà soát kín ngày 26/07/2026",
        "0,2834 m",
        "16,8 V",
        "0,18 rad/s²",
    )
    missing = [phrase for phrase in required_phrases if phrase not in text]
    if missing:
        raise RuntimeError(f"Báo cáo thiếu nội dung bắt buộc: {missing}")
    if "Plugin DiffDrive trong Gazebo dùng 0,2809 m" in text:
        raise RuntimeError("Báo cáo vẫn còn giá trị hiệu chuẩn Gazebo cũ 0,2809 m")
    if len(document.inline_shapes) != 35:
        raise RuntimeError(
            "Báo cáo cuối phải có 34 vị trí hình gốc đã được cập nhật đúng "
            "nguồn và một hình before–after mới"
        )
    if len(document.tables) < 29:
        raise RuntimeError("Các bảng cập nhật chưa được chèn đầy đủ")
    if any(not final_data[environment].get("physically_settled") for environment in ENVIRONMENTS):
        raise RuntimeError("Có final trace chưa settled")
    return {
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "figures": len(document.inline_shapes),
        "characters": len(text),
        "words": len(text.split()),
    }


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--master", type=Path, default=DEFAULT_MASTER)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    args = parser.parse_args()

    final_data = {name: load_json(path) for name, path in FINAL_TRACES.items()}
    audit_data = {name: load_json(path) for name, path in BEFORE_AFTER_TRACES.items()}
    validate_audit(final_data, audit_data)
    validate_source_contract()
    validate_current_figure_manifest()

    document = Document(args.master)
    validate_master(document, args.master)
    preservation_snapshot = master_preservation_snapshot(document)
    insert_revision_note(document)
    insert_controller_update(document)
    insert_hardware_update(document)
    update_map_results(document, final_data)
    insert_benchmark_contract(document)
    insert_current_audit_results(document, final_data, audit_data)
    replace_data_dependent_figures(document)
    prune_unused_main_document_images(document)
    update_conclusion_and_limits(document)
    update_appendices(document)
    set_document_metadata(document)
    validate_master_preservation(document, preservation_snapshot)
    summary = validate_output(document, final_data)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(args.output)
    saved = Document(args.output)
    saved_summary = validate_output(saved, final_data)
    if saved_summary != summary:
        raise RuntimeError(
            f"Round-trip DOCX changed document structure: {summary} -> {saved_summary}"
        )
    print(
        json.dumps(
            {
                "master": str(args.master),
                "master_sha256": sha256(args.master),
                "output": str(args.output),
                "output_sha256": sha256(args.output),
                **summary,
            },
            ensure_ascii=False,
            indent=2,
        )
    )


if __name__ == "__main__":
    main()
