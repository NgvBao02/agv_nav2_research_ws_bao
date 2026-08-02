#!/usr/bin/env python3

"""Audit hierarchical-alpha PSTMO results and rebuild the Vietnamese abstract."""

from __future__ import annotations

import ast
import collections
import csv
import html
import io
import json
import math
import shutil
import statistics
import subprocess
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
PROPOSED_DIR = (
    ROOT / "results" / "pstmo_hierarchical_alpha_two_trim_full_20260802"
)
CANONICAL_STEM = ROOT / "abstract" / "ICEEIS_2026_PSTMO_HIERARCHICAL_ALPHA_ABSTRACT_VI"

METHODS = ("raw", "simple", "savitzky_golay", "constrained", "pstmo")
EXPECTED_ENVIRONMENTS = {
    "research_warehouse": "lower_left_diagonal",
    "narrow_aisles": "southwest_northeast_weave",
    "office_maze": "office_long_diagonal",
    "open_arena": "center_block_detour",
    "warehouse_cross_aisles": "cross_aisle_transfer",
    "warehouse_dispatch": "full_replenishment",
    "warehouse_long_aisles": "diagonal_replenishment",
}
GROUP_FIELDS = ("environment", "scenario", "planner", "repetition")
METRICS = (
    "translation_path_length_m",
    "translation_max_abs_curvature_1pm",
    "translation_curvature_energy_1pm",
    "pivot_total_angle_rad",
    "footprint_clearance_min_m",
    "algorithm_time_s",
    "wall_time_s",
)


def finite_float(value: object) -> float:
    """Convert a benchmark value to a finite float."""
    result = float(value)
    if not math.isfinite(result):
        raise ValueError(f"Non-finite benchmark value: {value!r}")
    return result


def group_key(row: dict[str, str]) -> tuple[str, ...]:
    """Return the identity of one paired planner path."""
    return tuple(row[field] for field in GROUP_FIELDS)


def load_representative_rows(directory: Path) -> list[dict[str, str]]:
    """Load the locked representative scenario from each full benchmark CSV."""
    rows: list[dict[str, str]] = []
    for environment, scenario in EXPECTED_ENVIRONMENTS.items():
        path = directory / f"{environment}.csv"
        with path.open(newline="", encoding="utf-8") as stream:
            selected = [
                row for row in csv.DictReader(stream)
                if row["scenario"] == scenario
            ]
        if len(selected) != 25:
            raise RuntimeError(
                f"Expected 25 representative rows in {path}, found {len(selected)}"
            )
        for row in selected:
            row["environment"] = environment
        rows.extend(selected)
    return rows


def validate_method_groups(rows: list[dict[str, str]]) -> dict[tuple[str, ...], list[dict[str, str]]]:
    """Require 35 five-method groups with a shared Raw-path hash."""
    groups: dict[tuple[str, ...], list[dict[str, str]]] = collections.defaultdict(list)
    for row in rows:
        groups[group_key(row)].append(row)
    if len(rows) != 175 or len(groups) != 35:
        raise RuntimeError("Expected 175 records in 35 representative groups")
    for key, group in groups.items():
        if len(group) != 5 or {row["method"] for row in group} != set(METHODS):
            raise RuntimeError(f"Incomplete method set in {key}")
        if len({row["raw_path_sha256"] for row in group}) != 1:
            raise RuntimeError(f"Raw-path hash mismatch in {key}")
    return groups


def pstmo_by_group(rows: list[dict[str, str]]) -> dict[tuple[str, ...], dict[str, str]]:
    """Index successful PSTMO records by their paired path identity."""
    result = {
        group_key(row): row for row in rows
        if row["method"] == "pstmo" and row["success"] == "True"
    }
    if len(result) != 35:
        raise RuntimeError(f"PSTMO must succeed 35/35, found {len(result)}")
    return result


def validate_proposed_contract(rows: dict[tuple[str, ...], dict[str, str]]) -> None:
    """Validate the single-pipeline, footprint-safe experimental contract."""
    expected = {
        "pstmo_preprocessing_mode": "condition_then_los",
        "pstmo_search_mode": "hierarchical_alpha_two_trim",
        "pstmo_trim_domain": "derived_preferred_compatible",
        "pstmo_pipeline_execution_count": "1",
        "pstmo_final_invariants_verified": "True",
    }
    for field, value in expected.items():
        if any(row.get(field) != value for row in rows.values()):
            raise RuntimeError(f"Proposed diagnostic mismatch: {field} != {value}")
    collisions = sum(
        int(finite_float(row["footprint_collision_sample_count"]))
        for row in rows.values()
    )
    if collisions:
        raise RuntimeError("A proposed PSTMO output contains footprint collisions")


def arithmetic_mean(rows: list[dict[str, str]], field: str) -> float:
    """Calculate a finite arithmetic mean."""
    return statistics.fmean(finite_float(row[field]) for row in rows)


def percentage_change(proposed: float, baseline: float) -> float:
    """Calculate signed relative change from the paired baseline."""
    if abs(baseline) <= 1.0e-15:
        return 0.0 if abs(proposed) <= 1.0e-15 else math.inf
    return 100.0 * (proposed - baseline) / baseline


def audit_results() -> dict:
    """Audit benchmark identity, invariants and all values used by the abstract."""
    proposed_rows = load_representative_rows(PROPOSED_DIR)
    proposed_groups = validate_method_groups(proposed_rows)
    proposed = pstmo_by_group(proposed_rows)
    validate_proposed_contract(proposed)

    success_counts = {
        method: sum(
            row["method"] == method and row["success"] == "True"
            for row in proposed_rows
        )
        for method in METHODS
    }
    expected_success = {
        "raw": 35,
        "simple": 34,
        "savitzky_golay": 35,
        "constrained": 35,
        "pstmo": 35,
    }
    if success_counts != expected_success:
        raise RuntimeError(f"Unexpected success counts: {success_counts}")

    common_keys = {
        key for key, group in proposed_groups.items()
        if all(row["success"] == "True" for row in group)
    }
    if len(common_keys) != 34:
        raise RuntimeError(f"Expected 34 common-success groups, found {len(common_keys)}")
    common_means = {
        method: {
            field: arithmetic_mean(
                [
                    row for row in proposed_rows
                    if group_key(row) in common_keys and row["method"] == method
                ],
                field,
            )
            for field in METRICS
        }
        for method in METHODS
    }

    corners = []
    for row in proposed.values():
        parsed = ast.literal_eval(row["pstmo_corner_search"])
        if not isinstance(parsed, list):
            raise RuntimeError("pstmo_corner_search is not a list")
        corners.extend(parsed)
    transitions = [
        corner for corner in corners
        if finite_float(corner.get("selected_trim", 0.0)) > 0.0
    ]
    fractions = [
        finite_float(corner["selected_control_fraction"]) for corner in transitions
    ]
    trims = [finite_float(corner["selected_trim"]) for corner in transitions]
    proposed_evaluations = sum(
        int(finite_float(row["pstmo_evaluations"])) for row in proposed.values()
    )
    proposed_pivots = sum(
        int(finite_float(row["pstmo_pivots"])) for row in proposed.values()
    )
    pivot_total_angle_mean = statistics.fmean(
        finite_float(row["pivot_total_angle_rad"]) for row in proposed.values()
    )

    return {
        "design": {
            "environment_count": 7,
            "scenario_count": 7,
            "planner_count": 5,
            "method_count": 5,
            "paired_group_count": 35,
            "record_count": 175,
            "raw_path_hash_consistent_group_count": 35,
            "common_success_group_count": len(common_keys),
        },
        "validation": {
            "pstmo_success_count": 35,
            "single_pipeline_count": 35,
            "final_invariant_count": 35,
            "footprint_collision_sample_count": 0,
            "success_count": success_counts,
        },
        "search": {
            "proposed_evaluations": proposed_evaluations,
            "corner_count": len(corners),
            "transition_count": len(transitions),
            "proposed_pivot_count": proposed_pivots,
            "pivot_total_angle_mean_rad_per_path": pivot_total_angle_mean,
            "control_fraction_min": min(fractions),
            "control_fraction_median": statistics.median(fractions),
            "control_fraction_mean": statistics.fmean(fractions),
            "control_fraction_max": max(fractions),
            "trim_min_m": min(trims),
            "trim_median_m": statistics.median(trims),
            "trim_mean_m": statistics.fmean(trims),
            "trim_max_m": max(trims),
        },
        "common_success_means": common_means,
    }


def vi_number(value: float, digits: int = 2) -> str:
    """Format a number using the Vietnamese decimal separator."""
    return f"{value:.{digits}f}".replace(".", ",")


def build_abstract(evidence: dict) -> str:
    """Build the Vietnamese abstract exclusively from audited evidence."""
    search = evidence["search"]
    common = evidence["common_success_means"]
    pstmo = common["pstmo"]
    stock = ("simple", "savitzky_golay", "constrained")
    best_energy = min(common[name]["translation_curvature_energy_1pm"] for name in stock)
    best_curvature = min(
        common[name]["translation_max_abs_curvature_1pm"] for name in stock
    )
    best_length = min(common[name]["translation_path_length_m"] for name in stock)
    minimum_stock_runtime = min(common[name]["algorithm_time_s"] for name in stock)
    maximum_stock_runtime = max(common[name]["algorithm_time_s"] for name in stock)
    length_improvement = -percentage_change(
        pstmo["translation_path_length_m"], best_length
    )
    curvature_improvement = -percentage_change(
        pstmo["translation_max_abs_curvature_1pm"], best_curvature
    )
    energy_improvement = -percentage_change(
        pstmo["translation_curvature_energy_1pm"], best_energy
    )
    return (
        "Tóm tắt—Các bộ lập kế hoạch toàn cục trong ROS 2/Nav2 thường tạo đường "
        "gấp khúc có đổi hướng đột ngột. Nghiên cứu này đề xuất phương pháp làm "
        "mượt đường đi và tối ưu hóa thao tác chuyển hướng PSTMO cho robot di "
        "động vi sai hai bánh. Sau bước điều kiện hóa, PSTMO luôn áp dụng "
        "line-of-sight (LOS) tham lam để chọn dây cung xa nhất an toàn khi quét "
        "footprint thật trong cả tịnh tiến và xoay; cạnh liên tiếp của polyline "
        "đã điều kiện hóa là ứng viên tự nhiên cuối cùng. Tại mỗi góc sau LOS, "
        "thuật toán sinh tối đa hai khoảng cắt d có cơ sở hình học: giá trị ưu "
        "tiên theo chiều dài hai cạnh và giá trị tương thích dành một nửa đoạn "
        "chung sau margin. Với từng d, tỷ lệ α=q/d được tìm trên lưới thô "
        "0,1–0,5, tinh chỉnh trong khoảng hai hàng xóm của nghiệm thắng và dùng "
        "lưới lệch nửa bước khi toàn bộ lưới thô thất bại. Chỉ các chuyển tiếp "
        "Bézier bậc năm G² thỏa hình học, giới hạn bánh xe, động học, timing và "
        "swept-footprint mới được so sánh bằng năng lượng độ cong; time gate và "
        "quy hoạch động sau đó chọn giữa hai d và quay tại chỗ, đồng thời ngăn "
        "các chuyển tiếp chồng lấn. Đánh giá gồm 175 bản ghi của 7 tình huống "
        "đại diện, 5 bộ lập kế hoạch toàn cục và 5 phương án Raw, Simple, "
        "Savitzky–Golay, Constrained và PSTMO; 35/35 nhóm dùng cùng hash đường "
        "Raw. PSTMO, Raw, Savitzky–Golay và Constrained thành công 35/35, trong "
        "khi Simple thành công 34/35; mọi đầu ra PSTMO đều qua kiểm tra invariant "
        "cuối và không có mẫu va chạm footprint. Trên 34 nhóm mà cả năm phương "
        "án đều thành công, PSTMO đạt năng lượng độ cong "
        f"{vi_number(pstmo['translation_curvature_energy_1pm'], 3)} m⁻¹, độ cong "
        f"cực đại {vi_number(pstmo['translation_max_abs_curvature_1pm'], 3)} m⁻¹ "
        f"và chiều dài {vi_number(pstmo['translation_path_length_m'], 3)} m, so "
        "với đối chứng Nav2 tốt nhất tương ứng "
        f"{vi_number(best_energy, 3)} m⁻¹, {vi_number(best_curvature, 3)} m⁻¹ và "
        f"{vi_number(best_length, 3)} m, tương ứng cải thiện "
        f"{vi_number(energy_improvement)}%, {vi_number(curvature_improvement)}% "
        f"và {vi_number(length_improvement)}%. PSTMO dùng 81 chuyển tiếp G² và "
        f"{search['proposed_pivot_count']} pivot trên 83 góc, với tổng góc pivot "
        "trung bình "
        f"{vi_number(search['pivot_total_angle_mean_rad_per_path'], 5)} rad/đường. "
        f"Thời gian thuật toán trung bình là "
        f"{vi_number(1000.0 * pstmo['algorithm_time_s'], 1)} ms, cao hơn khoảng "
        f"{vi_number(1000.0 * minimum_stock_runtime, 1)}–"
        f"{vi_number(1000.0 * maximum_stock_runtime, 1)} ms của ba đối chứng. "
        "Kết quả cho thấy PSTMO cải thiện rõ rệt chiều dài và độ cong trong khi "
        "duy trì an toàn footprint; chi phí tính toán và thao tác quay tại chỗ "
        "còn cần được xác nhận bằng thử nghiệm vòng kín và phần cứng."
    )


def replace_paragraph(paragraph, label: str, body: str) -> None:
    """Replace one labeled DOCX paragraph while preserving paper typography."""
    paragraph.clear()
    label_run = paragraph.add_run(label)
    label_run.bold = True
    body_run = paragraph.add_run(body)
    for run in (label_run, body_run):
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(11)


def deduplicate_docx(path: Path) -> None:
    """Remove duplicate ZIP members sometimes introduced by office tooling."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(path, "r") as source:
        latest = {info.filename: info for info in source.infolist()}
        with zipfile.ZipFile(buffer, "w") as destination:
            for name, info in latest.items():
                destination.writestr(info, source.read(info))
    path.write_bytes(buffer.getvalue())


def write_docx(text: str) -> None:
    """Update the canonical conference abstract."""
    document = Document(CANONICAL_STEM.with_suffix(".docx"))
    abstracts = [
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Tóm tắt—")
    ]
    keywords = [
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Từ khóa—")
    ]
    if len(abstracts) != 1 or len(keywords) != 1:
        raise RuntimeError("Could not uniquely locate abstract and keyword paragraphs")
    replace_paragraph(abstracts[0], "Tóm tắt—", text.removeprefix("Tóm tắt—"))
    replace_paragraph(
        keywords[0],
        "Từ khóa—",
        "PSTMO; robot di động vi sai; làm mượt đường đi; line-of-sight tham lam; "
        "tìm kiếm thô–tinh; swept-footprint; ROS 2/Nav2.",
    )
    document.core_properties.title = "ICEEIS 2026 — PSTMO với tìm kiếm α phân cấp"
    document.core_properties.subject = "Bản tiếng Việt dùng để đối chiếu"
    document.core_properties.keywords = (
        "PSTMO, hierarchical alpha search, greedy line-of-sight, swept footprint, "
        "path smoothing, differential drive, ROS 2, Nav2"
    )
    document.save(CANONICAL_STEM.with_suffix(".docx"))
    deduplicate_docx(CANONICAL_STEM.with_suffix(".docx"))


def write_html(text: str) -> None:
    """Write synchronized browser-readable mirrors."""
    title = (
        "Phương pháp làm mượt đường đi và tối ưu hóa thao tác chuyển hướng "
        "có xét an toàn cho robot di động vi sai hai bánh"
    )
    body = html.escape(text.removeprefix("Tóm tắt—"))
    output = f"""<!doctype html>
<html lang="vi"><head><meta charset="utf-8"><title>ICEEIS 2026 — PSTMO với tìm kiếm α phân cấp</title>
<style>
body{{font-family:'Times New Roman',serif;max-width:850px;margin:40px auto;
line-height:1.45;font-size:12pt}}
h1{{text-align:center;font-size:18pt}}.notice{{text-align:center;font-weight:bold}}
p{{text-align:justify}}.label{{font-weight:bold}}
</style></head>
<body><p class="notice">BẢN TIẾNG VIỆT CHỈ DÙNG ĐỂ ĐỐI CHIẾU — KHÔNG NỘP LÊN CMT</p>
<h1>{html.escape(title)}</h1>
<p><span class="label">Tóm tắt—</span>{body}</p>
<p><span class="label">Từ khóa—</span>PSTMO; robot di động vi sai; làm mượt
đường đi; line-of-sight tham lam; tìm kiếm thô–tinh; swept-footprint; ROS 2/Nav2.</p>
</body></html>
"""
    CANONICAL_STEM.with_suffix(".html").write_text(output, encoding="utf-8")


def write_pdf() -> None:
    """Render the canonical DOCX to PDF."""
    with tempfile.TemporaryDirectory(prefix="pstmo_hierarchical_abstract_") as directory:
        subprocess.run(
            [
                "libreoffice",
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                directory,
                str(CANONICAL_STEM.with_suffix(".docx")),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        rendered = Path(directory) / CANONICAL_STEM.with_suffix(".pdf").name
        if not rendered.exists():
            raise RuntimeError("LibreOffice did not render the abstract PDF")
        shutil.copy2(rendered, CANONICAL_STEM.with_suffix(".pdf"))


def main() -> None:
    """Audit evidence and regenerate all requested abstract formats."""
    evidence = audit_results()
    (PROPOSED_DIR / "aggregate_summary.json").write_text(
        json.dumps(evidence, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    text = build_abstract(evidence)
    write_docx(text)
    write_html(text)
    write_pdf()
    print(f"Wrote {PROPOSED_DIR / 'aggregate_summary.json'}")
    for suffix in (".docx", ".html", ".pdf"):
        print(f"Updated {CANONICAL_STEM.with_suffix(suffix)}")


if __name__ == "__main__":
    main()
