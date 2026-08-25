#!/usr/bin/env python3
"""Insert redesigned Figures 6–10 and aligned captions into the revised paper."""

from __future__ import annotations

import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "final_bao_ICEEIS/ICEEIS_2026_PSTMO_final_section_V_revised.docx"
FIGURE_DIR = ROOT / "final_bao_ICEEIS/section_v_redesigned_figures"
OUTPUT = ROOT / "final_bao_ICEEIS/ICEEIS_2026_PSTMO_final_section_V_redesigned.docx"

MEDIA_REPLACEMENTS = {
    "word/media/image6.png": FIGURE_DIR / "figure_06_open_arena.png",
    "word/media/image7.png": FIGURE_DIR / "figure_07_narrow_aisles.png",
    "word/media/image8.png": FIGURE_DIR / "figure_08_warehouse_cross_aisles.png",
    "word/media/image9.png": FIGURE_DIR / "figure_09_paired_curvature_ratio.png",
    "word/media/image10.png": FIGURE_DIR / "figure_10_runtime_distribution.png",
}

CAPTIONS = {
    "Hình 6.": (
        "Hình 6. ",
        "Ca Theta* đại diện trong không gian mở: (a) bối cảnh bản đồ; (b) vùng chuyển tiếp phóng đại với các tư thế hình bao "
        "robot. Ô bên dưới báo cáo số liệu của riêng ca minh họa, không phải trung bình môi trường.",
    ),
    "Hình 7.": (
        "Hình 7. ",
        "Ca Theta* đại diện trong lối đi hẹp: (a) bối cảnh bản đồ; (b) vùng chuyển tiếp phóng đại với các tư thế hình bao "
        "robot. Khung nét đứt trong (a) xác định miền được phóng đại ở (b).",
    ),
    "Hình 8.": (
        "Hình 8. ",
        "Ca Theta* đại diện trong kho có lối giao cắt: (a) bối cảnh bản đồ; (b) vùng chuyển tiếp phóng đại với các tư thế hình "
        "bao robot. Đường Raw và PSTMO được hiển thị trên cùng bản đồ chi phí.",
    ),
    "Hình 9.": (
        "Hình 9. ",
        "Tỷ số ghép cặp Eκ(PSTMO)/Eκ(Simple) trên 15 đường đầu vào; giá trị < 1 biểu thị PSTMO có Eκ thấp hơn. "
        "Viền đỏ đánh dấu ba trường hợp > 1.",
    ),
    "Hình 10.": (
        "Hình 10. ",
        "Phân bố thời gian xử lý của 15 nhóm ghép cặp (thang logarit). Đường xám: khoảng biến thiên; đoạn đậm: khoảng tứ phân vị; "
        "hình thoi: trung vị; điểm rỗng tại 3 ms: phép đo dưới độ phân giải.",
    ),
}

SCENARIO_TEXT = (
    "Ba bản đồ được chọn để đại diện cho ba cấu trúc hình học khác nhau: (i) không gian mở với một khối cản trung tâm, từ "
    "(−2,20; −0,60) m đến (1,20; −0,60) m; (ii) lối đi hẹp theo đường chéo tây nam–đông bắc, từ "
    "(−5,00; −3,00) m đến (5,00; 3,00) m; và (iii) kho có lối giao cắt, từ (−2,00; −2,80) m đến "
    "(2,00; 2,80) m. Cả năm phương án đều tạo được đầu ra trong toàn bộ 15 nhóm ghép cặp. Hình 6–8 minh họa ca Theta* "
    "đại diện trong mỗi môi trường; mỗi hình gồm bối cảnh bản đồ, một vùng chuyển tiếp phóng đại và số liệu riêng của ca minh họa."
)

FIGURE_DISCUSSION = (
    "Hình 9 cho thấy lợi thế trung bình của PSTMO so với Simple không xuất hiện đồng đều trên mọi đường đầu vào. PSTMO có "
    "Eκ thấp hơn trong 12/15 nhóm; ba ngoại lệ gồm Theta* và SmacHybrid ở lối đi hẹp, cùng SmacHybrid ở kho có lối giao cắt. "
    "Thông tin ghép cặp này không thể hiện trong các trung bình của Bảng IV. Hình 10 cho thấy thời gian PSTMO trải từ 39 đến "
    "237 ms, trong khi Constrained nằm trong khoảng 6–45 ms; các phép đo của Simple và Savitzky–Golay chủ yếu nằm tại hoặc "
    "dưới ngưỡng phân giải 3 ms. Vì mỗi điểm chỉ dựa trên một lần chạy, hai hình được dùng để mô tả độ biến thiên giữa các "
    "đầu vào chứ không hỗ trợ suy luận thống kê."
)


def replace_two_runs(paragraph, first_text: str, second_text: str) -> None:
    if len(paragraph.runs) < 2:
        raise RuntimeError(f"Expected two formatted runs in caption: {paragraph.text!r}")
    first_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs[0]._r.rPr is not None else None
    second_properties = deepcopy(paragraph.runs[1]._r.rPr) if paragraph.runs[1]._r.rPr is not None else None
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    first_run = paragraph.add_run(first_text)
    second_run = paragraph.add_run(second_text)
    if first_properties is not None:
        first_run._r.insert(0, first_properties)
    if second_properties is not None:
        second_run._r.insert(0, second_properties)


def replace_body_text(paragraph, text: str) -> None:
    run_properties = None
    for run in paragraph.runs:
        if run._r.rPr is not None:
            run_properties = deepcopy(run._r.rPr)
            break
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def revise_text(document: Document) -> None:
    caption_done = set()
    scenario_done = False
    discussion_done = False
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        for prefix, (first, second) in CAPTIONS.items():
            if text.startswith(prefix):
                replace_two_runs(paragraph, first, second)
                caption_done.add(prefix)
                break
        if text.startswith("Ba bản đồ được chọn để đại diện cho ba cấu trúc hình học khác nhau:"):
            replace_body_text(paragraph, SCENARIO_TEXT)
            scenario_done = True
        elif text.startswith("Hình 9 và Hình 10 trực quan hóa"):
            replace_body_text(paragraph, FIGURE_DISCUSSION)
            discussion_done = True
    if caption_done != set(CAPTIONS):
        raise RuntimeError(f"Missing captions: {sorted(set(CAPTIONS) - caption_done)}")
    if not scenario_done or not discussion_done:
        raise RuntimeError("Could not locate Section V paragraphs aligned with the redesigned figures")


def apply() -> None:
    for path in MEDIA_REPLACEMENTS.values():
        if not path.exists():
            raise FileNotFoundError(path)
    document = Document(SOURCE)
    revise_text(document)
    document.core_properties.comments = (
        "Phần V và Hình 6–10 đã được biên tập lại. Hình 9–10 dùng dữ liệu từng nhóm ghép cặp để bổ sung cho các bảng trung bình."
    )
    with tempfile.TemporaryDirectory(prefix="iceeis_section_v_figures_") as temp_name:
        intermediate = Path(temp_name) / "text_revised.docx"
        document.save(intermediate)
        with zipfile.ZipFile(intermediate, "r") as source_zip, zipfile.ZipFile(OUTPUT, "w") as output_zip:
            for item in source_zip.infolist():
                replacement = MEDIA_REPLACEMENTS.get(item.filename)
                data = replacement.read_bytes() if replacement is not None else source_zip.read(item.filename)
                output_zip.writestr(item, data)
    print(OUTPUT)


if __name__ == "__main__":
    apply()
