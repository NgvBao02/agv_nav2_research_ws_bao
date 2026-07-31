#!/usr/bin/env python3

"""Create the smoother-only Vietnamese abstract brief as an illustrated DOCX."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "DE_CUONG_ABSTRACT_SMOOTHER_PIVOT_G2.docx"
BLUE = "17365D"
LIGHT_BLUE = "DDEBF7"
LIGHT_YELLOW = "FFF2CC"
LIGHT_GRAY = "F2F2F2"


def font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    prop = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    prop.append(node)


def margins(cell, value=90):
    prop = cell._tc.get_or_add_tcPr()
    node = prop.first_child_found_in("w:tcMar")
    if node is None:
        node = OxmlElement("w:tcMar")
        prop.append(node)
    for name in ("top", "start", "bottom", "end"):
        item = OxmlElement(f"w:{name}")
        item.set(qn("w:w"), str(value))
        item.set(qn("w:type"), "dxa")
        node.append(item)


def page_number(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for tag, attr, value in (
        ("w:fldChar", "w:fldCharType", "begin"),
        ("w:instrText", "xml:space", "preserve"),
        ("w:fldChar", "w:fldCharType", "separate"),
        ("w:t", None, None),
        ("w:fldChar", "w:fldCharType", "end"),
    ):
        node = OxmlElement(tag)
        if attr:
            node.set(qn(attr), value)
        if tag == "w:instrText":
            node.text = " PAGE "
        if tag == "w:t":
            node.text = "1"
        run._r.append(node)


def configure(document):
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(1.9)
    page_number(section)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)
    for name, size in (("Title", 18), ("Heading 1", 15), ("Heading 2", 13)):
        style = document.styles[name]
        style.font.name = "Times New Roman"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True


def paragraph(document, text, *, italic=False, center=False):
    item = document.add_paragraph()
    if center:
        item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = item.add_run(text)
    font(run, italic=italic)
    return item


def formula(document, text):
    item = document.add_paragraph()
    item.alignment = WD_ALIGN_PARAGRAPH.CENTER
    item.paragraph_format.space_before = Pt(3)
    item.paragraph_format.space_after = Pt(5)
    run = item.add_run(text)
    font(run, size=11.5, italic=True)
    return item


def callout(document, title, text, fill=LIGHT_YELLOW):
    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, 150)
    p = cell.paragraphs[0]
    first = p.add_run(title + " ")
    font(first, bold=True)
    second = p.add_run(text)
    font(second)
    document.add_paragraph()


def bullets(document, items, numbered=False):
    for text in items:
        p = document.add_paragraph(style="List Number" if numbered else "List Bullet")
        font(p.add_run(text))


def table(document, headers, rows, widths):
    result = document.add_table(rows=1, cols=len(headers))
    result.style = "Table Grid"
    result.alignment = WD_TABLE_ALIGNMENT.CENTER
    result.autofit = False
    for index, text in enumerate(headers):
        cell = result.rows[0].cells[index]
        cell.text = text
        cell.width = Cm(widths[index])
        shade(cell, BLUE)
        margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                font(run, size=9.5, bold=True, color="FFFFFF")
    for row_index, values in enumerate(rows):
        cells = result.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = str(value)
            cell.width = Cm(widths[index])
            margins(cell)
            if row_index % 2:
                shade(cell, LIGHT_GRAY)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    font(run, size=9.2)
    document.add_paragraph()


def figure(document, relative_path, caption, width=16.4):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(ROOT / relative_path), width=Cm(width))
    cp = document.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(cp.add_run(caption), size=10, italic=True)


def page_break(document):
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def create():
    doc = Document()
    configure(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(title.add_run("ĐỀ CƯƠNG SƠ BỘ GỬI GIẢNG VIÊN HƯỚNG DẪN"), 18, True, color=BLUE)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        subtitle.add_run(
            "ADAPTIVE HYBRID PIVOT–G²\n"
            "THUẬT TOÁN HẬU XỬ LÝ VÀ LÀM MƯỢT ĐƯỜNG ĐI\n"
            "CHO ROBOT VI SAI TRONG ROS 2/NAV2"
        ),
        15,
        True,
        color=BLUE,
    )
    paragraph(
        doc,
        "Adaptive Hybrid Pivot–G2 Path Smoothing for Differential-Drive Robots in ROS 2/Nav2",
        italic=True,
        center=True,
    )
    callout(
        doc,
        "Phạm vi hiện tại:",
        "Tài liệu chỉ trình bày smoother và đánh giá hình học đường đi; các phần "
        "thực thi robot không nằm trong phạm vi abstract. Luật chọn nhánh và tham "
        "số vẫn đang được hiệu chỉnh.",
    )

    doc.add_heading("1. Đoạn nhắn ngắn có thể gửi cô", level=1)
    paragraph(
        doc,
        "Thưa cô, em dự kiến nghiên cứu một thuật toán hậu xử lý đường đi cho "
        "robot vi sai trong ROS 2/Nav2. Đường do global planner sinh ra thường "
        "là một polyline gồm nhiều đoạn thẳng nối nhau tại các góc gấp. Nếu giữ "
        "nguyên, robot có thể phải quay nhiều; nếu làm tròn tất cả các góc, đường "
        "mới có thể cắt gần hoặc đi vào vật cản. Vì vậy, tại mỗi góc em dự kiến "
        "tạo hai loại trạng thái: giữ góc và đánh dấu thao tác quay tại chỗ "
        "(Pivot), hoặc thay góc bằng một đoạn Bézier bậc năm liên tục hình học "
        "G² khi còn đủ không gian."
    )
    paragraph(
        doc,
        "Với mỗi góc, thuật toán tìm kiếm thích nghi giá trị cắt góc/bán kính, "
        "sinh nhiều ứng viên và loại các ứng viên không an toàn bằng kiểm tra "
        "toàn footprint trên costmap. Các ứng viên còn lại được ghép bằng quy "
        "hoạch động trên toàn đường để tránh hai đoạn cong kề nhau chồng lấn và "
        "chọn chuỗi có chi phí hình học thấp. Sau đó một cổng Hybrid dự kiến so "
        "sánh nhánh Pivot–G² với một smoother nền; nếu các đường làm mượt không "
        "bảo đảm an toàn thì hệ thống trả về đường Raw."
    )
    paragraph(
        doc,
        "Em dự kiến thử nghiệm trên 7 môi trường mô phỏng Gazebo/Nav2, tổng cộng "
        "60 cặp điểm đầu–đích, 5 global planner, 8 biến thể đường/phương pháp và "
        "3 lần lặp. Như vậy có 900 nhóm cùng đầu vào Raw và 7.200 bản ghi đường "
        "hình học để so sánh. Các điểm đánh giá gồm tỷ lệ smoother tạo được "
        "đường, va chạm footprint, khoảng hở tới vật cản, chiều dài đường, độ "
        "lệch khỏi Raw, độ cong cực đại, tích phân bình phương độ cong, số trạng "
        "thái Pivot/G² và thời gian xử lý. Kết quả kỳ vọng là giảm góc gấp và "
        "dao động độ cong nhưng vẫn giữ an toàn và khả năng trả về đường gốc khi "
        "làm mượt không phù hợp. Hiện em chưa chốt phần trăm cải thiện vì một số "
        "bước lựa chọn và xử lý ứng viên vẫn cần sửa và chạy lại."
    )

    page_break(doc)
    doc.add_heading("2. Bản tóm tắt sơ bộ để cô phát triển thành abstract", level=1)
    abstract = doc.add_table(rows=1, cols=1)
    cell = abstract.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    margins(cell, 160)
    text = (
        "Đường đi do các bộ lập kế hoạch toàn cục sinh ra thường có dạng polyline "
        "với các góc gấp, gây bất lợi cho chuyển động của robot vi sai. Các "
        "phương pháp làm mượt thuần túy có thể giảm độ gấp khúc nhưng đồng thời "
        "làm đường lệch khỏi hành lang an toàn hoặc cắt gần vật cản. Nghiên cứu "
        "này đề xuất Adaptive Hybrid Pivot–G2, một thuật toán hậu xử lý đường đi "
        "tích hợp trong ROS 2/Nav2. Tại mỗi điểm đổi hướng, thuật toán xem xét "
        "trạng thái Pivot, giữ nguyên đỉnh góc để robot có thể đổi hướng tại chỗ, "
        "và trạng thái G², thay góc bằng đoạn chuyển tiếp Bézier bậc năm liên tục "
        "hình học với hai đoạn thẳng lân cận. Miền tham số cắt góc được xác định "
        "từ hình học cục bộ và được lấy mẫu thích nghi. Mỗi ứng viên được kiểm "
        "tra độ hợp lệ hình học và vùng quét footprint trên costmap. Các ứng "
        "viên an toàn sau đó được ghép bằng quy hoạch động để tránh chồng lấn "
        "giữa các góc kề và giảm một hàm chi phí gồm độ cong, chiều dài, độ lệch "
        "đường và phạt Pivot. Một cổng Hybrid so sánh các nhánh làm mượt và sử "
        "dụng đường Raw làm phương án dự phòng khi không có ứng viên an toàn. "
        "Kế hoạch đánh giá gồm 7 môi trường mô phỏng, 60 tình huống điểm đầu–"
        "đích, 5 global planner, 8 phương pháp và 3 lần lặp, tạo 7.200 bản ghi "
        "đường hình học trong 900 nhóm đầu vào ghép cặp. Các đặc tính được đánh "
        "giá gồm tỷ lệ thành công, an toàn footprint, khoảng hở, chiều dài, độ "
        "lệch, độ cong và thời gian xử lý. Nghiên cứu kỳ vọng xác định được sự "
        "đánh đổi giữa độ trơn, mức thay đổi đường và độ an toàn; kết quả định "
        "lượng sẽ được cập nhật sau khi thuật toán được ổn định và ma trận thử "
        "nghiệm được chạy lại."
    )
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    font(p.add_run(text), 11.3)
    doc.add_paragraph()
    keyword = doc.add_paragraph()
    font(keyword.add_run("Từ khóa dự kiến: "), bold=True)
    font(keyword.add_run("robot vi sai; làm mượt đường đi; Bézier bậc năm; liên tục G²; quay tại chỗ; Nav2; Gazebo."))

    page_break(doc)
    doc.add_heading("3. Kiến thức nền tảng và giải thích ký hiệu", level=1)
    doc.add_heading("3.1 Đường đi polyline và góc đổi hướng", level=2)
    paragraph(
        doc,
        "Đường Raw từ global planner được xem là một dãy điểm trong hệ tọa độ map:"
    )
    formula(doc, "P = {p₀, p₁, …, pₙ},     pᵢ = [xᵢ, yᵢ]ᵀ")
    bullets(
        doc,
        [
            "P là toàn bộ đường đi dạng polyline; pᵢ là điểm thứ i.",
            "xᵢ và yᵢ là tọa độ của pᵢ trong map, đơn vị mét; n+1 là tổng số điểm.",
        ],
    )
    paragraph(doc, "Vector chỉ phương và chiều dài đoạn thứ i:")
    formula(doc, "eᵢ = pᵢ₊₁ − pᵢ,          lᵢ = ||eᵢ||₂")
    paragraph(
        doc,
        "eᵢ là vector từ pᵢ đến pᵢ₊₁; lᵢ là chiều dài Euclid, đơn vị mét; "
        "||·||₂ là chuẩn Euclid. Hướng của đoạn và góc đổi hướng tại đỉnh pᵢ là:"
    )
    formula(
        doc,
        "ψᵢ = atan2(yᵢ₊₁ − yᵢ, xᵢ₊₁ − xᵢ),     "
        "Δψᵢ = wrapToPi(ψᵢ − ψᵢ₋₁)"
    )
    paragraph(
        doc,
        "ψᵢ₋₁ và ψᵢ là hướng đi vào và đi ra; wrapToPi đưa góc về [−π,π]. "
        "Theo quy ước tọa độ, Δψᵢ dương biểu diễn rẽ trái và âm biểu diễn rẽ phải."
    )

    doc.add_heading("3.2 Liên tục G⁰, G¹ và G²", level=2)
    bullets(
        doc,
        [
            "G⁰ – liên tục vị trí: hai đoạn gặp nhau tại cùng một điểm.",
            "G¹ – liên tục tiếp tuyến: hai đoạn có cùng hướng tiếp tuyến tại điểm nối.",
            "G² – liên tục độ cong hình học: vị trí, hướng tiếp tuyến và độ cong ở hai phía phù hợp.",
        ],
    )
    paragraph(
        doc,
        "Hai đoạn kề của polyline là đường thẳng nên có độ cong bằng 0. Để nối "
        "G², đoạn Bézier chuyển tiếp được thiết kế sao cho:"
    )
    formula(doc, "κ(0) = 0,          κ(1) = 0")
    callout(
        doc,
        "Phân biệt G² và C²:",
        "C² yêu cầu các đạo hàm theo cùng một tham số bằng nhau. G² là điều kiện "
        "hình học, cho phép đổi tham số nhưng vẫn giữ vị trí, hướng tiếp tuyến và "
        "độ cong tương thích.",
        fill=LIGHT_BLUE,
    )

    doc.add_heading("3.3 Đường Bézier bậc năm", level=2)
    paragraph(
        doc,
        "Đường Bézier bậc năm được xác định bởi sáu điểm điều khiển B₀…B₅:"
    )
    formula(doc, "r(u) = Σᵢ₌₀⁵ bᵢ,₅(u)Bᵢ,      0 ≤ u ≤ 1")
    formula(doc, "bᵢ,₅(u) = C(5,i)(1−u)⁵⁻ⁱuⁱ")
    bullets(
        doc,
        [
            "r(u)=[x(u),y(u)]ᵀ là một điểm trên đường; u là tham số chuẩn hóa.",
            "Bᵢ là điểm điều khiển thứ i, đơn vị mét; bᵢ,₅ là hàm Bernstein.",
            "C(5,i)=5!/[i!(5−i)!] là hệ số tổ hợp.",
        ],
    )
    paragraph(doc, "Hai đạo hàm đầu của đường cong:")
    formula(doc, "r′(u) = 5Σᵢ₌₀⁴ bᵢ,₄(u)(Bᵢ₊₁ − Bᵢ)")
    formula(doc, "r″(u) = 20Σᵢ₌₀³ bᵢ,₃(u)(Bᵢ₊₂ − 2Bᵢ₊₁ + Bᵢ)")
    paragraph(
        doc,
        "r′ xác định hướng tiếp tuyến; r″ mô tả mức thay đổi tiếp tuyến. Sáu "
        "điểm điều khiển cung cấp đủ bậc tự do để đặt ràng buộc ở cả hai đầu."
    )

    page_break(doc)
    doc.add_heading("3.4 Khoảng cắt góc và bán kính tham chiếu", level=2)
    paragraph(
        doc,
        "Khi hai điểm nối cách đỉnh một khoảng bằng nhau d, bán kính của cung "
        "tròn tiếp xúc tương đương dùng làm đại lượng tham chiếu là:"
    )
    formula(doc, "R = d/tan(|Δψ|/2),          d = R tan(|Δψ|/2)")
    bullets(
        doc,
        [
            "d là khoảng trim/cắt dọc mỗi cạnh, đơn vị mét.",
            "R là bán kính hình học tham chiếu, đơn vị mét.",
            "Δψ là góc đổi hướng, đơn vị radian.",
        ],
    )
    paragraph(
        doc,
        "Công thức trên dùng diễn giải và giới hạn miền tìm kiếm. Đường thực tế "
        "là Bézier nên bán kính tức thời có thể thay đổi dọc đoạn chuyển tiếp."
    )

    doc.add_heading("3.5 Độ cong, chiều dài và năng lượng độ cong", level=2)
    formula(
        doc,
        "κ(u) = [x′(u)y″(u) − y′(u)x″(u)]/"
        "[x′(u)² + y′(u)²]³ᐟ²"
    )
    bullets(
        doc,
        [
            "x′, y′ và x″, y″ là đạo hàm bậc nhất và bậc hai theo u.",
            "κ là độ cong có dấu, đơn vị m⁻¹; |κ| càng lớn thì góc cua càng gắt.",
            "Khi κ khác 0, bán kính cong tức thời là ρ = 1/|κ|.",
        ],
    )
    paragraph(doc, "Phần tử chiều dài cung và chiều dài đoạn cong:")
    formula(doc, "ds = ||r′(u)||₂du,          L = ∫₀¹||r′(u)||₂du")
    paragraph(doc, "Chỉ tiêu năng lượng độ cong:")
    formula(doc, "Eκ = ∫κ(s)²ds = ∫₀¹κ(u)²||r′(u)||₂du")
    paragraph(
        doc,
        "s là tọa độ chiều dài cung; L có đơn vị mét; Eκ có đơn vị m⁻¹. Eκ nhỏ "
        "thường biểu diễn đường ít gấp và ít dao động độ cong hơn."
    )
    callout(
        doc,
        "Lưu ý:",
        "Eκ là thước đo hình học, không phải năng lượng điện hoặc mức tiêu thụ pin.",
    )

    doc.add_heading("3.6 Footprint, swept-footprint và clearance", level=2)
    paragraph(
        doc,
        "Gọi F là footprint trong hệ tọa độ thân xe. Tại pose "
        "q(s)=[x(s),y(s),θ(s)], footprint trong map là:"
    )
    formula(doc, "F(q(s)) = Q(θ(s))F + [x(s), y(s)]ᵀ")
    formula(doc, "Q(θ) = [[cosθ, −sinθ], [sinθ, cosθ]]")
    paragraph(doc, "Vùng quét dọc đường và điều kiện không va chạm:")
    formula(doc, "S = ⋃ₛF(q(s)),          S ∩ O = ∅")
    paragraph(doc, "Clearance nhỏ nhất:")
    formula(
        doc,
        "cmin = minₛ min_{a∈F(q(s)), b∈O} ||a−b||₂"
    )
    paragraph(
        doc,
        "Q(θ) là ma trận quay phẳng, ký hiệu khác R để không nhầm với bán kính. "
        "θ(s) là hướng tiếp tuyến; O là vùng vật cản; S là swept-footprint; "
        "cmin là khoảng hở nhỏ nhất, đơn vị mét. Chương trình xấp xỉ các đại "
        "lượng này bằng lấy mẫu đường và truy vấn costmap."
    )

    page_break(doc)
    doc.add_heading("3.7 Hàm chi phí và quy hoạch động", level=2)
    paragraph(doc, "Một hàm chi phí thiết kế cho toàn đường có dạng:")
    formula(doc, "J = wκEκ + wL L + wdev D + wpivot Npivot + wobs Jobs")
    bullets(
        doc,
        [
            "J là tổng chi phí; Eκ là năng lượng độ cong; L là chiều dài.",
            "D là độ lệch so với Raw; Npivot là số trạng thái Pivot.",
            "Jobs là phạt gần vật cản; các w là trọng số không âm.",
        ],
    )
    paragraph(
        doc,
        "Do các thành phần có đơn vị và thang giá trị khác nhau, cần chuẩn hóa "
        "hoặc chọn trọng số phù hợp. Đây là cấu trúc thiết kế, chưa phải bộ "
        "trọng số đã chốt."
    )
    paragraph(
        doc,
        "Gọi Cᵢ(j) là chi phí cục bộ của ứng viên j tại góc i; Tᵢ(k,j) là "
        "chi phí chuyển, hoặc vô cùng nếu hai ứng viên không tương thích:"
    )
    formula(
        doc,
        "Vᵢ(j) = Cᵢ(j) + minₖ[Vᵢ₋₁(k) + Tᵢ(k,j)]"
    )
    bullets(
        doc,
        [
            "Vᵢ(j) là chi phí nhỏ nhất từ góc đầu tới trạng thái j tại góc i.",
            "k duyệt các trạng thái của góc i−1; truy vết cho chuỗi Pivot/G² được chọn.",
            "Với N góc và tối đa K trạng thái mỗi góc, độ phức tạp là O(NK²).",
        ],
    )

    doc.add_heading("3.8 Bảng ký hiệu tổng hợp", level=2)
    table(
        doc,
        ["Ký hiệu", "Ý nghĩa", "Đơn vị"],
        [
            ("pᵢ=[xᵢ,yᵢ]ᵀ", "Điểm thứ i của polyline", "m"),
            ("eᵢ, lᵢ", "Vector và chiều dài đoạn thứ i", "lᵢ: m"),
            ("ψᵢ, Δψᵢ", "Hướng đoạn và góc đổi hướng", "rad"),
            ("u", "Tham số chuẩn hóa của Bézier", "không đơn vị"),
            ("B₀…B₅", "Sáu điểm điều khiển Bézier", "m"),
            ("d", "Khoảng cắt/trim tại góc", "m"),
            ("R, ρ", "Bán kính tham chiếu và tức thời", "m"),
            ("κ", "Độ cong có dấu", "m⁻¹"),
            ("s, L", "Chiều dài cung và tổng chiều dài", "m"),
            ("Eκ", "Tích phân bình phương độ cong", "m⁻¹"),
            ("q(s), θ(s), Q(θ)", "Pose, hướng tiếp tuyến, ma trận quay", "m, rad, –"),
            ("F, S, O", "Footprint, vùng quét, vùng vật cản", "tập hình học"),
            ("cmin", "Clearance nhỏ nhất", "m"),
            ("D", "Độ lệch so với Raw", "m"),
            ("Npivot", "Số trạng thái Pivot", "số đếm"),
            ("J, Cᵢ, Tᵢ, Vᵢ", "Các chi phí trong tối ưu và DP", "phụ thuộc chuẩn hóa"),
            ("N, K", "Số góc và số trạng thái tối đa", "số đếm"),
        ],
        [3.5, 9.0, 3.5],
    )

    page_break(doc)
    doc.add_heading("4. Ý tưởng thuật toán", level=1)
    doc.add_heading("4.1 Hai trạng thái tại mỗi góc", level=2)
    figure(
        doc,
        "docs/bao_cao_toan_dien_assets/figure_05_pivot_g2_geometry.png",
        "Hình 1. Một góc có thể được thay bằng đoạn chuyển tiếp Bézier G² hoặc giữ làm trạng thái Pivot.",
    )
    bullets(
        doc,
        [
            "Pivot: giữ đỉnh góc, không chèn đoạn cong; phù hợp khi hành lang hẹp hoặc không đủ chiều dài cắt góc.",
            "G²: cắt một khoảng d trên hai cạnh kề và nối bằng Bézier bậc năm có tiếp tuyến và độ cong nối êm với đoạn thẳng.",
        ],
    )
    paragraph(
        doc,
        "Với góc đổi hướng Δψ và khoảng cắt d, bán kính hình học tham chiếu là "
        "R = d/tan(|Δψ|/2). Đoạn chuyển tiếp được biểu diễn bởi "
        "r(u) = ΣBᵢ,₅(u)Pᵢ, 0 ≤ u ≤ 1. Độ cong được tính từ "
        "κ(u) = (x′y″ − y′x″)/(x′² + y′²)^(3/2)."
    )
    doc.add_heading("4.2 Sinh và kiểm tra ứng viên", level=2)
    paragraph(
        doc,
        "Miền khả thi của d được giới hạn bởi chiều dài hai đoạn kề, khoảng cách "
        "tới góc lân cận và vùng trống trên costmap. Tìm kiếm từ thô đến tinh "
        "được dùng để lấy thêm mẫu gần biên safe/unsafe và vùng có chi phí thấp."
    )
    bullets(
        doc,
        [
            "Loại đoạn chuyển tiếp tự cắt hoặc có hình học bất thường.",
            "Loại ứng viên đổi dấu độ cong ngoài ý muốn hoặc vượt giới hạn độ cong.",
            "Không cho hai đoạn chuyển tiếp kề nhau chồng lấn.",
            "Kiểm tra swept-footprint dọc toàn đoạn; loại ứng viên đi vào ô chiếm dụng.",
        ],
    )
    doc.add_heading("4.3 Ghép quyết định trên toàn đường", level=2)
    figure(
        doc,
        "docs/bao_cao_toan_dien_assets/figure_06_search_dp.png",
        "Hình 2. Tìm kiếm thích nghi tạo các trạng thái an toàn; quy hoạch động chọn chuỗi trạng thái trên toàn đường.",
    )
    paragraph(
        doc,
        "Mỗi góc giữ một số ứng viên G² an toàn và một trạng thái Pivot. Hai "
        "trạng thái kề chỉ được nối nếu tổng khoảng cắt và margin không vượt "
        "chiều dài đoạn chung. Với N góc và tối đa K trạng thái mỗi góc, bước "
        "ghép có độ phức tạp dự kiến O(NK²). Hàm chi phí thiết kế có dạng "
        "J = wκ∫κ²ds + wL L + wdev D(raw,smooth) + wpivot Npivot + Jobs. "
        "Các trọng số chưa được xem là cố định cho tới khi hoàn tất kiểm tra."
    )
    doc.add_heading("4.4 Cổng Hybrid và fallback", level=2)
    paragraph(
        doc,
        "Hybrid chỉ nhận một đường sau khi đường đó vượt kiểm tra footprint và "
        "các điều kiện hình học. Việc so sánh các nhánh dự kiến dựa trên peak "
        "cost/clearance và effort hình học. Nếu không có nhánh làm mượt hợp lệ, "
        "thuật toán trả Raw thay vì cố tạo một đường cong không an toàn. Luật "
        "chọn nhánh này đang được hiệu chỉnh."
    )

    page_break(doc)
    doc.add_heading("5. Kế hoạch thử nghiệm", level=1)
    table(
        doc,
        ["Thành phần", "Số lượng", "Ý nghĩa"],
        [
            ("Môi trường Gazebo/Nav2", "7", "Kho nghiên cứu, lối hẹp, văn phòng, không gian mở và ba bố trí kho"),
            ("Cặp điểm đầu–đích", "60", "Đường thẳng, rẽ gấp, chữ S, hành lang hẹp và đường vòng vật cản"),
            ("Global planner", "5", "NavFn A*, NavFn Dijkstra, Theta*, Smac 2D, Smac Hybrid"),
            ("Đường/phương pháp", "8", "Raw, Simple, SG, Constrained, Pivot fixed/adaptive, Hybrid fixed/adaptive"),
            ("Lần lặp", "3", "Mọi smoother trong một nhóm nhận cùng Raw"),
            ("Nhóm đầu vào", "900", "60 tình huống × 5 planner × 3 lần lặp"),
            ("Bản ghi hình học", "7.200", "900 nhóm × 8 phương pháp"),
        ],
        [4.2, 2.1, 9.7],
    )
    callout(
        doc,
        "Cách gọi chính xác:",
        "7.200 là số bản ghi đường đi/quỹ đạo hình học sau lập kế hoạch và làm "
        "mượt, không phải 7.200 lượt robot thực thi trong Gazebo.",
        fill=LIGHT_BLUE,
    )
    doc.add_heading("6. Điểm đánh giá và đặc tính dự kiến thu được", level=1)
    table(
        doc,
        ["Nhóm", "Chỉ tiêu", "Đặc tính cần rút ra"],
        [
            ("Tạo đường", "Tỷ lệ trả kết quả; timeout", "Độ bền trên nhiều dạng đầu vào"),
            ("An toàn", "Va chạm footprint; clearance; peak cost", "Khả năng không cắt vật cản"),
            ("Độ trơn", "κmax; ∫κ²ds; biến thiên κ", "Mức giảm góc gấp và dao động"),
            ("Thay đổi đường", "Chiều dài; lệch max/RMS so với Raw", "Đánh đổi giữa trơn và giữ hành lang"),
            ("Quyết định", "Số Pivot/G²; d được chọn; nhánh Hybrid", "Mức thích nghi theo không gian"),
            ("Tính toán", "Runtime trung bình, P95, cực đại", "Khả năng dùng trong Nav2"),
            ("Công bằng", "Raw hash; cùng planner/scenario/repetition", "Các smoother nhận cùng đầu vào"),
        ],
        [2.8, 6.1, 7.1],
    )
    callout(
        doc,
        "Lưu ý:",
        "Tích phân bình phương độ cong là thước đo effort hình học, không phải "
        "năng lượng điện hoặc mức tiêu thụ pin.",
    )

    page_break(doc)
    doc.add_heading("7. Ảnh môi trường mô phỏng Gazebo", level=1)
    figure(
        doc,
        "results/gui_validation_20260724/gazebo_research_warehouse.png",
        "Hình 3. Kho nghiên cứu trong Gazebo; world này được ghép với map/costmap tương ứng để kiểm tra smoother.",
        12.0,
    )
    figure(
        doc,
        "results/gui_validation_20260724/gazebo_warehouse_cross_aisles_active.png",
        "Hình 4. Kho giao cắt dùng tạo các đường có nhiều lần đổi hướng giữa các dãy kệ.",
        8.2,
    )

    page_break(doc)
    figure(
        doc,
        "results/gui_validation_20260724/gazebo_warehouse_dispatch_final.png",
        "Hình 5. Kho điều phối có kệ và kiện hàng, dùng kiểm tra làm mượt gần nhiều loại vật cản.",
        9.2,
    )
    figure(
        doc,
        "docs/bao_cao_toan_dien_assets/map_narrow_aisles.png",
        "Hình 6. Môi trường lối hẹp đại diện cho trường hợp phải cân nhắc giữa đoạn G² và trạng thái Pivot.",
        13.2,
    )

    page_break(doc)
    doc.add_heading("8. Ảnh quan sát và so sánh đường trong RViz2", level=1)
    figure(
        doc,
        "results/neutral_hybrid_20260727/rviz_neutral_selector.png",
        "Hình 7. RViz2 chồng đường Raw và các kết quả smoother trên cùng costmap; panel bên phải chọn planner và từng phương pháp.",
    )
    figure(
        doc,
        "docs/rev_ecit_2026_assets/figure_10_rviz_all_methods.png",
        "Hình 8. Gazebo và RViz2 hiển thị đồng thời để đối chiếu world, costmap và các đường sau xử lý.",
    )

    page_break(doc)
    doc.add_heading("9. Kết quả dự kiến và giới hạn phát biểu", level=1)
    paragraph(
        doc,
        "Kết quả kỳ vọng là nhánh G² giảm độ gấp và dao động độ cong tại những "
        "góc có đủ không gian; nhánh Pivot giữ khả năng xử lý các góc hẹp; kiểm "
        "tra swept-footprint loại các đoạn cong có nguy cơ cắt vật cản; và Raw "
        "fallback giúp hệ thống vẫn trả đường khi các ứng viên làm mượt không "
        "hợp lệ. Dữ liệu cũng dự kiến cho thấy sự đánh đổi: độ trơn có thể tốt "
        "hơn nhưng độ lệch khỏi Raw hoặc thời gian tính toán có thể tăng."
    )
    paragraph(
        doc,
        "Không nên tuyên bố thuật toán luôn tạo đường ngắn nhất, nhanh nhất hoặc "
        "tối ưu toàn cục. Abstract cuối chỉ nên đưa phần trăm cải thiện sau khi "
        "khóa phiên bản code, chốt trọng số, chạy lại đầy đủ ma trận và báo cáo "
        "cả các trường hợp thất bại."
    )
    doc.add_heading("10. Việc cần hoàn thiện trước abstract cuối", level=1)
    bullets(
        doc,
        [
            "Chốt tiêu chí góc cần xử lý và miền tìm kiếm d.",
            "Chốt cách lấy mẫu footprint và xử lý ô unknown.",
            "Chốt công thức chi phí và trọng số của quy hoạch động.",
            "Chốt luật so sánh các nhánh trong Hybrid.",
            "Khóa code/config/seed rồi chạy lại toàn bộ dữ liệu.",
            "Không trộn dữ liệu pilot của các phiên bản thuật toán khác nhau.",
            "Báo cáo cả failure và timeout, không chỉ các đường thành công.",
        ],
    )
    callout(
        doc,
        "Câu kết gợi ý khi gửi cô:",
        "Em gửi cô bản mô tả sơ bộ để cô xem giúp hướng thuật toán smoother, "
        "cách bố trí thử nghiệm và các đặc tính em dự kiến đánh giá. Hiện em "
        "chưa chốt các con số cải thiện vì phần sinh ứng viên và luật chọn "
        "Hybrid vẫn đang được hiệu chỉnh. Sau khi cô góp ý về phạm vi, em sẽ "
        "khóa cấu hình, chạy lại ma trận và cập nhật kết quả định lượng.",
        fill=LIGHT_BLUE,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    create()
