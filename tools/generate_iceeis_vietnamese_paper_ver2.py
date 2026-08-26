#!/usr/bin/env python3
"""Generate the complete Vietnamese PSTMO ver2 paper in the IEEE A4 layout.

The manuscript keeps the terminology, mathematical model, method scope, and
recorded measurements of ver1.  Its argument and section order follow the
teacher-reviewed ver2 outline: related work is integrated into Introduction,
the compact model is part of the proposed method, and each experimental
scenario is introduced before its evidence.
"""

from __future__ import annotations

import subprocess
import shutil
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
import generate_iceeis_vietnamese_paper as base


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "final_bao_ICEEIS/ver2"
ASSET_DIR = OUT_DIR / "assets"
OUTPUT_DOCX = OUT_DIR / "ICEEIS_2026_PSTMO_ver2_tieng_Viet.docx"
ABSTRACT_TEMPLATE = ROOT / "abstract/abstract final 2.docx"
VER1_FIGURE_DIR = ROOT / "final_bao_ICEEIS/ver1/ICEEIS_2026_PSTMO_drawio_assets"


def body(doc, text: str, first: bool = False):
    return base.add_body(doc, text, first=first)


def h1(doc, text: str):
    return base.add_heading1(doc, text)


def h2(doc, text: str):
    return base.add_heading2(doc, text)


def eq(doc, text, number: int):
    return base.add_equation(doc, text, number)


def fig(doc, path: Path, width: float, caption: str, number: int):
    return base.add_figure(doc, path, width, caption, number)


def compact_table(doc, headers, rows, widths, font_size=7.15):
    table = base.add_ieee_table(doc, headers, rows, widths=widths, font_size=font_size)
    for row in table.rows:
        for cell in row.cells:
            base.set_cell_margins(cell, top=25, start=38, bottom=25, end=38)
    return table


def prepare_assets(metrics) -> dict[str, Path]:
    """Copy, byte for byte, the revised figures already delivered in ver1."""
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    keys = ("corner", "drive", "swept", "bezier", "edge", "open", "narrow", "warehouse")
    assets = {}
    for index, key in enumerate(keys, 1):
        source = VER1_FIGURE_DIR / f"image{index}.png"
        destination = ASSET_DIR / f"ver1_hinh_{index}.png"
        shutil.copyfile(source, destination)
        assets[key] = destination
    return assets


def configure_title_and_abstract(doc: Document) -> None:
    title = next(p for p in doc.paragraphs if p.text.startswith("Path Smoothing via"))
    base.clear_paragraph_content(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run(
        "Làm mượt đường đi bằng tối ưu chuyển tiếp tại góc "
        "có xét hình bao cho robot di động vi sai"
    )
    base.set_run_font(run, 24)

    abstract_text = (
        "Đường do bộ lập kế hoạch toàn cục sinh ra có thể là đa tuyến gồm các đoạn thẳng liên tiếp; dù không phát hiện "
        "va chạm trên bản đồ chi phí, đường này chưa bảo đảm chuyển hướng liên tục hoặc khả thi với robot vi sai. Bài báo "
        "đề xuất PSTMO, một bộ hậu xử lý thay lân cận góc bằng đoạn chuyển tiếp Bézier bậc năm liên tục hình học bậc hai "
        "(G²). Ứng viên được loại theo độ cong, giới hạn vận tốc thân và hai bánh, gia tốc dọc/ngang, gia tốc góc và vùng "
        "quét hình bao; "
        "sau đó quy hoạch động chọn chuỗi xử lý không chồng lấn trên toàn đường. PSTMO được đánh giá trong ROS 2 "
        "Jazzy/Navigation2 và Gazebo Harmonic trên ba bản đồ tĩnh với năm bộ lập kế hoạch, tạo 15 nhóm đầu vào ghép cặp. "
        "So với Simple và Constrained, PSTMO giảm trung bình tích phân bình phương độ cong lần lượt 75,42% và 90,75%, "
        "đồng thời giảm chiều dài 0,98% và 2,01%. Giá trị trung bình của PSTMO là 8,645 m, 2,887 m⁻¹ và 96,0 ms cho "
        "chiều dài, tích phân bình phương độ cong và thời gian xử lý. Kết quả cho thấy cải thiện hình học rõ rệt nhưng "
        "đổi lại chi phí xử lý cao hơn các bộ làm mượt đối chứng."
    )
    abstract = next(p for p in doc.paragraphs if p.text.startswith("Abstract—"))
    base.clear_paragraph_content(abstract)
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.first_line_indent = Inches(0)
    abstract.paragraph_format.space_after = Pt(3)
    run = abstract.add_run("Tóm tắt—")
    base.set_run_font(run, 9, bold=True, italic=True)
    run = abstract.add_run(abstract_text)
    base.set_run_font(run, 9, bold=True)

    keywords = next(p for p in doc.paragraphs if p.text.startswith("Keywords—"))
    base.clear_paragraph_content(keywords)
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.first_line_indent = Inches(0)
    keywords.paragraph_format.space_after = Pt(4)
    run = keywords.add_run("Từ khóa—")
    base.set_run_font(run, 9, bold=True, italic=True)
    run = keywords.add_run(
        "robot vi sai, làm mượt đường đi, Bézier bậc năm, hình bao robot, Navigation2, PSTMO."
    )
    base.set_run_font(run, 9, italic=True)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    base.configure_section(body_section, columns=2)
    section_break_paragraph = doc.paragraphs[-1]
    abstract._p.addprevious(section_break_paragraph._p)


def add_introduction(doc: Document, assets: dict[str, Path]) -> None:
    h1(doc, "I. GIỚI THIỆU")
    body(doc, (
        "Robot tự hành dẫn đường và robot di động tự hành ngày càng được dùng trong kho vận, sản xuất và vận chuyển nội "
        "bộ, nơi chuyển động tự chủ phải đồng thời an toàn và ổn định. Trong ROS 2 Navigation2 (Nav2), bộ lập kế hoạch "
        "toàn cục tạo đường, Smoother Server có thể cải thiện hình học của đường, và bộ điều khiển sinh lệnh vận tốc để "
        "bám đường [1]. Vì vậy, đầu vào của bộ điều khiển không chỉ cần tránh vật cản mà còn cần có hình học phù hợp với "
        "khả năng chuyển động của robot."), first=True)
    body(doc, (
        "Dijkstra tìm đường chi phí nhỏ nhất trên đồ thị [2], còn A* dùng heuristic để hướng quá trình tìm kiếm [3]. Trên "
        "lưới chiếm chỗ, Theta* cho phép nối any-angle có kiểm tra đường nhìn nhằm giảm sự phụ thuộc vào các hướng lưới "
        "[4]. Hybrid A* mở rộng trạng thái tìm kiếm bằng hướng và mô hình chuyển động để tạo đường phù hợp hơn cho xe có "
        "ràng buộc động học [5]. Trong thí nghiệm của bài báo, NavFn A*, NavFn Dijkstra, Theta*, Smac2D và SmacHybrid được "
        "dùng để tạo năm dạng đường đầu vào. Các planner này giải quyết nhiệm vụ tìm tuyến hợp lệ hoặc có chi phí thấp theo "
        "mô hình riêng; tuy nhiên, đường rời rạc sau lấy mẫu vẫn có thể chứa các đỉnh đổi hướng cần hậu xử lý trước khi bám."))
    body(doc, (
        "Tại một đỉnh của đa tuyến, hướng tiếp tuyến thay đổi gián đoạn. Nếu robot vi sai đi qua đỉnh với vận tốc tịnh tiến "
        "khác không, yêu cầu vận tốc góc và chênh lệch vận tốc hai bánh có thể tăng mạnh; bộ điều khiển khi đó có thể phải "
        "giảm tốc, lệch khỏi đường hoặc dừng để quay tại chỗ. Do đó, bộ làm mượt phải phân bố biến thiên hướng trên một đoạn "
        "hữu hạn, nhưng không được đổi tuyến toàn cục hoặc tạo va chạm mới. Hình 1 minh họa đúng phạm vi biến đổi cục bộ này."))
    fig(doc, assets["corner"], 3.12,
        "Đỉnh gãy của đường đầu vào và đoạn chuyển tiếp cục bộ phân bố biến thiên hướng trên một đoạn hữu hạn.", 1)
    body(doc, (
        "Các kỹ thuật làm mượt đường đi có thể được phân loại theo biểu diễn hình học, cơ chế biến dạng/tối ưu không gian và "
        "tối ưu quỹ đạo không–thời gian [6]. Ở nhóm thứ nhất, Fleury và cộng sự dùng clothoid cùng anticlothoid làm phần tử "
        "nối [7]; Brezak và Petrović phát triển làm mượt clothoid trực tuyến cho robot vi sai [8]. Clothoid có quy luật độ "
        "cong tuyến tính theo chiều dài cung nhưng thường cần tích phân Fresnel, xấp xỉ hoặc giải số, đồng thời vẫn cần một "
        "cơ chế riêng để kiểm tra khoảng hở và hình bao robot."))
    body(doc, (
        "Đường Bézier có biểu thức đóng, đạo hàm dễ tính và điểm điều khiển trực quan. Yang và Sukkarieh xây dựng thuật toán "
        "Bézier bậc ba giải tích với độ cong liên tục và giới hạn độ cong [9]; Bu và cộng sự phát triển lượt rẽ Bézier liên "
        "tục độ cong cho xe car-like [10]. Xu và cộng sự ghép các đoạn Bézier liên tục cho đường toàn cục [11], còn Parque "
        "và Miyashita dùng Differential Evolution để khớp đường cong [12]. Các kết quả này xác nhận khả năng tạo hình học "
        "trơn của Bézier, nhưng tính liên tục phụ thuộc cấu trúc điểm điều khiển và điều kiện biên; làm mượt đường tâm cũng "
        "không tự bảo đảm vùng quét của toàn thân robot không giao vật cản."))
    body(doc, (
        "Ở nhóm biến dạng và tối ưu không gian, Elastic Bands co đường bằng lực nội tại và lực đẩy vật cản [13]. Bộ lọc "
        "Savitzky–Golay bắt nguồn từ xấp xỉ đa thức bình phương tối thiểu cục bộ [14]. Trong Nav2, Simple Smoother cập nhật "
        "lặp từng điểm bằng hai thành phần bám đường gốc và làm trơn lân cận [15]; Savitzky–Golay Smoother áp bộ lọc cục bộ "
        "[16]; Constrained Smoother xây dựng bài toán tối ưu Ceres với các thành phần độ trơn, costmap, bám đường và bán kính "
        "quay [17]. Ba plugin mã nguồn mở này đại diện lần lượt cho làm mượt lặp chi phí thấp, lọc cục bộ và tối ưu có ràng "
        "buộc. Đổi lại, bộ lọc nhanh không tự bảo đảm mọi giới hạn động học hoặc swept footprint, còn nghiệm tối ưu phụ thuộc "
        "hàm mục tiêu, trọng số, khởi tạo và rời rạc hóa."))
    body(doc, (
        "CHOMP tối ưu đồng thời độ trơn và chi phí vật cản bằng gradient hiệp biến [18]. GRIPS kết hợp gradient trường khoảng "
        "cách với nối tắt có kiểm tra [19], trong khi Andreasson và cộng sự tối ưu trạng thái liên tục trong các miền lồi "
        "không va chạm [20]. Các phương pháp này linh hoạt hơn phép lọc cục bộ nhưng cần thông tin môi trường và giải một bài "
        "toán tối ưu lớn hơn. Ở lớp rộng hơn, TEB đưa khoảng thời gian giữa các trạng thái vào dải đàn hồi [21], sau đó được "
        "phát triển cho điều khiển dự đoán tối ưu thời gian [22] và tối ưu kinodynamic cho xe car-like [23]. Các phương pháp "
        "không–thời gian có thể xét đồng thời vận tốc, gia tốc và vật cản, nhưng đầu ra và mục tiêu rộng hơn một bộ hậu xử lý "
        "đường hình học; PSTMO không nhằm thay thế lớp trajectory optimizer này."))
    body(doc, (
        "Như vậy, phương pháp hình học có cấu trúc và dễ giải thích nhưng cần cổng an toàn; bộ lọc có chi phí thấp nhưng ít "
        "bảo đảm tường minh; tối ưu toàn đường linh hoạt nhưng tăng số biến và độ nhạy tham số; còn tối ưu không–thời gian giải "
        "một bài toán rộng hơn. Trong tập tài liệu được khảo sát, chúng tôi chưa thấy một bộ hậu xử lý Nav2 cho robot vi sai "
        "phối hợp đồng thời: (i) mối nối thẳng–cong có độ cong đầu cuối bằng không; (ii) giới hạn hai bánh và kiểm tra vùng quét "
        "hình bao; và (iii) lựa chọn tương thích giữa các góc kề nhau trong khi vẫn giữ tuyến toàn cục."))
    body(doc, (
        "Bài báo đề xuất Path Smoothing via Footprint-Aware Corner-Transition Optimization for Differential-Drive Mobile "
        "Robots (PSTMO), đặt giữa global planner và controller trong Nav2. PSTMO tạo các chuyển tiếp Bézier bậc năm G², loại "
        "ứng viên không khả thi và dùng quy hoạch động để chọn chuỗi xử lý không chồng lấn. Phạm vi đánh giá là costmap tĩnh, "
        "mô phỏng ROS 2 Jazzy/Nav2 và Gazebo Harmonic, ba môi trường quy mô nhỏ và năm nguồn đường toàn cục; robot vật lý, "
        "vật cản động, sai số bám đường, năng lượng, nhiễu, trượt, tải trọng và phần cứng nhúng nằm ngoài phạm vi hiện tại."))
    body(doc, (
        "Ba đóng góp chính là: (1) cấu trúc Bézier bậc năm tạo mối nối thẳng–cong G² với độ cong bằng không ở hai đầu; "
        "(2) cổng khả thi kết hợp giới hạn chuyển động robot vi sai, swept footprint và ngân sách cạnh chung, cùng quy hoạch "
        "động chọn phương án trên toàn chuỗi góc; và (3) đánh giá ghép cặp Raw, Simple, Savitzky–Golay, Constrained và PSTMO "
        "trên 15 tổ hợp bản đồ–planner bằng chiều dài L, tích phân bình phương độ cong Eκ và thời gian xử lý T. PSTMO đạt Eκ "
        "trung bình thấp nhất mà không làm tăng chiều dài trung bình, đổi lại thời gian xử lý cao hơn các đối chứng. Phần II "
        "trình bày mô hình và phương pháp; Phần III mô tả thử nghiệm và thảo luận; Phần IV kết luận và nêu giới hạn."))


def add_method(doc: Document, assets: dict[str, Path]) -> None:
    h1(doc, "II. PHƯƠNG PHÁP PSTMO ĐỀ XUẤT")
    h2(doc, "A. Bài toán và mô hình toán học rút gọn")
    body(doc, (
        "Trạng thái robot là z=(x,y,ψ), trong đó (x,y) là vị trí tâm hình học và ψ là góc hướng trong hệ quy chiếu bản đồ. "
        "Với vận tốc tịnh tiến v và vận tốc góc ω, động học robot vi sai được viết như (1). Phương trình này liên hệ hướng "
        "tiếp tuyến của đường với chuyển động thân robot."), first=True)
    eq(doc, "ẋ = v cos ψ,   ẏ = v sin ψ,   ψ̇ = ω.", 1)
    body(doc, (
        "Đường đầu vào là P={p₀,…,pₙ}, với pᵢ=(rᵢ,ψᵢ) và rᵢ=(xᵢ,yᵢ). Chiều dài được tính trên tọa độ phẳng theo (2). "
        "Một điểm trong là góc cần xem xét khi hai vectơ đơn vị của cạnh vào u và cạnh ra w tạo góc chuyển hướng có dấu "
        "θᵢ=atan2(uₓwᵧ−uᵧwₓ,u·w) thỏa |θᵢ|≥θmin; các góc nhỏ hơn ngưỡng được giữ nguyên."))
    eq(doc, "L(P) = Σⁿ⁻¹ᵢ₌₀ ‖rᵢ₊₁ − rᵢ‖₂.", 2)
    body(doc, (
        "Với đường tham số r(t)=(x(t),y(t)), độ cong có dấu được xác định bởi (3). Hai đại lượng hình học trong (4) là "
        "độ cong lớn nhất κmax và tích phân bình phương độ cong Eκ. Eκ có đơn vị m⁻¹, phản ánh mức uốn tích lũy theo giao "
        "thức đánh giá và không phải năng lượng tiêu thụ."))
    eq(doc, "κ = (x′y″ − y′x″)/(x′² + y′²)^(3/2).", 3)
    eq(doc, ["κmax = maxₛ∈[0,L] |κ(s)|;", "Eκ = ∫₀ᴸ κ²(s) ds."], 4)
    body(doc, (
        "Liên tục hình học bậc một G¹ bảo đảm hướng tiếp tuyến trùng nhau; G² yêu cầu thêm độ cong liên tục. Vì đoạn thẳng "
        "có κ=0, đoạn chuyển tiếp phải có độ cong bằng không tại hai đầu. Với b là khoảng cách hai bánh chủ động, quan hệ "
        "giữa độ cong và vận tốc trong (5) được dùng để kiểm tra vận tốc góc, vận tốc từng bánh và gia tốc ngang của ứng viên; "
        "profile vận tốc nội bộ tiếp tục kiểm tra gia tốc dọc và gia tốc góc."))
    eq(doc, ["vL = v(1 − bκ/2),   vR = v(1 + bκ/2);", "ω = vκ."], 5)
    fig(doc, assets["drive"], 2.64,
        "Mô hình động học robot hai bánh vi sai và quy ước vận tốc tuyến tính tại bánh.", 2)
    body(doc, (
        "Gọi Fbody là đa giác hình bao trong hệ thân xe. Tại z, hình bao trên bản đồ là Fmap(z)=[x,y]ᵀ+R(ψ)Fbody; hợp các "
        "hình bao được lấy mẫu dọc đường tạo swept footprint."))
    fig(doc, assets["swept"], 2.82,
        "Vùng quét được tạo từ hợp các hình bao robot lấy mẫu dọc đường chuyển động.", 3)
    body(doc, (
        "Đầu vào của bài toán gồm P, costmap M, Fbody và tập giới hạn "
        "chuyển động. Cần tìm P* giữ vị trí đầu, vị trí đích và orientation tại đích, không đổi thứ tự hành lang, nối G² tại "
        "các vùng được thay, thỏa (5), không phát hiện va chạm theo swept footprint và không chồng lấn giữa các góc. Các điều "
        "kiện này xác định miền khả thi F; trong F, Eκ và L được dùng để so sánh chất lượng, còn T được báo cáo riêng thay vì "
        "gộp ba đại lượng khác đơn vị vào một khẳng định tối ưu duy nhất."))

    h2(doc, "B. Chuyển tiếp Bézier bậc năm")
    body(doc, (
        "PSTMO chỉ sửa lân cận góc nhưng lựa chọn phương án ở cấp toàn đường. Sau bước điều kiện hóa an toàn và phát hiện "
        "góc, mỗi góc có thể giữ nguyên, quay tại chỗ hoặc dùng một chuyển tiếp tịnh tiến. Xét đỉnh V, u và w lần lượt là "
        "vectơ đơn vị của cạnh vào và ra. Hai điểm tiếp giáp A=V−du và B=V+dw được xác định bởi khoảng cắt d. Đặt q=αd, "
        "trong đó α là tỷ lệ hình dạng không thứ nguyên. Sáu điểm điều khiển và đường Bézier bậc năm được cho bởi (6)–(7)."),
        first=True)
    eq(doc, [
        "P₀=A,  P₁=A+qu,  P₂=A+2qu;",
        "P₃=B−2qw,  P₄=B−qw,  P₅=B."
    ], 6)
    eq(doc, "B(t)=Σ⁵ᵢ₌₀ C(5,i)(1−t)⁵⁻ⁱtⁱPᵢ,   0≤t≤1.", 7)
    fig(doc, assets["bezier"], 3.02,
        "Cấu trúc sáu điểm điều khiển của đoạn chuyển tiếp Bézier bậc năm tại một góc.", 4)
    body(doc, (
        "Từ (6), B′(0)=5qu và B′(1)=5qw nên tiếp tuyến hai đầu cùng hướng với hai cạnh; đồng thời B″(0)=B″(1)=0. Do đó "
        "κ(0)=κ(1)=0 và đoạn Bézier nối G² với hai đoạn thẳng. Tham số d quyết định chiều dài cạnh dành cho chuyển tiếp và "
        "bị giới hạn bởi vật cản cũng như góc kề; α điều chỉnh phân bố độ cong khi d cố định. PSTMO suy ra tối đa hai giá trị "
        "d đại diện, gồm giá trị ưu tiên theo hình học và giá trị tương thích với ngân sách cạnh chung. Với mỗi d, α được khảo "
        "sát theo lưới nhiều mức; chỉ các cặp (d,α) vượt qua toàn bộ điều kiện bắt buộc mới được giữ lại."))

    h2(doc, "C. Cổng khả thi và lựa chọn tương thích toàn đường")
    body(doc, (
        "Ứng viên trước hết phải có đạo hàm không suy biến và độ cong hữu hạn. Từ κ và v, thuật toán suy ra ω, vL và vR "
        "theo (5), rồi loại ứng viên vi phạm giới hạn vận tốc thân, vận tốc góc, vận tốc bánh hoặc gia tốc ngang. Một profile "
        "vận tốc nội bộ được tham số hóa trên ứng viên để kiểm tra thêm gia tốc dọc, giảm tốc dọc và gia tốc góc; profile này "
        "chỉ phục vụ cổng khả thi và không phải đầu ra của smoother. Các mẫu có đảo dấu độ cong ngoài ý muốn, yêu cầu bánh "
        "trong quay ngược trong trạng thái chuyển tiếp tịnh tiến, hoặc không còn vận tốc khả thi dương cũng bị loại. Tiếp theo, "
        "Fbody được đặt theo vị trí và góc hướng dọc ứng viên; bất kỳ giao cắt "
        "nào với ô cấm trên M đều làm ứng viên không khả thi."), first=True)
    body(doc, (
        "Hai góc kề nhau có thể cùng chiếm một cạnh. Với dᵢ và dᵢ₊₁ là phần cạnh chung bị chiếm bởi hai phương án, m là "
        "biên không chồng lấn và Lᵢ là chiều dài cạnh chung, hai phương án chỉ tương thích khi thỏa (8). Đây là ngân sách "
        "hình học trên cạnh, không phải biên khoảng hở vật cản."))
    eq(doc, "dᵢ + dᵢ₊₁ + m ≤ Lᵢ.", 8)
    fig(doc, assets["edge"], 3.02,
        "Ngân sách cạnh chung dùng để kiểm tra tính tương thích giữa hai góc kề nhau.", 5)
    body(doc, (
        "Mỗi phương án tại góc i là một trạng thái k có chi phí cục bộ Jᵢ(k), chỉ được tính sau cổng khả thi. Chi phí này "
        "tổng hợp các thành phần đã chuẩn hóa về costmap, vận tốc góc cực đại và Eκ của ứng viên. Gọi Cᵢ(k) là tập trạng "
        "thái ở góc i−1 tương thích với k theo (8), quy hoạch động dùng truy hồi (9), lưu trạng thái trước và truy vết từ "
        "trạng thái cuối có chi phí nhỏ nhất. Với N góc và tối đa K trạng thái mỗi góc, bước lựa chọn có độ phức tạp O(NK²)."))
    eq(doc, "F₁(k)=J₁(k),    Fᵢ(k)=Jᵢ(k)+minⱼ∈Cᵢ(k) Fᵢ₋₁(j).", 9)

    h2(doc, "D. Ghép và hậu kiểm đường đầu ra")
    body(doc, (
        "Chuỗi trạng thái được ghép theo đúng thứ tự từ các đoạn thẳng còn lại, chuyển tiếp Bézier và trạng thái quay tại "
        "chỗ, sau đó lấy mẫu lại với khoảng cách đầu ra cấu hình. Vị trí đầu, vị trí đích và orientation tại đích được bảo "
        "toàn; orientation đầu được suy theo đoạn chuyển động đầu, nên bài báo không tuyên bố bảo toàn orientation tại start. "
        "Một hậu kiểm độc lập trên toàn đường xác nhận giá trị hữu hạn, đầu mút, thứ tự tuyến, các giới hạn chuyển động, ngân "
        "sách cạnh chung và swept footprint. Nếu bất kỳ bất biến nào không thỏa, PSTMO báo thất bại thay vì trả một đường "
        "làm mượt không an toàn. Đầu ra vẫn là nav_msgs/Path hình học; các giới hạn chuyển động chỉ dùng để sàng lọc và không "
        "biến đầu ra thành quỹ đạo đã tham số hóa theo thời gian."), first=True)


def format_metric(value: float, digits: int) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def metric_rows(metrics, env: str):
    labels = {
        "raw": "Raw",
        "simple": "Simple",
        "savitzky_golay": "Savitzky–Golay",
        "constrained": "Constrained",
        "pstmo": "PSTMO",
    }
    rows = []
    for method in ("raw", "simple", "savitzky_golay", "constrained", "pstmo"):
        m = metrics[env][method]
        rows.append((labels[method], format_metric(m["L"], 3), format_metric(m["E"], 3), format_metric(m["T"], 1)))
    return rows


def add_experiments(doc: Document, assets: dict[str, Path], metrics) -> None:
    h1(doc, "III. ĐÁNH GIÁ THỰC NGHIỆM")
    h2(doc, "A. Điều kiện, phương án so sánh và giao thức ghép cặp")
    body(doc, (
        "Thử nghiệm được thực hiện trên Ubuntu 24.04, ROS 2 Jazzy, Nav2 và Gazebo Harmonic với CPU Intel Core "
        "i5-12450HX. Để mọi phương pháp được đánh giá dưới cùng điều kiện hình học và chuyển động, các thông số trong Bảng I "
        "được giữ cố định. Robot có hình bao chữ nhật 0,44×0,34 m. Khoảng cách bánh 0,2548 m được dùng trong mô hình và cổng "
        "động học; plugin DiffDrive của Gazebo dùng giá trị hiệu dụng 0,2834 m sau hiệu chỉnh tiếp xúc bánh–mặt sàn và odometry."),
        first=True)
    base.add_table_caption(doc, "I", "Thông số chính của hệ thống thử nghiệm")
    setup_rows = [
        ("Nền tảng", "Ubuntu 24.04; ROS 2 Jazzy; Nav2; Gazebo Harmonic"),
        ("CPU", "Intel Core i5-12450HX"),
        ("Costmap", "0,05 m/ô"),
        ("Hình bao", "Hình chữ nhật 0,44×0,34 m"),
        ("Khoảng cách bánh", "Mô hình 0,2548 m; hiệu dụng Gazebo 0,2834 m"),
        ("Giới hạn vận tốc", "vmax=0,30 m/s; |ω|max=0,80 rad/s; |vbánh|max=0,36 m/s"),
        ("Giới hạn gia tốc", "dọc +0,35/−0,45 m/s²; ngang 0,18 m/s²; góc 1,20 rad/s²"),
        ("Khoảng mẫu", "Kiểm tra 0,02 m; đầu ra 0,05 m"),
    ]
    compact_table(doc, ["Thành phần", "Giá trị"], setup_rows, [1.08, 2.17], font_size=6.85)
    body(doc, (
        "Năm phương án xử lý trong Bảng II gồm Raw và bốn smoother. Raw là mốc đầu vào, không phải bộ làm mượt. Simple là "
        "đối chứng lặp chi phí thấp; Savitzky–Golay là đối chứng lọc cục bộ; Constrained là đối chứng tối ưu có ràng buộc và "
        "costmap; PSTMO là phương pháp đề xuất. Đặc biệt, trong cấu hình Constrained đã thử, wcurve=wdist=0 nên Eκ không "
        "phải thành phần được tối thiểu hóa trực tiếp. Đây là so sánh giữa các cấu hình cố định đã thử, không phải hiệu năng "
        "cực đại sau khi hiệu chỉnh riêng từng phương pháp."))
    base.add_table_caption(doc, "II", "Các phương án được dùng trong phép so sánh")
    baseline_rows = [
        ("Raw", "Đường từ planner, chưa làm mượt", "Mốc đầu vào"),
        ("Simple", "max_its=1000; refinement=2", "Lặp chi phí thấp"),
        ("Savitzky–Golay", "Cửa sổ 7 điểm; refinement=2", "Lọc cục bộ"),
        ("Constrained", "wsmooth=200000; wcost=0,015; wcurve=wdist=0", "Tối ưu có ràng buộc"),
        ("PSTMO", "Bézier G²; động học; footprint; DP", "Đề xuất"),
    ]
    compact_table(doc, ["Phương án", "Cấu hình/nguyên lý", "Vai trò"], baseline_rows,
                  [0.82, 1.61, 0.82], font_size=6.55)
    body(doc, (
        "Trong mỗi môi trường, NavFn A*, NavFn Dijkstra, Theta*, Smac2D và SmacHybrid tạo năm đường Raw. Bốn smoother nhận "
        "đúng cùng một Raw path, costmap và tập giới hạn trong từng nhóm. Ba môi trường nhân năm planner tạo 15 nhóm đầu vào "
        "ghép cặp; mỗi tổ hợp chỉ chạy một lần và cả năm phương án đều tạo đầu ra. Vì vậy, kết quả là thống kê mô tả, không "
        "phải 15 lần lặp độc lập và không được dùng để lập khoảng tin cậy."))

    h2(doc, "B. Chỉ số đánh giá")
    body(doc, (
        "Ba chỉ số duy nhất là chiều dài L theo (2), tích phân bình phương độ cong Eκ theo (4) và thời gian xử lý thuật toán "
        "T từ lúc nhận đến lúc trả đường. Các đường có mật độ điểm khác nhau được lấy mẫu lại theo cùng chiều dài cung trước "
        "khi ước lượng độ cong rời rạc và tích phân số. Thời gian Raw chỉ phản ánh chuyển tiếp/sao chép trong cùng bộ đo, không "
        "gồm global planning. Bộ đo có lượng tử khoảng 3 ms; giá trị trung bình dưới mức này có nghĩa là dưới độ phân giải "
        "đáng tin cậy, không phải chi phí bằng không."), first=True)

    h2(doc, "C. Kịch bản 1—Không gian mở")
    body(doc, (
        "Không gian mở có một khối cản trung tâm; robot đi từ (−2,20;−0,60) m đến (1,20;−0,60) m. Kịch bản này cho phép "
        "quan sát chuyển tiếp khi vùng lân cận góc tương đối rộng. Hình 6 minh họa ca Theta*; các số liệu của Bảng III là "
        "trung bình trên năm planner, không phải riêng ca minh họa."), first=True)
    fig(doc, assets["open"], 3.04,
        "Ca Theta* trong không gian mở; đường Raw, PSTMO và các tư thế hình bao được đặt trên cùng bản đồ chi phí.", 6)
    base.add_table_caption(doc, "III", "Kết quả trung bình—Không gian mở")
    compact_table(doc, ["Phương pháp", "L (m)", "Eκ (m⁻¹)", "T (ms)"], metric_rows(metrics, "open_arena"),
                  [1.25, 0.64, 0.78, 0.58], font_size=7.15)
    body(doc, (
        "PSTMO đạt Eκ=1,954 m⁻¹, thấp nhất trong môi trường này và thấp hơn Simple 84,36%. Chiều dài PSTMO là 4,224 m, "
        "không tăng so với Raw hoặc ba smoother; đổi lại thời gian 54,0 ms cao hơn Constrained 10,2 ms và các đối chứng "
        "nhanh. Kết quả phù hợp với việc không gian rộng cho phép dùng chuyển tiếp dài hơn, nhưng chưa đủ để suy ra quan hệ "
        "nhân quả giữa khoảng hở và chất lượng."))

    h2(doc, "D. Kịch bản 2—Lối đi hẹp")
    body(doc, (
        "Lối đi hẹp chạy theo đường chéo tây nam–đông bắc, với điểm đầu (−5,00;−3,00) m và điểm đích (5,00;3,00) m. Các "
        "góc liên tiếp phải chia sẻ chiều dài cạnh và không gian swept footprint bị hạn chế hơn; đây là tình huống trực tiếp "
        "tác động đến cổng khả thi và điều kiện (8). Hình 7 cho thấy ca Theta* đại diện."), first=True)
    fig(doc, assets["narrow"], 3.04,
        "Ca Theta* trong lối đi hẹp; hình bao được dùng cho kiểm tra va chạm nhị phân trên costmap cấu hình.", 7)
    base.add_table_caption(doc, "IV", "Kết quả trung bình—Lối đi hẹp")
    compact_table(doc, ["Phương pháp", "L (m)", "Eκ (m⁻¹)", "T (ms)"], metric_rows(metrics, "narrow_aisles"),
                  [1.25, 0.64, 0.78, 0.58], font_size=7.15)
    body(doc, (
        "PSTMO vẫn thấp nhất với Eκ=4,326 m⁻¹, nhưng mức giảm so với Simple còn 51,03%, nhỏ hơn hai môi trường còn lại. "
        "Đồng thời, thời gian 154,2 ms là lớn nhất trong ba kịch bản. Điều này nhất quán với không gian ứng viên bị thu hẹp và "
        "nhiều quan hệ cạnh chung phải đánh giá, nhưng thí nghiệm chưa tách riêng số góc, chiều dài cạnh và khoảng hở nên chỉ "
        "được xem là diễn giải phù hợp, không phải bằng chứng nhân quả. Kết quả an toàn chỉ có nghĩa không phát hiện va chạm "
        "theo footprint và costmap đã cấu hình; bài báo không tuyên bố một biên clearance định lượng."))

    h2(doc, "E. Kịch bản 3—Kho có lối giao cắt")
    body(doc, (
        "Bản đồ kho có các lối giao cắt; robot đi từ (−2,00;−2,80) m đến (2,00;2,80) m. Kịch bản nhằm minh họa lựa chọn "
        "chuỗi chuyển tiếp qua nhiều góc trong cấu trúc kho, không phải phép đo độ ổn định G²—tính chất G² đã được suy ra "
        "từ cấu trúc điểm điều khiển ở Phần II. Hình 8 trình bày ca Theta* và Bảng V tổng hợp năm planner."), first=True)
    fig(doc, assets["warehouse"], 3.04,
        "Ca Theta* trong kho có lối giao cắt; PSTMO giữ thứ tự hành lang do planner quyết định.", 8)
    base.add_table_caption(doc, "V", "Kết quả trung bình—Kho có lối giao cắt")
    compact_table(doc, ["Phương pháp", "L (m)", "Eκ (m⁻¹)", "T (ms)"], metric_rows(metrics, "warehouse_cross_aisles"),
                  [1.25, 0.64, 0.78, 0.58], font_size=7.15)
    body(doc, (
        "PSTMO đạt L=7,541 m và Eκ=2,380 m⁻¹; Eκ thấp hơn Simple 82,89%, trong khi thời gian 79,8 ms cao hơn Constrained "
        "15,6 ms. Chiều dài nhỏ hơn không có nghĩa PSTMO tìm tuyến toàn cục tốt hơn: thứ tự hành lang vẫn do planner quyết "
        "định, còn PSTMO chỉ phân bố lại chuyển hướng trong cùng tuyến."))

    h2(doc, "F. Kết quả tổng hợp, đánh đổi và giới hạn")
    body(doc, (
        "Bảng VI tổng hợp trung bình mô tả của 15 nhóm ghép cặp. Do mọi smoother trong một nhóm nhận cùng Raw path, chênh "
        "lệch phản ánh bước xử lý đường trong đúng cấu hình đã nêu; tuy nhiên, một lần chạy mỗi tổ hợp không cho phép ước "
        "lượng độ bất định."), first=True)
    base.add_table_caption(doc, "VI", "Trung bình trên 15 nhóm bản đồ–bộ lập kế hoạch")
    compact_table(doc, ["Phương pháp", "L (m)", "Eκ (m⁻¹)", "T (ms)"], metric_rows(metrics, "all"),
                  [1.25, 0.64, 0.78, 0.58], font_size=7.2)
    body(doc, (
        "PSTMO đạt L=8,645 m, Eκ=2,887 m⁻¹ và T=96,0 ms. So với Raw, Simple, Savitzky–Golay và Constrained, Eκ giảm "
        "lần lượt 98,26%, 75,42%, 91,06% và 90,75%; chiều dài giảm lần lượt 1,97%, 0,98%, 1,62% và 2,01%. Như vậy, trong "
        "tập thử, giảm mức uốn không đi kèm việc tăng chiều dài. Mặt khác, PSTMO chậm hơn Simple 1,0 ms, Savitzky–Golay "
        "0,2 ms và Constrained 19,0 ms. Không thể kết luận PSTMO vượt trội toàn diện; đóng góp quan sát được nằm ở chất lượng "
        "hình học, với đánh đổi là thời gian xử lý."))
    body(doc, (
        "Eκ lớn của Raw xuất phát từ đổi hướng rời rạc sau khi mọi đường được đánh giá bằng cùng quy trình lấy mẫu. Simple "
        "giảm mạnh Eκ nhờ dịch chuyển lặp các điểm; Savitzky–Golay vẫn giữ nhiều biến thiên cục bộ; Constrained không tối "
        "thiểu trực tiếp Eκ khi wcurve=0. PSTMO buộc κ về không tại hai đầu chuyển tiếp, nên xu hướng kết quả phù hợp với mục "
        "tiêu thiết kế. Tuy nhiên, các phương pháp không được hiệu chỉnh để tối ưu cùng một hàm mục tiêu, vì thế kết quả không "
        "đại diện cho mọi cấu hình có thể có của từng plugin."))
    body(doc, (
        "Bằng chứng hiện chỉ bao gồm ba bản đồ tĩnh quy mô nhỏ, một cặp start–goal cho mỗi bản đồ, năm nguồn đường và một lần "
        "chạy mỗi tổ hợp. Lượng tử bộ đo làm các thời gian dưới vài mili giây kém phân giải; kiểm tra va chạm nhị phân không "
        "thay thế đánh giá clearance hoặc sai số mô hình. Eκ không trực tiếp phản ánh tracking error, năng lượng hay tải động "
        "lực học. PSTMO tìm trong tập tham số rời rạc nên không bảo đảm tối ưu liên tục toàn cục; G² hình học loại bước nhảy "
        "độ cong tại mối nối nhưng không tự bảo đảm jerk theo thời gian bằng không. Do đó, phạm vi thử nghiệm chưa đủ để đặc "
        "trưng giới hạn làm việc hoặc khả năng khái quát của phương pháp."))


def add_conclusion_and_references(doc: Document) -> None:
    h1(doc, "IV. KẾT LUẬN")
    body(doc, (
        "Bài báo đã xử lý bước chuyển từ đa tuyến có góc gãy của global planner sang đường phù hợp hơn cho robot vi sai trong "
        "Nav2. PSTMO sử dụng chuyển tiếp Bézier bậc năm G², cổng động học và swept footprint, rồi dùng quy hoạch động phối "
        "hợp các góc kề nhau. Trên 15 nhóm ghép cặp, Eκ giảm 75,42% so với Simple và 90,75% so với Constrained, trong khi "
        "chiều dài trung bình không tăng; đánh đổi là thời gian xử lý trung bình 96,0 ms. Bằng chứng hiện còn giới hạn ở mô "
        "phỏng tĩnh quy mô nhỏ và chưa đủ đặc trưng giới hạn làm việc. Nghiên cứu tiếp theo sẽ lặp lại trên nhiều tuyến và "
        "robot thật, đo sai số bám, clearance và năng lượng, khảo sát nhiễu, trượt, tải trọng và vật cản động, đồng thời xem "
        "xét tối ưu tham số liên tục và tham số hóa theo thời gian."), first=True)

    h1(doc, "TÀI LIỆU THAM KHẢO")
    refs = [
        "S. Macenski, F. Martín, R. White, and J. Ginés Clavero, “The Marathon 2: A Navigation System,” in Proc. IEEE/RSJ Int. Conf. Intelligent Robots and Systems (IROS), 2020, pp. 2718–2725, doi: 10.1109/IROS45743.2020.9341207.",
        "E. W. Dijkstra, “A Note on Two Problems in Connexion with Graphs,” Numerische Mathematik, vol. 1, pp. 269–271, 1959, doi: 10.1007/BF01386390.",
        "P. E. Hart, N. J. Nilsson, and B. Raphael, “A Formal Basis for the Heuristic Determination of Minimum Cost Paths,” IEEE Trans. Systems Science and Cybernetics, vol. 4, no. 2, pp. 100–107, 1968, doi: 10.1109/TSSC.1968.300136.",
        "A. Nash, K. Daniel, S. Koenig, and A. Felner, “Theta*: Any-Angle Path Planning on Grids,” in Proc. AAAI Conf. Artificial Intelligence, 2007, pp. 1177–1183.",
        "D. Dolgov, S. Thrun, M. Montemerlo, and J. Diebel, “Practical Search Techniques in Path Planning for Autonomous Driving,” in Proc. 1st Int. Symp. Search Techniques in Artificial Intelligence and Robotics, 2008.",
        "A. Ravankar, A. A. Ravankar, Y. Kobayashi, Y. Hoshino, and C.-C. Peng, “Path Smoothing Techniques in Robot Navigation: State-of-the-Art, Current and Future Challenges,” Sensors, vol. 18, no. 9, Art. no. 3170, 2018, doi: 10.3390/s18093170.",
        "S. Fleury, P. Souères, J.-P. Laumond, and R. Chatila, “Primitives for Smoothing Mobile Robot Trajectories,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 1993, pp. 832–839.",
        "M. Brezak and I. Petrović, “Path Smoothing Using Clothoids for Differential Drive Mobile Robots,” IFAC Proc. Volumes, vol. 44, no. 1, pp. 1133–1138, 2011, doi: 10.3182/20110828-6-IT-1002.02944.",
        "K. Yang and S. Sukkarieh, “An Analytical Continuous-Curvature Path-Smoothing Algorithm,” IEEE Trans. Robotics, vol. 26, no. 3, pp. 561–568, 2010, doi: 10.1109/TRO.2010.2042990.",
        "X. Bu, H. Su, W. Zou, and P. Wang, “Curvature Continuous Path Smoothing Based on Cubic Bezier Curves for Car-Like Vehicles,” in Proc. IEEE Int. Conf. Robotics and Biomimetics (ROBIO), 2015, pp. 1453–1458, doi: 10.1109/ROBIO.2015.7418975.",
        "L. Xu, D. Wang, B. Song, and M. Cao, “Global Smooth Path Planning for Mobile Robots Based on Continuous Bezier Curve,” in Proc. Chinese Automation Congress (CAC), 2017, pp. 2081–2085, doi: 10.1109/CAC.2017.8243114.",
        "V. Parque and T. Miyashita, “Smooth Curve Fitting of Mobile Robot Trajectories Using Differential Evolution,” IEEE Access, vol. 8, pp. 82855–82866, 2020, doi: 10.1109/ACCESS.2020.2991003.",
        "S. Quinlan and O. Khatib, “Elastic Bands: Connecting Path Planning and Control,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 1993, pp. 802–807.",
        "A. Savitzky and M. J. E. Golay, “Smoothing and Differentiation of Data by Simplified Least Squares Procedures,” Analytical Chemistry, vol. 36, no. 8, pp. 1627–1639, 1964, doi: 10.1021/ac60214a047.",
        "Navigation2 Project, “Simple Smoother,” Nav2 Configuration Guide. [Online]. Available: https://docs.nav2.org/configuration/packages/configuring-simple-smoother.html. Accessed: Aug. 26, 2026.",
        "Navigation2 Project, “Savitzky-Golay Smoother,” Nav2 Configuration Guide. [Online]. Available: https://docs.nav2.org/configuration/packages/configuring-savitzky-golay-smoother.html. Accessed: Aug. 26, 2026.",
        "Navigation2 Project, “Constrained Smoother,” Nav2 Configuration Guide. [Online]. Available: https://docs.nav2.org/configuration/packages/configuring-constrained-smoother.html. Accessed: Aug. 26, 2026.",
        "N. Ratliff, M. Zucker, J. A. Bagnell, and S. Srinivasa, “CHOMP: Gradient Optimization Techniques for Efficient Motion Planning,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 2009, pp. 489–494.",
        "E. Heiden, L. Palmieri, S. Koenig, K. O. Arras, and G. S. Sukhatme, “Gradient-Informed Path Smoothing for Wheeled Mobile Robots,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 2018, pp. 1710–1717, doi: 10.1109/ICRA.2018.8460818.",
        "H. Andreasson, J. Saarinen, M. Cirillo, T. Stoyanov, and A. J. Lilienthal, “Fast, Continuous State Path Smoothing to Improve Navigation Accuracy,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 2015, pp. 662–669, doi: 10.1109/ICRA.2015.7139250.",
        "C. Rösmann, W. Feiten, T. Wösch, F. Hoffmann, and T. Bertram, “Trajectory Modification Considering Dynamic Constraints of Autonomous Robots,” in ROBOTIK 2012, pp. 74–79, 2012.",
        "C. Rösmann, F. Hoffmann, and T. Bertram, “Timed-Elastic-Bands for Time-Optimal Point-to-Point Nonlinear Model Predictive Control,” in Proc. European Control Conf. (ECC), 2015, pp. 3352–3357, doi: 10.1109/ECC.2015.7331052.",
        "C. Rösmann, F. Hoffmann, and T. Bertram, “Kinodynamic Trajectory Optimization and Control for Car-Like Robots,” in Proc. IEEE/RSJ Int. Conf. Intelligent Robots and Systems (IROS), 2017, pp. 5681–5686, doi: 10.1109/IROS.2017.8206458.",
    ]
    for index, reference in enumerate(refs, 1):
        paragraph = base.add_reference(doc, index, reference)
        for run in paragraph.runs:
            base.set_run_font(run, 7.35)


def build_document(template_path: Path, assets: dict[str, Path], metrics) -> None:
    doc = Document(str(template_path))
    base.ensure_custom_styles(doc)
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in footer.paragraphs:
                base.clear_paragraph_content(paragraph)

    configure_title_and_abstract(doc)
    add_introduction(doc, assets)
    add_method(doc, assets)
    add_experiments(doc, assets, metrics)
    add_conclusion_and_references(doc)

    doc.core_properties.title = (
        "Làm mượt đường đi bằng tối ưu chuyển tiếp tại góc có xét hình bao cho robot di động vi sai"
    )
    doc.core_properties.subject = "Bản thảo tiếng Việt ver2 theo mẫu IEEE A4 cho ICEEIS 2026"
    doc.core_properties.author = (
        "Hai Linh Pham; Tien Cuong Nguyen; Viet Bao Nguyen; "
        "Thi Phuong Thao Nguyen; Thi Ly Pham"
    )
    doc.core_properties.keywords = (
        "PSTMO; path smoothing; differential drive; quintic Bezier; swept footprint; Nav2"
    )
    doc.core_properties.comments = (
        "Ver2: related work integrated into Introduction; evidence from the paired 15-group benchmark."
    )
    doc.save(str(OUTPUT_DOCX))


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    _, metrics = base.load_results()
    assets = prepare_assets(metrics)
    with tempfile.TemporaryDirectory(prefix="iceeis_ver2_") as temp_name:
        temp = Path(temp_name)
        converted = temp / ABSTRACT_TEMPLATE.name
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx", "--outdir", str(temp), str(ABSTRACT_TEMPLATE)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
        if not converted.exists():
            raise RuntimeError("LibreOffice did not create the Transitional OOXML template copy")
        build_document(converted, assets, metrics)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
