#!/usr/bin/env python3

"""Audit the final PSTMO data and rebuild the Vietnamese ICEEIS abstract."""

from __future__ import annotations

import ast
import collections
import csv
import hashlib
import html
import io
import json
import math
from pathlib import Path
import statistics
import zipfile

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
BASELINE_DIR = (
    ROOT / "results" / "current_pstmo_footprint_padding15_nav2_comparison_20260802"
)
GEOMETRY_DIR = (
    ROOT / "results" / "pstmo_direct_dq_local08_adaptive_los_full_20260802"
)
CLOSED_LOOP_DIR = (
    ROOT
    / "results"
    / "pstmo_direct_dq_local08_adaptive_los_closed_loop_20260802"
)
REQUESTED_DOCX = (
    ROOT / "abstract" / "ICEEIS_2026_ADAPTIVE_PIVOT_G2_ABSTRACT_VI.docx"
)
OUTPUT_DOCX = ROOT / "abstract" / "ICEEIS_2026_PSTMO_FOOTPRINT_LOS_ABSTRACT_VI.docx"
REQUESTED_HTML = (
    ROOT / "abstract" / "ICEEIS_2026_ADAPTIVE_PIVOT_G2_ABSTRACT_VI.html"
)
OUTPUT_HTML = ROOT / "abstract" / "ICEEIS_2026_PSTMO_FOOTPRINT_LOS_ABSTRACT_VI.html"
NAV2_PARAMS = ROOT / "src" / "vacuum_robot_gazebo" / "config" / "nav2_params.yaml"

METHODS = ("raw", "simple", "savitzky_golay", "constrained", "pstmo")
METHOD_LABELS = {
    "raw": "Raw",
    "simple": "Simple",
    "savitzky_golay": "Savitzky–Golay",
    "constrained": "Constrained",
    "pstmo": "PSTMO",
}
EXPECTED_ENVIRONMENTS = {
    "research_warehouse": "lower_left_diagonal",
    "narrow_aisles": "southwest_northeast_weave",
    "office_maze": "office_long_diagonal",
    "open_arena": "center_block_detour",
    "warehouse_cross_aisles": "cross_aisle_transfer",
    "warehouse_dispatch": "full_replenishment",
    "warehouse_long_aisles": "diagonal_replenishment",
}


def as_float(value: object) -> float:
    result = float(value)
    if not math.isfinite(result):
        raise ValueError(f"Non-finite numeric value: {value!r}")
    return result


def mean(rows: list[dict], key: str) -> float:
    return statistics.fmean(as_float(row[key]) for row in rows)


def percent_change(value: float, reference: float) -> float:
    return 100.0 * (value - reference) / reference


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def literal_dict(value: str) -> dict:
    parsed = ast.literal_eval(value)
    if not isinstance(parsed, dict):
        raise ValueError("Expected a diagnostic dictionary")
    return parsed


def load_geometry() -> tuple[list[dict[str, str]], dict]:
    """Combine unchanged Nav2 baselines with the final PSTMO rows."""
    rows: list[dict[str, str]] = []
    for environment, scenario in EXPECTED_ENVIRONMENTS.items():
        baseline_path = BASELINE_DIR / f"{environment}.csv"
        final_path = GEOMETRY_DIR / f"{environment}.csv"
        with baseline_path.open(newline="", encoding="utf-8") as stream:
            baseline = list(csv.DictReader(stream))
        with final_path.open(newline="", encoding="utf-8") as stream:
            final = list(csv.DictReader(stream))
        selected = [row for row in baseline if row["method"] != "pstmo"]
        selected.extend(row for row in final if row["method"] == "pstmo")
        if len(selected) != 25 or {row["scenario"] for row in selected} != {scenario}:
            raise RuntimeError(f"Unexpected comparison design in {environment}")
        for row in selected:
            row["environment"] = environment
        rows.extend(selected)

    if len(rows) != 175 or {row["method"] for row in rows} != set(METHODS):
        raise RuntimeError("The final comparison must contain 175 rows and five methods")
    if any(row["method"] == "adaptive_hybrid" for row in rows):
        raise RuntimeError("The excluded method appeared in the abstract data")

    groups: dict[tuple[str, str, str, str], list[dict[str, str]]] = (
        collections.defaultdict(list)
    )
    for row in rows:
        key = (
            row["environment"],
            row["scenario"],
            row["planner"],
            row["repetition"],
        )
        groups[key].append(row)
    if len(groups) != 35:
        raise RuntimeError(f"Expected 35 paired groups, found {len(groups)}")
    for key, group in groups.items():
        if len(group) != 5 or {row["method"] for row in group} != set(METHODS):
            raise RuntimeError(f"Incomplete method set in {key}")
        if len({row["raw_path_sha256"] for row in group}) != 1:
            raise RuntimeError(f"Raw-path hash mismatch in {key}")

    pstmo_rows = [row for row in rows if row["method"] == "pstmo"]
    if len(pstmo_rows) != 35 or any(row["success"] != "True" for row in pstmo_rows):
        raise RuntimeError("Final PSTMO must succeed in all 35 paired cases")
    required_contract = {
        "pstmo_los_selection_enabled": "True",
        "pstmo_los_completed": "True",
        "pstmo_los_no_los_completed": "True",
        "pstmo_search_mode": "joint_d_q",
        "pstmo_trim_domain": "direct_metric",
        "pstmo_fallback": "none",
    }
    for field, expected in required_contract.items():
        if any(row.get(field) != expected for row in pstmo_rows):
            raise RuntimeError(f"Final PSTMO contract mismatch for {field}")
    if any(
        not math.isclose(
            as_float(row["pstmo_los_footprint_padding_m"]),
            0.15,
            rel_tol=0.0,
            abs_tol=1.0e-9,
        )
        for row in pstmo_rows
    ):
        raise RuntimeError("Footprint-aware LOS did not use 0.15 m padding")
    collision_samples = sum(
        int(as_float(row["footprint_collision_sample_count"])) for row in pstmo_rows
    )
    if collision_samples != 0:
        raise RuntimeError("A final PSTMO path contains a footprint collision")

    los_selected = [row for row in pstmo_rows if row["pstmo_los_selected"] == "True"]
    no_los_selected = [
        row for row in pstmo_rows if row["pstmo_los_selected"] == "False"
    ]
    if len(los_selected) != 25 or len(no_los_selected) != 10:
        raise RuntimeError("Unexpected adaptive LOS decision counts")

    quality_triplets = []
    for row in pstmo_rows:
        no_los = literal_dict(row["pstmo_los_no_los_quality"])
        los = literal_dict(row["pstmo_los_quality"])
        selected = los if row["pstmo_los_selected"] == "True" else no_los
        if not all(branch.get("valid") and branch.get("safe") for branch in (no_los, los)):
            raise RuntimeError("An internal LOS comparison branch was invalid or unsafe")
        quality_triplets.append((no_los, los, selected))

    quality_keys = (
        "length_m",
        "max_curvature_1pm",
        "curvature_energy_1pm",
        "pivot_rotation_rad",
        "peak_proximity_cost",
    )
    branch_quality = {}
    for name, index in (("no_los", 0), ("los", 1), ("adaptive", 2)):
        branch_quality[name] = {
            key: statistics.fmean(as_float(item[index][key]) for item in quality_triplets)
            for key in quality_keys
        }
    adaptive_vs_no_los = {}
    for key in quality_keys:
        before = branch_quality["no_los"][key]
        after = branch_quality["adaptive"][key]
        lower_is_better = key != "peak_proximity_cost"
        adaptive_vs_no_los[key] = {
            "no_los_mean": before,
            "adaptive_mean": after,
            "relative_change_percent": percent_change(after, before),
            "improved_pair_count": sum(
                (item[2][key] < item[0][key] - 1.0e-12)
                if lower_is_better
                else (item[2][key] < item[0][key] - 1.0e-12)
                for item in quality_triplets
            ),
        }

    margins = [
        as_float(row["pstmo_los_no_los_score"]) - as_float(row["pstmo_los_score"])
        for row in los_selected
    ]
    q_values: list[float] = []
    d_values: list[float] = []
    for row in pstmo_rows:
        corners = ast.literal_eval(row["pstmo_corner_search"])
        for corner in corners:
            if as_float(corner.get("selected_trim", 0.0)) > 0.0:
                d_values.append(as_float(corner["selected_trim"]))
                q_values.append(as_float(corner["selected_control_fraction"]))

    common_groups = [
        group for group in groups.values() if all(row["success"] == "True" for row in group)
    ]
    if len(common_groups) != 34:
        raise RuntimeError(f"Expected 34 common-success groups, found {len(common_groups)}")
    common_rows = [row for group in common_groups for row in group]
    metric_keys = (
        "translation_path_length_m",
        "translation_max_abs_curvature_1pm",
        "translation_curvature_energy_1pm",
        "footprint_clearance_min_m",
        "algorithm_time_s",
        "wall_time_s",
    )
    common_means = {
        method: {
            key: mean(
                [row for row in common_rows if row["method"] == method], key
            )
            for key in metric_keys
        }
        for method in METHODS
    }
    success_counts = {
        method: sum(
            row["method"] == method and row["success"] == "True" for row in rows
        )
        for method in METHODS
    }
    if success_counts != {
        "raw": 35,
        "simple": 34,
        "savitzky_golay": 35,
        "constrained": 35,
        "pstmo": 35,
    }:
        raise RuntimeError(f"Unexpected success counts: {success_counts}")

    summary = {
        "design": {
            "environment_count": 7,
            "scenario_count": 7,
            "planner_count": 5,
            "methods": list(METHODS),
            "record_count": 175,
            "paired_group_count": 35,
            "common_success_group_count": 34,
        },
        "configuration": {
            "line_of_sight_mode": "adaptive_full_path_selection",
            "line_of_sight_footprint_padding_m": 0.15,
            "minimum_trim_distance_m": 0.02,
            "maximum_trim_distance_m": 0.8,
            "control_fraction_reference": 0.35,
            "los_minimum_score_improvement": 0.005,
            "nav2_params_sha256": sha256(NAV2_PARAMS),
        },
        "validation": {
            "raw_path_hash_consistent_group_count": 35,
            "pstmo_success_count": 35,
            "pstmo_fallback_count": 0,
            "pstmo_footprint_collision_sample_count": 0,
            "both_internal_branches_safe_count": 35,
        },
        "success_count": success_counts,
        "los_selection": {
            "los_selected_count": len(los_selected),
            "no_los_selected_count": len(no_los_selected),
            "minimum_selected_score_improvement": min(margins),
            "mean_selected_score_improvement": statistics.fmean(margins),
            "maximum_selected_score_improvement": max(margins),
            "branch_quality_means": branch_quality,
            "adaptive_vs_no_los": adaptive_vs_no_los,
        },
        "joint_d_q": {
            "selected_transition_count": len(q_values),
            "angle_aware_q_count": sum(abs(value - 0.35) > 1.0e-8 for value in q_values),
            "control_fraction_mean": statistics.fmean(q_values),
            "control_fraction_min": min(q_values),
            "control_fraction_max": max(q_values),
            "trim_distance_mean_m": statistics.fmean(d_values),
            "trim_distance_min_m": min(d_values),
            "trim_distance_max_m": max(d_values),
        },
        "common_success_means": common_means,
    }
    return rows, summary


def load_closed_loop() -> dict:
    records: list[dict] = []
    for environment in ("research_warehouse", "narrow_aisles", "warehouse_dispatch"):
        summaries = list((CLOSED_LOOP_DIR / environment).glob("*_summary.json"))
        if len(summaries) != 1:
            raise RuntimeError(f"Expected one closed-loop summary for {environment}")
        selected = json.loads(summaries[0].read_text(encoding="utf-8"))["records"]
        if len(selected) != 2 or {row["method"] for row in selected} != {"raw", "pstmo"}:
            raise RuntimeError(f"Unexpected closed-loop design in {summaries[0]}")
        if len({row["raw_path_sha256"] for row in selected}) != 1:
            raise RuntimeError(f"Closed-loop raw-path mismatch in {summaries[0]}")
        for row in selected:
            row["environment"] = environment
        records.extend(selected)
    if len(records) != 6 or any(row.get("success") is not True for row in records):
        raise RuntimeError("Every final closed-loop trial must succeed")
    if any(
        int(row.get("planned_footprint_collision_sample_count", 0)) != 0
        or int(row.get("collision_monitor_interventions", 0)) != 0
        for row in records
    ):
        raise RuntimeError("A closed-loop safety event occurred")

    metric_keys = (
        "execution_time_s",
        "tracking_rmse_m",
        "tracking_max_error_m",
        "executed_curvature_energy_1pm",
        "planned_translation_path_length_m",
        "final_position_error_m",
        "planned_footprint_clearance_min_m",
    )
    means = {}
    for method in ("raw", "pstmo"):
        selected = [row for row in records if row["method"] == method]
        means[method] = {key: mean(selected, key) for key in metric_keys}
    changes = {
        key: percent_change(means["pstmo"][key], means["raw"][key])
        for key in metric_keys
    }
    return {
        "trial_count": 6,
        "paired_group_count": 3,
        "raw_path_hash_match_count": 3,
        "raw_success_count": 3,
        "pstmo_success_count": 3,
        "planned_footprint_collision_sample_count": 0,
        "collision_monitor_interventions": 0,
        "method_means": means,
        "pstmo_vs_raw_change_percent": changes,
    }


def vi_number(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def abstract_text(geometry: dict, closed: dict) -> str:
    los = geometry["los_selection"]
    comparison = los["adaptive_vs_no_los"]
    dq = geometry["joint_d_q"]
    means = geometry["common_success_means"]
    pstmo = means["pstmo"]
    closed_means = closed["method_means"]
    raw_closed = closed_means["raw"]
    pstmo_closed = closed_means["pstmo"]
    closed_change = closed["pstmo_vs_raw_change_percent"]
    other = METHODS[:-1]
    best_other_energy = min(
        means[method]["translation_curvature_energy_1pm"] for method in other
    )
    best_other_curvature = min(
        means[method]["translation_max_abs_curvature_1pm"] for method in other
    )
    best_other_length = min(
        means[method]["translation_path_length_m"] for method in other
    )

    length = comparison["length_m"]
    curvature = comparison["max_curvature_1pm"]
    energy = comparison["curvature_energy_1pm"]
    pivot = comparison["pivot_rotation_rad"]
    proximity = comparison["peak_proximity_cost"]
    return (
        "Tóm tắt—Các bộ lập kế hoạch toàn cục trong ROS 2/Nav2 thường tạo đường "
        "gấp khúc có đổi hướng đột ngột. Nghiên cứu này đề xuất phương pháp làm "
        "mượt đường đi và tối ưu hóa thao tác chuyển hướng PSTMO cho robot di "
        "động vi sai hai bánh. Mỗi góc được biểu diễn bằng quay tại chỗ hoặc "
        "chuyển tiếp Bézier bậc năm liên tục hình học G². Khác với cách sinh tham "
        "số trước đây, khoảng cắt d được tìm trực tiếp theo mét trong miền "
        "[0,02; min(0,8, Lᵢₙ, Lₒᵤₜ)] thay vì suy ra từ bán kính của một cung tròn "
        "không đồng nhất với đường Bézier. Với từng d, khoảng điều khiển q so sánh "
        "tỷ lệ q/d theo góc với đúng mốc tham chiếu 0,35; các ứng viên được kiểm "
        "tra giới hạn bánh xe, tham số hóa thời gian, vùng quét footprint và tối "
        "ưu hóa toàn cục giữa các góc. Trong "
        f"{dq['selected_transition_count']} chuyển tiếp được chọn, "
        f"{dq['angle_aware_q_count']} chuyển tiếp ({vi_number(100.0 * dq['angle_aware_q_count'] / dq['selected_transition_count'], 1)}%) "
        "dùng q/d khác 0,35. Tiền xử lý line-of-sight (LOS) dùng footprint đang "
        "được Nav2 công bố, cộng biên vận hành 0,15 m và quét cả dịch chuyển lẫn "
        "quay. Để LOS không làm giảm kết quả, PSTMO tạo hai đầu ra hoàn chỉnh từ "
        "cùng đường Raw: một nhánh không LOS và một nhánh LOS xét footprint; LOS "
        "chỉ được giữ khi điểm chất lượng toàn đường giảm ít nhất 0,005, nếu không "
        "thuật toán giữ nhánh không LOS. Đánh giá hình học gồm 175 bản ghi trên 7 "
        "tình huống, 5 global planner và 5 phương án Raw, Simple, Savitzky–Golay, "
        "Constrained và PSTMO. Cả 35 nhóm ghép cặp có cùng hash đường Raw. Bộ chọn "
        f"dùng LOS ở {los['los_selected_count']}/35 ca và giữ không LOS ở "
        f"{los['no_los_selected_count']}/35 ca; PSTMO thành công 35/35, không "
        "fallback và không có mẫu va chạm footprint. So với luôn dùng nhánh không "
        "LOS của chính cấu hình mới, lựa chọn thích nghi giảm chiều dài trung bình "
        f"từ {vi_number(length['no_los_mean'], 3)} xuống "
        f"{vi_number(length['adaptive_mean'], 3)} m "
        f"({vi_number(abs(length['relative_change_percent']))}%), giảm độ cong cực "
        f"đại từ {vi_number(curvature['no_los_mean'], 3)} xuống "
        f"{vi_number(curvature['adaptive_mean'], 3)} m⁻¹ "
        f"({vi_number(abs(curvature['relative_change_percent']))}%), giảm năng "
        f"lượng độ cong từ {vi_number(energy['no_los_mean'], 3)} xuống "
        f"{vi_number(energy['adaptive_mean'], 3)} m⁻¹ "
        f"({vi_number(abs(energy['relative_change_percent']))}%) và giảm tổng góc "
        f"quay tại chỗ {vi_number(abs(pivot['relative_change_percent']))}%. Đánh "
        "đổi là chi phí lân cận vật cản cực đại trung bình tăng từ "
        f"{vi_number(proximity['no_los_mean'], 1)} lên "
        f"{vi_number(proximity['adaptive_mean'], 1)} "
        f"({vi_number(proximity['relative_change_percent'])}%). Trên 34 nhóm mà "
        "mọi phương án đều thành công, PSTMO đạt năng lượng độ cong tịnh tiến "
        f"{vi_number(pstmo['translation_curvature_energy_1pm'], 3)} m⁻¹, độ cong "
        f"cực đại {vi_number(pstmo['translation_max_abs_curvature_1pm'], 3)} m⁻¹ "
        f"và chiều dài {vi_number(pstmo['translation_path_length_m'], 3)} m, tốt "
        f"hơn phương án đối chứng tốt nhất tương ứng {vi_number(best_other_energy, 3)} m⁻¹, "
        f"{vi_number(best_other_curvature, 3)} m⁻¹ và {vi_number(best_other_length, 3)} m; "
        f"thời gian thuật toán là {vi_number(1000.0 * pstmo['algorithm_time_s'], 1)} ms. "
        "Trong ba cặp thử nghiệm vòng kín, Raw và PSTMO đều hoàn thành 3/3 mà "
        "không kích hoạt bộ giám sát va chạm. PSTMO giảm thời gian hoàn thành từ "
        f"{vi_number(raw_closed['execution_time_s'], 3)} xuống "
        f"{vi_number(pstmo_closed['execution_time_s'], 3)} s "
        f"({vi_number(abs(closed_change['execution_time_s']))}%), giảm năng lượng "
        f"độ cong thực thi từ {vi_number(raw_closed['executed_curvature_energy_1pm'], 3)} "
        f"xuống {vi_number(pstmo_closed['executed_curvature_energy_1pm'], 3)} m⁻¹ "
        f"({vi_number(abs(closed_change['executed_curvature_energy_1pm']))}%) và "
        "giảm nhẹ sai số bám cực đại, nhưng RMSE bám tăng từ "
        f"{vi_number(100.0 * raw_closed['tracking_rmse_m'], 3)} lên "
        f"{vi_number(100.0 * pstmo_closed['tracking_rmse_m'], 3)} cm. Kết quả cho "
        "thấy LOS xét footprint có lợi rõ rệt khi được chọn theo chất lượng, nhưng "
        "không nên ép dùng cho mọi đường; cần tăng số lần lặp và thử nghiệm phần "
        "cứng để xác nhận khả năng khái quát."
    )


def replace_paragraph(paragraph, label: str, body: str) -> None:
    paragraph.clear()
    first = paragraph.add_run(label)
    first.bold = True
    second = paragraph.add_run(body)
    for run in (first, second):
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(11)


def deduplicate_docx_archive(path: Path) -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(path, "r") as source:
        latest = {info.filename: info for info in source.infolist()}
        with zipfile.ZipFile(buffer, "w") as destination:
            for name, info in latest.items():
                destination.writestr(info, source.read(info))
    path.write_bytes(buffer.getvalue())


def write_docx(text: str) -> None:
    document = Document(REQUESTED_DOCX)
    abstracts = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Tóm tắt—")
    ]
    if len(abstracts) != 1:
        raise RuntimeError("Could not uniquely locate the Vietnamese abstract")
    replace_paragraph(abstracts[0], "Tóm tắt—", text.removeprefix("Tóm tắt—"))
    keywords = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Từ khóa—")
    ]
    if len(keywords) != 1:
        raise RuntimeError("Could not uniquely locate the keyword paragraph")
    replace_paragraph(
        keywords[0],
        "Từ khóa—",
        "PSTMO; robot di động vi sai; làm mượt đường đi; line-of-sight thích nghi; "
        "an toàn footprint; ROS 2/Nav2.",
    )
    document.core_properties.title = "ICEEIS 2026 — PSTMO với LOS xét footprint"
    document.core_properties.subject = "Bản tiếng Việt dùng để đối chiếu"
    document.core_properties.keywords = (
        "PSTMO, adaptive footprint-aware line-of-sight, path smoothing, "
        "differential drive, ROS 2, Nav2"
    )
    document.save(OUTPUT_DOCX)
    deduplicate_docx_archive(OUTPUT_DOCX)
    REQUESTED_DOCX.write_bytes(OUTPUT_DOCX.read_bytes())


def write_html(text: str) -> None:
    title = (
        "Phương pháp làm mượt đường đi và tối ưu hóa thao tác chuyển hướng "
        "có xét an toàn cho robot di động vi sai hai bánh"
    )
    body = html.escape(text.removeprefix("Tóm tắt—"))
    output = f"""<!doctype html>
<html lang="vi"><head><meta charset="utf-8"><title>ICEEIS 2026 — PSTMO với LOS xét footprint</title>
<style>body{{font-family:'Times New Roman',serif;max-width:850px;margin:40px auto;line-height:1.45;font-size:12pt}}h1{{text-align:center;font-size:18pt}}.notice{{text-align:center;font-weight:bold}}p{{text-align:justify}}.label{{font-weight:bold}}</style></head>
<body><p class="notice">BẢN TIẾNG VIỆT CHỈ DÙNG ĐỂ ĐỐI CHIẾU — KHÔNG NỘP LÊN CMT</p>
<h1>{html.escape(title)}</h1>
<p><span class="label">Tóm tắt—</span>{body}</p>
<p><span class="label">Từ khóa—</span>PSTMO; robot di động vi sai; làm mượt đường đi; line-of-sight thích nghi; an toàn footprint; ROS 2/Nav2.</p>
</body></html>\n"""
    OUTPUT_HTML.write_text(output, encoding="utf-8")
    REQUESTED_HTML.write_text(output, encoding="utf-8")


def write_results_readme(geometry: dict, closed: dict) -> None:
    means = geometry["common_success_means"]
    selection = geometry["los_selection"]
    comparison = selection["adaptive_vs_no_los"]
    dq = geometry["joint_d_q"]
    lines = [
        "# Kết quả cuối PSTMO với LOS thích nghi xét footprint",
        "",
        "LOS là tiền xử lý nội tại của PSTMO. Thuật toán đánh giá cả nhánh "
        "không LOS và nhánh LOS trên cùng đường Raw, sau đó chỉ chọn LOS khi "
        "điểm chất lượng toàn đường tốt hơn ít nhất 0,005.",
        "",
        "## Thiết kế và kiểm định",
        "",
        "- 7 môi trường, 7 tình huống, 5 global planner và 5 phương án;",
        "- 35 nhóm ghép cặp, tổng cộng 175 bản ghi;",
        "- 35/35 nhóm có cùng `raw_path_sha256` giữa các phương án;",
        f"- LOS được chọn {selection['los_selected_count']}/35 ca; không LOS được "
        f"chọn {selection['no_los_selected_count']}/35 ca;",
        "- PSTMO thành công 35/35, không fallback và không có mẫu va chạm footprint;",
        "- cả hai nhánh nội bộ hợp lệ và an toàn trong 35/35 ca;",
        f"- d được tìm trực tiếp trong 0,02–0,8 m; {dq['angle_aware_q_count']}/"
        f"{dq['selected_transition_count']} chuyển tiếp chọn q/d khác 0,35.",
        "",
        "## LOS thích nghi so với luôn dùng nhánh không LOS",
        "",
        "| Chỉ số | Không LOS | Lựa chọn thích nghi | Thay đổi |",
        "|---|---:|---:|---:|",
    ]
    for key, label, unit in (
        ("length_m", "Chiều dài", "m"),
        ("max_curvature_1pm", "Độ cong cực đại", "1/m"),
        ("curvature_energy_1pm", "Năng lượng độ cong", "1/m"),
        ("pivot_rotation_rad", "Tổng góc quay tại chỗ", "rad"),
        ("peak_proximity_cost", "Chi phí lân cận cực đại", "cost"),
    ):
        row = comparison[key]
        lines.append(
            f"| {label} | {row['no_los_mean']:.4f} {unit} | "
            f"{row['adaptive_mean']:.4f} {unit} | "
            f"{row['relative_change_percent']:+.2f}% |"
        )
    lines.extend(
        [
            "",
            "Bốn chỉ số chuyển động đầu giảm; chi phí lân cận vật cản tăng là "
            "đánh đổi. LOS không bị ép dùng trong 10 ca mà điểm toàn đường không "
            "cải thiện đủ ngưỡng.",
            "",
            "## So sánh hình học trên 34 nhóm cùng thành công",
            "",
            "| Phương án | Thành công | Chiều dài (m) | Kmax (1/m) | Eκ (1/m) | Clearance (m) | Thời gian thuật toán (ms) |",
            "|---|---:|---:|---:|---:|---:|---:|",
        ]
    )
    for method in METHODS:
        row = means[method]
        lines.append(
            f"| {METHOD_LABELS[method]} | {geometry['success_count'][method]}/35 | "
            f"{row['translation_path_length_m']:.3f} | "
            f"{row['translation_max_abs_curvature_1pm']:.3f} | "
            f"{row['translation_curvature_energy_1pm']:.3f} | "
            f"{row['footprint_clearance_min_m']:.3f} | "
            f"{1000.0 * row['algorithm_time_s']:.1f} |"
        )
    raw = closed["method_means"]["raw"]
    pstmo = closed["method_means"]["pstmo"]
    change = closed["pstmo_vs_raw_change_percent"]
    lines.extend(
        [
            "",
            "## Kiểm chứng vòng kín trên ba cặp",
            "",
            "Raw và PSTMO đều hoàn thành 3/3; không có mẫu va chạm footprint "
            "trên đường kế hoạch và không có can thiệp của bộ giám sát va chạm.",
            "",
            "| Chỉ số | Raw | PSTMO | Thay đổi |",
            "|---|---:|---:|---:|",
        ]
    )
    for key, label, scale, unit in (
        ("execution_time_s", "Thời gian hoàn thành", 1.0, "s"),
        ("executed_curvature_energy_1pm", "Eκ thực thi", 1.0, "1/m"),
        ("planned_translation_path_length_m", "Chiều dài kế hoạch", 1.0, "m"),
        ("tracking_max_error_m", "Sai số bám cực đại", 100.0, "cm"),
        ("tracking_rmse_m", "RMSE bám", 100.0, "cm"),
        ("final_position_error_m", "Sai số đích", 100.0, "cm"),
        ("planned_footprint_clearance_min_m", "Clearance kế hoạch", 1.0, "m"),
    ):
        lines.append(
            f"| {label} | {scale * raw[key]:.3f} {unit} | "
            f"{scale * pstmo[key]:.3f} {unit} | {change[key]:+.2f}% |"
        )
    lines.extend(
        [
            "",
            "Kết luận: với tiêu chí hiện tại, LOS thích nghi tốt hơn không LOS "
            "về chiều dài, độ cong cực đại, năng lượng độ cong và lượng quay tại "
            "chỗ; đánh đổi là đi gần vùng chi phí cao hơn. Kết quả vòng kín có "
            "lợi về thời gian và năng lượng, nhưng RMSE bám và sai số đích tăng.",
            "",
        ]
    )
    (GEOMETRY_DIR / "README.md").write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    _, geometry = load_geometry()
    closed = load_closed_loop()
    geometry["closed_loop"] = closed
    (GEOMETRY_DIR / "aggregate_summary.json").write_text(
        json.dumps(geometry, ensure_ascii=False, indent=2, sort_keys=True) + "\n",
        encoding="utf-8",
    )
    text = abstract_text(geometry, closed)
    write_docx(text)
    write_html(text)
    write_results_readme(geometry, closed)
    print(f"Wrote {GEOMETRY_DIR / 'aggregate_summary.json'}")
    print(f"Wrote {GEOMETRY_DIR / 'README.md'}")
    print(f"Wrote {OUTPUT_DOCX}")
    print(f"Updated {REQUESTED_DOCX}")


if __name__ == "__main__":
    main()
