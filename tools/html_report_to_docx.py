#!/usr/bin/env python3

"""Convert the generated REV-ECIT HTML reports to styled A4 DOCX files.

Dependencies: ``python3-docx`` and ``python3-bs4``.  Ubuntu packages may be
installed with ``sudo apt install python3-docx python3-bs4``.  The converter is
kept deliberately small and supports only the semantic elements emitted by
``generate_rev_ecit_2026_report.py``.
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


ROOT = Path(__file__).resolve().parents[1]


def set_cell_text_size(cell, size):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_paragraph_shading(paragraph, fill):
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_repeat_table_header(row):
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_cell_margins(cell, top=45, start=45, bottom=45, end=45):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_columns(section, count, space_twips=360):
    properties = section._sectPr
    columns = properties.xpath("./w:cols")
    element = columns[0] if columns else OxmlElement("w:cols")
    if not columns:
        properties.append(element)
    element.set(qn("w:num"), str(count))
    element.set(qn("w:space"), str(space_twips))


def set_section_geometry(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(19)
    section.bottom_margin = Mm(19)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(8)


def add_page_number(section):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    finish = OxmlElement("w:fldChar")
    finish.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, finish):
        run._r.append(element)


def prevent_row_split(row):
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def configure_styles(document, paper):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(10 if paper else 10)
    normal.paragraph_format.space_after = Pt(1.5 if paper else 3)
    normal.paragraph_format.line_spacing = 1.0 if paper else 1.1
    # The two-column conference paper follows IEEE-style justification.
    # In the single-column supplement, Vietnamese text becomes more readable
    # left-aligned because LibreOffice does not hyphenate it reliably.
    normal.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
        if paper else WD_ALIGN_PARAGRAPH.LEFT
    )

    for style_name, size, alignment in (
        ("Report Title", 14 if paper else 18, WD_ALIGN_PARAGRAPH.CENTER),
        ("Report Subtitle", 10 if paper else 12, WD_ALIGN_PARAGRAPH.CENTER),
        ("Report Meta", 8 if paper else 9, WD_ALIGN_PARAGRAPH.CENTER),
        ("Report Callout", 9 if paper else 10, WD_ALIGN_PARAGRAPH.LEFT),
        ("Report Authors", 10 if paper else 10, WD_ALIGN_PARAGRAPH.CENTER),
        ("Report Heading 1", 10 if paper else 13, WD_ALIGN_PARAGRAPH.CENTER),
        ("Report Heading 2", 10 if paper else 11, WD_ALIGN_PARAGRAPH.LEFT),
        ("Report Caption", 8 if paper else 9, WD_ALIGN_PARAGRAPH.CENTER),
        (
            "Report References",
            8 if paper else 9,
            WD_ALIGN_PARAGRAPH.JUSTIFY
            if paper else WD_ALIGN_PARAGRAPH.LEFT,
        ),
        ("Report Equation", 9.5 if paper else 10, WD_ALIGN_PARAGRAPH.CENTER),
    ):
        if style_name not in document.styles:
            style = document.styles.add_style(
                style_name, WD_STYLE_TYPE.PARAGRAPH
            )
        else:
            style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = "Heading" in style_name or style_name == "Report Title"
        style.paragraph_format.alignment = alignment
        style.paragraph_format.space_before = Pt(4 if "Heading" in style_name else 0)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.keep_with_next = (
            "Heading" in style_name or style_name == "Report Caption"
        )


def add_inline(paragraph, node, bold=False, italic=False):
    if isinstance(node, NavigableString):
        # HTML collapses source formatting whitespace, whereas python-docx
        # turns embedded newlines into explicit line breaks.  Normalizing here
        # prevents every wrapped source line from being fully justified as if
        # it were a complete line in the two-column paper.
        text = re.sub(r"\s+", " ", str(node))
        if text:
            run = paragraph.add_run(text)
            run.bold = bold
            run.italic = italic
        return
    if not isinstance(node, Tag):
        return
    if node.name == "br":
        paragraph.add_run().add_break()
        return
    next_bold = bold or node.name in {"b", "strong"}
    next_italic = italic or node.name in {"i", "em"}
    for child in node.children:
        add_inline(paragraph, child, next_bold, next_italic)


def add_text_element(document, element, style=None, list_style=None):
    paragraph = document.add_paragraph(
        style=list_style or style or document.styles["Normal"]
    )
    add_inline(paragraph, element)
    return paragraph


def add_figure(document, element, base_directory, paper):
    image = element.find("img")
    if image is None or not image.get("src"):
        return
    image_path = (base_directory / image["src"]).resolve()
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    width = Inches(3.27 if paper else 6.35)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width)
    paragraph.paragraph_format.keep_with_next = True
    caption = element.find("figcaption")
    if caption is not None:
        caption_paragraph = add_text_element(
            document, caption, style="Report Caption"
        )
        caption_paragraph.paragraph_format.keep_with_next = False


def add_table(document, element, paper):
    html_rows = element.find_all("tr", recursive=True)
    if not html_rows:
        return
    column_count = max(len(row.find_all(["th", "td"], recursive=False)) for row in html_rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    column_width = Inches((3.22 if paper else 6.45) / column_count)
    for column in table.columns:
        column.width = column_width
    tiny = "tiny" in (element.get("class") or [])
    compact = "compact" in (element.get("class") or [])
    size = 5.8 if paper and tiny else 6.2 if paper else 7.0 if tiny else 8.0 if compact else 8.5
    for row_index, html_row in enumerate(html_rows):
        cells = html_row.find_all(["th", "td"], recursive=False)
        row = table.add_row()
        prevent_row_split(row)
        if row_index == 0 and cells and all(cell.name == "th" for cell in cells):
            set_repeat_table_header(row)
        for index, html_cell in enumerate(cells):
            cell = row.cells[index]
            cell.width = column_width
            cell.text = html_cell.get_text(" ", strip=True)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if html_cell.name == "th":
                set_cell_shading(cell, "E7E6E6")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            set_cell_text_size(cell, size)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_list(document, element, ordered):
    # Word/LibreOffice may continue the same built-in list numbering across
    # unrelated <ol> elements.  Emit the marker explicitly so every HTML list
    # restarts deterministically and the PDF matches the source document.
    start = int(element.get("start", 1)) if ordered else 1
    for offset, item in enumerate(element.find_all("li", recursive=False)):
        paragraph = document.add_paragraph(style=document.styles["Normal"])
        paragraph.paragraph_format.left_indent = Mm(6)
        paragraph.paragraph_format.first_line_indent = Mm(-4)
        marker = f"{start + offset}. " if ordered else "• "
        paragraph.add_run(marker)
        add_inline(paragraph, item)


def add_children(document, container, base_directory, paper):
    in_references = False
    for element in container.children:
        if isinstance(element, NavigableString):
            continue
        if not isinstance(element, Tag):
            continue
        if element.name == "h1":
            add_text_element(document, element, style="Report Title")
        elif element.name == "h2":
            in_references = "TÀI LIỆU THAM KHẢO" in element.get_text().upper()
            add_text_element(document, element, style="Report Heading 1")
        elif element.name == "h3":
            add_text_element(document, element, style="Report Heading 2")
        elif element.name == "h4":
            paragraph = add_text_element(document, element)
            paragraph.runs[0].bold = True
            paragraph.paragraph_format.keep_with_next = True
        elif element.name == "p":
            style = "Report References" if in_references else None
            add_text_element(document, element, style=style)
        elif element.name == "div":
            classes = element.get("class") or []
            if "title" in classes:
                add_text_element(document, element, style="Report Title")
            elif "subtitle" in classes:
                add_text_element(document, element, style="Report Subtitle")
            elif "meta" in classes:
                add_text_element(document, element, style="Report Meta")
            elif "mine" in classes:
                paragraph = add_text_element(
                    document, element, style="Report Callout"
                )
                paragraph.paragraph_format.left_indent = Mm(5)
                paragraph.paragraph_format.right_indent = Mm(5)
                paragraph.paragraph_format.space_before = Pt(4)
                paragraph.paragraph_format.space_after = Pt(4)
                set_paragraph_shading(paragraph, "E8F1FB")
            elif "authors" in classes:
                add_text_element(document, element, style="Report Authors")
            elif "abstract" in classes:
                paragraph = add_text_element(document, element)
                paragraph.paragraph_format.left_indent = Mm(4)
                paragraph.paragraph_format.right_indent = Mm(4)
            elif "eq" in classes:
                add_text_element(document, element, style="Report Equation")
            elif "page-break" in classes:
                paragraph = document.add_paragraph()
                paragraph.add_run().add_break(WD_BREAK.PAGE)
            else:
                add_children(document, element, base_directory, paper)
        elif element.name == "figure":
            add_figure(document, element, base_directory, paper)
        elif element.name == "table":
            add_table(document, element, paper)
        elif element.name == "ul":
            add_list(document, element, ordered=False)
        elif element.name == "ol":
            add_list(document, element, ordered=True)


def convert(source, output, paper):
    soup = BeautifulSoup(source.read_text(encoding="utf-8"), "html.parser")
    document = Document()
    configure_styles(document, paper)
    first_section = document.sections[0]
    set_section_geometry(first_section)
    add_page_number(first_section)
    set_columns(first_section, 1)
    body = soup.body
    if body is None:
        raise RuntimeError("HTML document has no body")

    if paper:
        columns = body.find("div", class_="columns")
        for child in body.children:
            if child is columns:
                break
            if isinstance(child, Tag):
                wrapper = BeautifulSoup("<div></div>", "html.parser").div
                wrapper.append(child)
                add_children(document, wrapper, source.parent, paper=True)
        second = document.add_section(WD_SECTION.CONTINUOUS)
        set_section_geometry(second)
        set_columns(second, 2, space_twips=360)
        if columns is not None:
            add_children(document, columns, source.parent, paper=True)
    else:
        add_children(document, body, source.parent, paper=False)

    for paragraph in document.paragraphs:
        if paragraph.style.name not in {
            "Report Title",
            "Report Authors",
            "Report Heading 1",
            "Report Heading 2",
            "Report Caption",
            "Report Equation",
        }:
            paragraph.paragraph_format.widow_control = True
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run._element.get_or_add_rPr().rFonts.set(
                qn("w:eastAsia"), "Times New Roman"
            )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("source")
    parser.add_argument("output")
    parser.add_argument(
        "--paper",
        action="store_true",
        help="Use the REV-ECIT two-column paper layout.",
    )
    options = parser.parse_args()
    convert(
        Path(options.source).resolve(),
        Path(options.output).resolve(),
        options.paper,
    )
    print(Path(options.output).resolve())


if __name__ == "__main__":
    main()
