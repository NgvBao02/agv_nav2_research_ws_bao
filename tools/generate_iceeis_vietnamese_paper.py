#!/usr/bin/env python3
"""Generate the Vietnamese ICEEIS PSTMO paper from the IEEE A4 template.

The experimental tables and map figure are generated from the recorded CSV,
PGM maps, and RViz2 JSON paths in this repository.  The original IEEE
template is Strict OOXML, so LibreOffice is first used to create a temporary
Transitional OOXML copy that python-docx can edit without changing the source
template.
"""

from __future__ import annotations

import csv
import json
import math
import os
import statistics
import subprocess
import tempfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "final_bao_ICEEIS"
TEMPLATE = OUT_DIR / "IEEE Xplore conference-template-a4.docx"
ABSTRACT_TEMPLATE = ROOT / "abstract/abstract final 2.docx"
OUTPUT_DOCX = OUT_DIR / "ICEEIS_2026_PSTMO_ban_tieng_Viet.docx"
CSV_PATH = ROOT / "docs/pstmo_bao_cao_toan_dien_assets/benchmark_hinh_hoc_175_luot.csv"
RVIZ_DIR = ROOT / "docs/pstmo_bao_cao_toan_dien_assets/rviz_cases"
MAP_DIR = ROOT / "src/vacuum_robot_gazebo/maps"
ASSET_DIR = ROOT / "docs/pstmo_bao_cao_toan_dien_assets"

FONT_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf"
FONT_ITALIC = "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf"
TNR = "Times New Roman"


def set_cell_margins(cell, top=60, start=60, bottom=60, end=60):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, attrs in edges.items():
        tag = f"w:{edge_name}"
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key, value in attrs.items():
            edge.set(qn(f"w:{key}"), str(value))


def set_columns(section, count: int, space_twips: int = 360):
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))
    if qn("w:equalWidth") in cols.attrib:
        del cols.attrib[qn("w:equalWidth")]


def configure_section(section, columns=2):
    section.page_width = Inches(8.2677)
    section.page_height = Inches(11.6929)
    section.top_margin = Inches(0.866)
    section.bottom_margin = Inches(0.866)
    section.left_margin = Inches(0.630)
    section.right_margin = Inches(0.630)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    set_columns(section, columns)


def remove_all_content(doc):
    body = doc._element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                p._element.getparent().remove(p._element)


def set_run_font(run, size=10, bold=None, italic=None, color=None):
    run.font.name = TNR
    run._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def format_paragraph_runs(paragraph, size=10, bold=None, italic=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic)


def ensure_custom_styles(doc):
    styles = doc.styles

    def get_or_add(name):
        try:
            return styles[name]
        except KeyError:
            return styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

    normal = styles["Normal"]
    normal.font.name = TNR
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    normal.font.size = Pt(10)

    body = get_or_add("ICEEIS Body")
    body.font.name = TNR
    body._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    body.font.size = Pt(10)
    body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.line_spacing = 1.0
    body.paragraph_format.space_after = Pt(0)

    h1 = get_or_add("ICEEIS Heading 1")
    h1.font.name = TNR
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    h1.font.size = Pt(10)
    h1.font.bold = True
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(4)
    h1.paragraph_format.keep_with_next = True

    h2 = get_or_add("ICEEIS Heading 2")
    h2.font.name = TNR
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    h2.font.size = Pt(10)
    h2.font.italic = True
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(2)
    h2.paragraph_format.keep_with_next = True

    cap = get_or_add("ICEEIS Figure Caption")
    cap.font.name = TNR
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    cap.font.size = Pt(8)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = False

    ref = get_or_add("ICEEIS Reference")
    ref.font.name = TNR
    ref._element.rPr.rFonts.set(qn("w:eastAsia"), TNR)
    ref.font.size = Pt(8)
    ref.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ref.paragraph_format.left_indent = Inches(0.18)
    ref.paragraph_format.first_line_indent = Inches(-0.18)
    ref.paragraph_format.space_after = Pt(0)
    ref.paragraph_format.line_spacing = 1.0


def add_body(doc, text, first=False):
    p = doc.add_paragraph(style="ICEEIS Body")
    if not first:
        p.paragraph_format.first_line_indent = Inches(0.14)
    p.add_run(text)
    format_paragraph_runs(p, 10)
    return p


def add_heading1(doc, text):
    p = doc.add_paragraph(style="ICEEIS Heading 1")
    p.add_run(text)
    format_paragraph_runs(p, 10, bold=True)
    return p


def add_heading2(doc, text):
    p = doc.add_paragraph(style="ICEEIS Heading 2")
    p.add_run(text)
    format_paragraph_runs(p, 10, italic=True)
    return p


def add_equation(doc, lines, number):
    if isinstance(lines, str):
        lines = [lines]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(2.94)
    table.columns[1].width = Inches(0.32)
    table.rows[0].cells[0].width = Inches(2.94)
    table.rows[0].cells[1].width = Inches(0.32)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, start={"val": "nil"}, end={"val": "nil"})
    p = table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, 9, italic=True)
        run.font.name = "DejaVu Math TeX Gyre"
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            run._element.rPr.rFonts.set(qn(f"w:{attr}"), "DejaVu Math TeX Gyre")
        if i != len(lines) - 1:
            run.add_break()
    p2 = table.rows[0].cells[1].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p2.add_run(f"({number})")
    set_run_font(r, 9)
    return table


def add_figure(doc, image_path, width, caption, number):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph(style="ICEEIS Figure Caption")
    cap.paragraph_format.keep_together = True
    r = cap.add_run(f"Hình {number}. ")
    set_run_font(r, 8, bold=True)
    r = cap.add_run(caption)
    set_run_font(r, 8)
    return cap


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def add_author_grid(doc):
    """IEEE author block: three authors in row 1 and two centered in row 2."""
    authors = [
        ("1st Hai Linh Pham", "linh241632055@lms.utc.edu.vn", None),
        ("2nd Thi Phuong Thao Nguyen", "thao241632077@lms.utc.edu.vn", None),
        ("3rd Tien Cuong Nguyen", "cuong241632016@lms.utc.edu.vn", None),
        ("4th Thi Ly Pham*", "ptyldk@utc.edu.vn", "ORCID: 0009-0004-1170-6057"),
        ("5th Viet Bao Nguyen", "bao241632013@lms.utc.edu.vn", None),
    ]
    table = doc.add_table(rows=2, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.rows[0].cells[0].merge(table.rows[0].cells[1])
    table.rows[0].cells[2].merge(table.rows[0].cells[3])
    table.rows[0].cells[4].merge(table.rows[0].cells[5])
    table.rows[1].cells[0].merge(table.rows[1].cells[2])
    table.rows[1].cells[3].merge(table.rows[1].cells[5])
    slots = [table.rows[0].cells[0], table.rows[0].cells[2], table.rows[0].cells[4],
             table.rows[1].cells[0], table.rows[1].cells[3]]
    for cell, (name, email, orcid) in zip(slots, authors):
        cell.width = Inches(2.31 if cell in slots[:3] else 3.47)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell, top=20, start=25, bottom=45, end=25)
        set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"},
                        start={"val": "nil"}, end={"val": "nil"})
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(name)
        set_run_font(r, 10.2)
        for line, italic, size in [
            ("Faculty of Electrical-Electronic Engineering", True, 8.2),
            ("University of Transport and Communications", True, 8.2),
            ("Hanoi, Vietnam", False, 8.2),
            (email, False, 7.8),
        ]:
            p.add_run().add_break()
            r = p.add_run(line)
            set_run_font(r, size, italic=italic)
        if orcid:
            p.add_run().add_break()
            r = p.add_run(orcid)
            set_run_font(r, 7.8)
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"},
                            start={"val": "nil"}, end={"val": "nil"})
    after = doc.add_paragraph()
    after.alignment = WD_ALIGN_PARAGRAPH.CENTER
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(5)
    r = after.add_run("*Corresponding author")
    set_run_font(r, 8, italic=True)
    return table


def add_table_caption(doc, number, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"BẢNG {number}\n")
    set_run_font(r, 8, bold=True)
    r = p.add_run(caption.upper())
    set_run_font(r, 8)
    return p


def add_ieee_table(doc, headers, rows, widths=None, font_size=7.5, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths is None:
        widths = [3.25 / len(headers)] * len(headers)
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(int(sum(widths) * 1440)))
    for column, grid_col, width in zip(table.columns, table._tbl.tblGrid.gridCol_lst, widths):
        column.width = Inches(width)
        grid_col.set(qn("w:w"), str(int(width * 1440)))
    hdr = table.rows[0]
    tr_pr = hdr._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)
    for j, header in enumerate(headers):
        cell = hdr.cells[j]
        cell.width = Inches(widths[j])
        shade_cell(cell, "D9E2F3")
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, font_size, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for j, value in enumerate(row):
            cells[j].width = Inches(widths[j])
            set_cell_margins(cells[j])
            p = cells[j].paragraphs[0]
            p.alignment = alignments[j] if alignments else (WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, font_size)
            cells[j].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    border = {"val": "single", "sz": "4", "space": "0", "color": "808080"}
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell in row.cells:
            set_cell_border(cell, top=border, bottom=border, start=border, end=border, insideH=border, insideV=border)
    return table


def add_reference(doc, number, text):
    p = doc.add_paragraph(style="ICEEIS Reference")
    p.add_run(f"[{number}] ")
    p.add_run(text)
    format_paragraph_runs(p, 8)
    return p


def load_results():
    selected = {"open_arena", "narrow_aisles", "warehouse_cross_aisles"}
    with CSV_PATH.open(encoding="utf-8-sig", newline="") as handle:
        rows = [r for r in csv.DictReader(handle) if r["environment"] in selected]
    methods = ["raw", "simple", "savitzky_golay", "constrained", "pstmo"]
    metrics = {}
    for env in sorted(selected | {"all"}):
        source = rows if env == "all" else [r for r in rows if r["environment"] == env]
        metrics[env] = {}
        for method in methods:
            subset = [r for r in source if r["method"] == method and r["success"] == "True"]
            metrics[env][method] = {
                "n": len(subset),
                "L": statistics.fmean(float(r["path_length_m"]) for r in subset),
                "E": statistics.fmean(float(r["curvature_energy_1pm"]) for r in subset),
                "T": 1000.0 * statistics.fmean(float(r["algorithm_time_s"]) for r in subset),
            }
    return rows, metrics


def draw_dashed(draw, points, fill, width, dash=18, gap=11):
    for a, b in zip(points, points[1:]):
        dx, dy = b[0] - a[0], b[1] - a[1]
        length = math.hypot(dx, dy)
        if length <= 1e-9:
            continue
        ux, uy = dx / length, dy / length
        s = 0.0
        while s < length:
            e = min(s + dash, length)
            draw.line((a[0] + ux * s, a[1] + uy * s, a[0] + ux * e, a[1] + uy * e), fill=fill, width=width)
            s += dash + gap


def draw_arrow(draw, start, end, fill="#222222", width=5, head=18):
    draw.line((start, end), fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    for delta in (2.55, -2.55):
        tip = (end[0] + head * math.cos(angle + delta), end[1] + head * math.sin(angle + delta))
        draw.line((end, tip), fill=fill, width=width)


def make_problem_figure(path: Path):
    """Compact one-column schematic of the corner-discontinuity problem."""
    canvas = Image.new("RGB", (1500, 690), "white")
    draw = ImageDraw.Draw(canvas)
    label = ImageFont.truetype(FONT_REGULAR, 31)
    italic = ImageFont.truetype(FONT_ITALIC, 30)
    bold = ImageFont.truetype(FONT_BOLD, 31)

    panels = [(55, 70, 690, 535), (785, 70, 660, 535)]
    for x, y, w, h in panels:
        draw.rectangle((x, y, x + w, y + h), outline="#555555", width=3)
    draw.text((400, 25), "(a) Đa tuyến tại góc", font=bold, fill="#111111", anchor="ma")
    draw.text((1115, 25), "(b) Chuyển tiếp cục bộ", font=bold, fill="#111111", anchor="ma")

    # Raw polyline: the tangent direction changes instantaneously at V.
    a, vtx, b = (135, 495), (505, 495), (655, 170)
    draw.line((a, vtx, b), fill="#555555", width=9, joint="curve")
    draw.ellipse((vtx[0] - 11, vtx[1] - 11, vtx[0] + 11, vtx[1] + 11), fill="#111111")
    draw.text((520, 525), "V", font=italic, fill="#111111")
    draw_arrow(draw, (300, 495), (455, 495), "#D55E00", 7, 24)
    draw_arrow(draw, (530, 440), (600, 285), "#D55E00", 7, 24)
    draw.arc((420, 325, 595, 510), 272, 336, fill="#7A3DF0", width=6)
    draw.text((462, 365), "Δψ", font=italic, fill="#7A3DF0")
    draw.text((400, 575), "đổi hướng tức thời ⇒ κ không bị chặn", font=label, fill="#222222", anchor="ma")

    # Local transition: keep the same corridor and replace only the corner vicinity.
    a2, v2, b2 = (850, 495), (1175, 495), (1355, 170)
    draw_dashed(draw, [a2, v2, b2], "#777777", 6, dash=22, gap=14)
    A, B = (1000, 495), (1245, 365)
    ctrl = [A, (1085, 495), (1145, 495), (1200, 450), (1220, 405), B]
    curve = []
    for j in range(201):
        t = j / 200
        x = y = 0.0
        for i, p in enumerate(ctrl):
            c = math.comb(5, i) * (1 - t) ** (5 - i) * t**i
            x += c * p[0]
            y += c * p[1]
        curve.append((x, y))
    draw.line((a2, A), fill="#0072B2", width=9)
    draw.line(curve, fill="#0072B2", width=11, joint="curve")
    draw.line((B, b2), fill="#0072B2", width=9)
    for p, name, off in [(A, "A", (-20, 24)), (B, "B", (10, -38))]:
        draw.ellipse((p[0] - 10, p[1] - 10, p[0] + 10, p[1] + 10), fill="#D55E00")
        draw.text((p[0] + off[0], p[1] + off[1]), name, font=italic, fill="#111111")
    draw.text((1115, 575), "giữ tuyến; phân bố chuyển hướng trên A–B", font=label, fill="#222222", anchor="ma")
    canvas.save(path, dpi=(300, 300))


def make_footprint_figure(path: Path):
    """Differential-drive kinematics and footprint-sweep schematic."""
    canvas = Image.new("RGB", (1500, 790), "white")
    draw = ImageDraw.Draw(canvas, "RGBA")
    label = ImageFont.truetype(FONT_REGULAR, 30)
    italic = ImageFont.truetype(FONT_ITALIC, 29)
    bold = ImageFont.truetype(FONT_BOLD, 31)

    draw.rectangle((45, 65, 715, 720), outline="#555555", width=3)
    draw.rectangle((785, 65, 1455, 720), outline="#555555", width=3)
    draw.text((380, 22), "(a) Robot hai bánh vi sai", font=bold, fill="#111111", anchor="ma")
    draw.text((1120, 22), "(b) Vùng quét hình bao", font=bold, fill="#111111", anchor="ma")

    # Plan view of robot.
    cx, cy = 380, 390
    draw.rounded_rectangle((235, 255, 525, 525), radius=28, fill="#D9EAF4", outline="#222222", width=5)
    draw.rectangle((205, 285, 240, 385), fill="#444444")
    draw.rectangle((520, 285, 555, 385), fill="#444444")
    draw.line((220, 235, 540, 235), fill="#555555", width=4)
    draw.line((220, 220, 220, 250), fill="#555555", width=4)
    draw.line((540, 220, 540, 250), fill="#555555", width=4)
    draw.text((380, 200), "b", font=italic, fill="#111111", anchor="ma")
    draw_arrow(draw, (380, 390), (380, 155), "#0072B2", 8, 26)
    draw.text((400, 165), "v", font=italic, fill="#0072B2")
    draw.arc((325, 325, 500, 505), 285, 85, fill="#7A3DF0", width=7)
    draw.text((485, 410), "ω", font=italic, fill="#7A3DF0")
    draw_arrow(draw, (220, 335), (220, 195), "#D55E00", 7, 23)
    draw_arrow(draw, (540, 335), (540, 135), "#009E73", 7, 23)
    draw.text((175, 170), "vL", font=italic, fill="#D55E00")
    draw.text((555, 125), "vR", font=italic, fill="#009E73")
    draw.text((380, 640), "vL,R = v(1 ∓ bκ/2)", font=label, fill="#222222", anchor="ma")

    # Swept footprint along a curve; obstacle deliberately outside centerline but near the body.
    curve = []
    for j in range(181):
        t = j / 180
        curve.append((850 + 500 * t, 590 - 390 * (3*t*t - 2*t*t*t)))
    draw.line(curve, fill="#0072B2", width=8)
    draw.rectangle((1230, 190, 1410, 335), fill="#4D004B", outline="#111111", width=4)
    draw.text((1320, 165), "vật cản", font=label, fill="#111111", anchor="ma")

    def pose_rect(center, theta, alpha):
        length, width = 150, 105
        pts = []
        for dx, dy in [(-length/2, -width/2), (length/2, -width/2), (length/2, width/2), (-length/2, width/2)]:
            x = center[0] + dx * math.cos(theta) - dy * math.sin(theta)
            y = center[1] + dx * math.sin(theta) + dy * math.cos(theta)
            pts.append((x, y))
        draw.polygon(pts, fill=(0, 114, 178, alpha), outline=(0, 72, 115, 230))
        return pts

    for t in (0.08, 0.28, 0.48, 0.68, 0.88):
        center = (850 + 500 * t, 590 - 390 * (3*t*t - 2*t*t*t))
        dx = 500
        dy = -390 * (6*t - 6*t*t)
        pose_rect(center, math.atan2(dy, dx), 55)
    draw.text((1120, 665), "hợp các hình bao dọc đường", font=label, fill="#222222", anchor="ma")
    canvas.save(path, dpi=(300, 300))


def make_dp_figure(path: Path):
    """Shared-edge budget and dynamic-programming state compatibility."""
    canvas = Image.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(canvas)
    label = ImageFont.truetype(FONT_REGULAR, 30)
    italic = ImageFont.truetype(FONT_ITALIC, 29)
    bold = ImageFont.truetype(FONT_BOLD, 31)

    draw.text((375, 28), "(a) Ngân sách cạnh chung", font=bold, fill="#111111", anchor="ma")
    draw.text((1125, 28), "(b) Chọn chuỗi tương thích", font=bold, fill="#111111", anchor="ma")
    draw.rectangle((40, 70, 730, 700), outline="#555555", width=3)
    draw.rectangle((770, 70, 1460, 700), outline="#555555", width=3)

    p0, vi, vip1, p3 = (95, 520), (255, 520), (560, 520), (680, 310)
    draw.line((p0, vi, vip1, p3), fill="#777777", width=7, joint="curve")
    draw.arc((175, 350, 365, 535), 88, 180, fill="#0072B2", width=10)
    draw.arc((465, 380, 645, 560), 270, 338, fill="#0072B2", width=10)
    for p, name in [(vi, "Vi"), (vip1, "Vi+1")]:
        draw.ellipse((p[0]-10, p[1]-10, p[0]+10, p[1]+10), fill="#111111")
        draw.text((p[0], p[1]+38), name, font=italic, fill="#111111", anchor="ma")
    ydim = 610
    for x in (255, 560):
        draw.line((x, ydim-12, x, ydim+12), fill="#333333", width=4)
    draw.line((255, ydim, 560, ydim), fill="#333333", width=4)
    draw.text((407, ydim+24), "ℓi", font=italic, fill="#111111", anchor="ma")
    draw.line((255, 570, 345, 570), fill="#D55E00", width=7)
    draw.text((300, 555), "di", font=italic, fill="#D55E00", anchor="ma")
    draw.line((475, 570, 560, 570), fill="#009E73", width=7)
    draw.text((520, 555), "di+1", font=italic, fill="#009E73", anchor="ma")
    draw.line((345, 570, 475, 570), fill="#7A3DF0", width=7)
    draw.text((410, 555), "m", font=italic, fill="#7A3DF0", anchor="ma")

    cols = [860, 1060, 1260, 1400]
    states = [("qua", "#777777"), ("cong", "#0072B2"), ("quay", "#D55E00")]
    for ci, x in enumerate(cols):
        draw.text((x, 115), f"góc {ci+1}", font=label, fill="#111111", anchor="ma")
        for si, (name, color) in enumerate(states):
            y = 240 + si * 165
            draw.ellipse((x-16, y-16, x+16, y+16), fill=color, outline="#111111", width=2)
            if ci == 0:
                draw.text((x-35, y), name, font=label, fill="#111111", anchor="rm")
            if ci < len(cols)-1:
                for sj in range(3):
                    # Deliberately omit a few incompatible transitions.
                    if (ci, si, sj) not in {(0, 1, 1), (1, 0, 1), (2, 1, 0)}:
                        y2 = 240 + sj * 165
                        draw.line((x+18, y, cols[ci+1]-18, y2), fill="#C6CED8", width=3)
    selected = [(860, 405), (1060, 570), (1260, 405), (1400, 240)]
    draw.line(selected, fill="#7A3DF0", width=8)
    for x, y in selected:
        draw.ellipse((x-20, y-20, x+20, y+20), outline="#7A3DF0", width=6)
    draw.text((1115, 670), "chuỗi chi phí nhỏ nhất trên các cạnh hợp lệ", font=label, fill="#222222", anchor="ma")
    canvas.save(path, dpi=(300, 300))


def make_environment_figure(path: Path, env: str, scenario: str):
    """One-column map figure with a full-map panel and an enlarged path corner."""
    pgm = Image.open(MAP_DIR / f"{env}.pgm").convert("L")
    meta = yaml.safe_load((MAP_DIR / f"{env}.yaml").read_text(encoding="utf-8"))
    resolution = float(meta["resolution"])
    xmin, ymin = float(meta["origin"][0]), float(meta["origin"][1])
    xmax, ymax = xmin + pgm.width * resolution, ymin + pgm.height * resolution
    data = json.loads((RVIZ_DIR / f"{env}__{scenario}__ThetaStar.json").read_text(encoding="utf-8"))
    raw_pose = data["paths"]["raw"]["poses"]
    pstmo_pose = data["paths"]["pstmo"]["poses"]

    # Center the inset on the raw point farthest from the smoothed path.
    focus = max(raw_pose, key=lambda p: min((p["x"]-s["x"])**2 + (p["y"]-s["y"])**2 for s in pstmo_pose))
    canvas = Image.new("RGB", (1600, 980), "white")
    draw = ImageDraw.Draw(canvas)
    label = ImageFont.truetype(FONT_REGULAR, 29)
    bold = ImageFont.truetype(FONT_BOLD, 30)
    small = ImageFont.truetype(FONT_REGULAR, 25)

    def render_panel(box, world_bounds, panel_label, show_endpoints=False):
        left, top, right, bottom = box
        wx0, wy0, wx1, wy1 = world_bounds
        span_x, span_y = wx1-wx0, wy1-wy0
        avail_w, avail_h = right-left, bottom-top
        scale = min(avail_w/span_x, avail_h/span_y)
        out_w, out_h = int(span_x*scale), int(span_y*scale)
        ox, oy = left+(avail_w-out_w)//2, top+(avail_h-out_h)//2
        px0 = max(0, int((wx0-xmin)/resolution)); px1 = min(pgm.width, int(math.ceil((wx1-xmin)/resolution)))
        py0_bottom = max(0, int((wy0-ymin)/resolution)); py1_bottom = min(pgm.height, int(math.ceil((wy1-ymin)/resolution)))
        # PGM row zero is the top of the displayed map.
        crop = pgm.crop((px0, pgm.height-py1_bottom, px1, pgm.height-py0_bottom)).convert("RGB")
        crop = crop.resize((out_w, out_h), Image.Resampling.NEAREST)
        canvas.paste(crop, (ox, oy))

        def cv(p):
            return (ox+(p["x"]-wx0)*scale, oy+out_h-(p["y"]-wy0)*scale)
        # Draw a metric grid; x/y use the same scale, so cells stay square.
        grid_step = 1.0 if max(span_x, span_y) > 4 else 0.25
        gx = math.ceil(wx0/grid_step)*grid_step
        while gx <= wx1+1e-9:
            x = ox+(gx-wx0)*scale
            draw.line((x, oy, x, oy+out_h), fill="#D9D9D9", width=2)
            gx += grid_step
        gy = math.ceil(wy0/grid_step)*grid_step
        while gy <= wy1+1e-9:
            y = oy+out_h-(gy-wy0)*scale
            draw.line((ox, y, ox+out_w, y), fill="#D9D9D9", width=2)
            gy += grid_step
        raw = [cv(p) for p in raw_pose if wx0 <= p["x"] <= wx1 and wy0 <= p["y"] <= wy1]
        sm = [cv(p) for p in pstmo_pose if wx0 <= p["x"] <= wx1 and wy0 <= p["y"] <= wy1]
        if len(sm) > 1:
            draw.line(sm, fill="#0072B2", width=9, joint="curve")
        if len(raw) > 1:
            draw_dashed(draw, raw, "#D55E00", 7, dash=20, gap=12)
        if show_endpoints:
            sx, sy, _ = data["start"]; gxv, gyv, _ = data["goal"]
            for xw, yw, color, shape in [(sx, sy, "#009E73", "circle"), (gxv, gyv, "#CC79A7", "diamond")]:
                if wx0 <= xw <= wx1 and wy0 <= yw <= wy1:
                    x, y = ox+(xw-wx0)*scale, oy+out_h-(yw-wy0)*scale
                    if shape == "circle": draw.ellipse((x-11,y-11,x+11,y+11), fill=color, outline="#111111", width=2)
                    else: draw.polygon([(x,y-14),(x+14,y),(x,y+14),(x-14,y)], fill=color, outline="#111111")
        draw.rectangle((ox, oy, ox+out_w, oy+out_h), outline="#222222", width=4)
        draw.text(((left+right)/2, top-10), panel_label, font=bold, fill="#111111", anchor="ma")
        return (ox, oy, ox+out_w, oy+out_h)

    full_box = render_panel((45, 85, 975, 800), (xmin, ymin, xmax, ymax), "(a) Toàn bản đồ", True)
    half = 0.75 if env == "open_arena" else 1.05
    zcx, zcy = focus["x"], focus["y"]
    zbox = render_panel((1025, 85, 1555, 800), (zcx-half, zcy-half, zcx+half, zcy+half), "(b) Phóng đại góc", False)
    # Connector from focus to inset without crossing the map data excessively.
    fx = full_box[0] + (zcx-xmin)/(xmax-xmin)*(full_box[2]-full_box[0])
    fy = full_box[3] - (zcy-ymin)/(ymax-ymin)*(full_box[3]-full_box[1])
    draw.rectangle((fx-22, fy-22, fx+22, fy+22), outline="#7A3DF0", width=5)
    draw.line((fx+24, fy, zbox[0]-12, zbox[1]+45), fill="#7A3DF0", width=3)

    y = 875
    draw_dashed(draw, [(285,y),(405,y)], "#D55E00", 7, dash=20, gap=12)
    draw.text((425,y), "Raw", font=label, fill="#111111", anchor="lm")
    draw.line((585,y,705,y), fill="#0072B2", width=9)
    draw.text((725,y), "PSTMO", font=label, fill="#111111", anchor="lm")
    draw.ellipse((935,y-10,955,y+10), fill="#009E73", outline="#111111", width=2)
    draw.text((970,y), "bắt đầu", font=small, fill="#111111", anchor="lm")
    draw.polygon([(1215,y-13),(1228,y),(1215,y+13),(1202,y)], fill="#CC79A7", outline="#111111")
    draw.text((1240,y), "đích", font=small, fill="#111111", anchor="lm")
    canvas.save(path, dpi=(300, 300))


def make_map_figure(path: Path):
    cases = [
        ("open_arena", "center_block_detour", "Không gian mở"),
        ("narrow_aisles", "southwest_northeast_weave", "Lối đi hẹp"),
        ("warehouse_cross_aisles", "cross_aisle_transfer", "Kho có lối giao cắt"),
    ]
    canvas = Image.new("RGB", (2400, 850), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = ImageFont.truetype(FONT_BOLD, 36)
    label_font = ImageFont.truetype(FONT_REGULAR, 27)
    small_font = ImageFont.truetype(FONT_REGULAR, 23)
    draw.text((1200, 28), "Dữ liệu bản đồ và đường đi ghi từ RViz2", font=title_font, fill="#111111", anchor="ma")

    legend_y = 80
    draw_dashed(draw, [(740, legend_y), (840, legend_y)], "#D55E00", 6, dash=20, gap=12)
    draw.text((858, legend_y), "Đường thô", font=label_font, fill="#111111", anchor="lm")
    draw.line((1090, legend_y, 1190, legend_y), fill="#0072B2", width=7)
    draw.text((1208, legend_y), "PSTMO", font=label_font, fill="#111111", anchor="lm")
    draw.ellipse((1430, legend_y - 9, 1448, legend_y + 9), fill="#009E73", outline="#111111", width=2)
    draw.text((1462, legend_y), "Bắt đầu", font=label_font, fill="#111111", anchor="lm")
    draw.polygon([(1675, legend_y - 11), (1686, legend_y), (1675, legend_y + 11), (1664, legend_y)], fill="#CC79A7", outline="#111111")
    draw.text((1698, legend_y), "Đích", font=label_font, fill="#111111", anchor="lm")

    panel_w = 730
    panel_gap = 55
    panel_xs = [60, 60 + panel_w + panel_gap, 60 + 2 * (panel_w + panel_gap)]
    top = 135
    map_w = 690
    for index, ((env, scenario, title), panel_x) in enumerate(zip(cases, panel_xs)):
        pgm = Image.open(MAP_DIR / f"{env}.pgm").convert("L")
        w, h = pgm.size
        map_meta = yaml.safe_load((MAP_DIR / f"{env}.yaml").read_text(encoding="utf-8"))
        resolution = float(map_meta["resolution"])
        origin = (float(map_meta["origin"][0]), float(map_meta["origin"][1]))
        map_h = round(map_w * h / w)
        left = panel_x + 20
        upper = top + 62
        map_img = pgm.resize((map_w, map_h), Image.Resampling.NEAREST).convert("RGB")
        canvas.paste(map_img, (left, upper))
        map_draw = ImageDraw.Draw(canvas)

        # A 0.5 m grid is shown with equal x/y scale; individual occupancy
        # cells remain square because the source is resized by nearest-neighbor.
        xmin, ymin = origin
        xmax, ymax = xmin + w * resolution, ymin + h * resolution
        for gx_i in range(math.ceil(xmin * 2), math.floor(xmax * 2) + 1):
            gx = gx_i / 2
            px = left + (gx - xmin) / (xmax - xmin) * map_w
            map_draw.line((px, upper, px, upper + map_h), fill="#E6E6E6", width=1)
        for gy_i in range(math.ceil(ymin * 2), math.floor(ymax * 2) + 1):
            gy = gy_i / 2
            py = upper + map_h - (gy - ymin) / (ymax - ymin) * map_h
            map_draw.line((left, py, left + map_w, py), fill="#E6E6E6", width=1)

        json_path = RVIZ_DIR / f"{env}__{scenario}__ThetaStar.json"
        data = json.loads(json_path.read_text(encoding="utf-8"))

        def transform(poses):
            pts = []
            for pose in poses:
                px = left + (pose["x"] - xmin) / (xmax - xmin) * map_w
                py = upper + map_h - (pose["y"] - ymin) / (ymax - ymin) * map_h
                pts.append((px, py))
            return pts

        raw = transform(data["paths"]["raw"]["poses"])
        pstmo = transform(data["paths"]["pstmo"]["poses"])
        map_draw.line(pstmo, fill="#0072B2", width=7, joint="curve")
        draw_dashed(map_draw, raw, "#D55E00", 6, dash=18, gap=10)
        sx, sy, _ = data["start"]
        gx, gy, _ = data["goal"]
        sp = (left + (sx - xmin) / (xmax - xmin) * map_w, upper + map_h - (sy - ymin) / (ymax - ymin) * map_h)
        gp = (left + (gx - xmin) / (xmax - xmin) * map_w, upper + map_h - (gy - ymin) / (ymax - ymin) * map_h)
        map_draw.ellipse((sp[0] - 10, sp[1] - 10, sp[0] + 10, sp[1] + 10), fill="#009E73", outline="#111111", width=2)
        map_draw.polygon([(gp[0], gp[1] - 13), (gp[0] + 13, gp[1]), (gp[0], gp[1] + 13), (gp[0] - 13, gp[1])], fill="#CC79A7", outline="#111111")
        map_draw.rectangle((left, upper, left + map_w, upper + map_h), outline="#222222", width=3)

        # Major ticks outside the map so labels never cover the data.
        for tick in (-6, -3, 0, 3, 6):
            if xmin <= tick <= xmax:
                px = left + (tick - xmin) / (xmax - xmin) * map_w
                map_draw.line((px, upper + map_h, px, upper + map_h + 8), fill="#222222", width=2)
                map_draw.text((px, upper + map_h + 10), str(tick), font=small_font, fill="#111111", anchor="ma")
        for tick in (-4, -2, 0, 2, 4):
            if ymin <= tick <= ymax:
                py = upper + map_h - (tick - ymin) / (ymax - ymin) * map_h
                map_draw.line((left - 8, py, left, py), fill="#222222", width=2)
                map_draw.text((left - 12, py), str(tick), font=small_font, fill="#111111", anchor="rm")
        map_draw.text((left + map_w / 2, upper + map_h + 43), "x (m)", font=small_font, fill="#111111", anchor="ma")
        map_draw.text((left - 52, upper + map_h / 2), "y (m)", font=small_font, fill="#111111", anchor="mm")
        map_draw.text((left + map_w / 2, top + 18), f"({chr(97 + index)}) {title}", font=label_font, fill="#111111", anchor="ma")

    canvas.save(path, dpi=(300, 300))


def make_bezier_figure(path: Path):
    canvas = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(canvas)
    label = ImageFont.truetype(FONT_REGULAR, 34)
    small = ImageFont.truetype(FONT_ITALIC, 33)
    V = (0.0, 0.0)
    angle = math.radians(58)
    u = (1.0, 0.0)
    v = (math.cos(angle), math.sin(angle))
    d = 1.0
    q = 0.34
    A = (-d, 0.0)
    B = (d * v[0], d * v[1])
    points = [
        A,
        (A[0] + q * u[0], A[1] + q * u[1]),
        (A[0] + 2 * q * u[0], A[1] + 2 * q * u[1]),
        (B[0] - 2 * q * v[0], B[1] - 2 * q * v[1]),
        (B[0] - q * v[0], B[1] - q * v[1]),
        B,
    ]

    def cv(p):
        return (790 + 580 * p[0], 520 - 500 * p[1])

    draw.line((cv((-1.25, 0)), cv((0.08, 0))), fill="#555555", width=5)
    draw.line((cv((0, 0)), cv((0.72 * v[0], 0.72 * v[1]))), fill="#555555", width=5)
    ctrl = [cv(p) for p in points]
    draw.line(ctrl, fill="#8C8C8C", width=4)

    curve = []
    for j in range(301):
        t = j / 300
        x = y = 0.0
        for i, p in enumerate(points):
            bcoef = math.comb(5, i) * (1 - t) ** (5 - i) * t**i
            x += bcoef * p[0]
            y += bcoef * p[1]
        curve.append(cv((x, y)))
    draw.line(curve, fill="#0072B2", width=11, joint="curve")

    label_positions = [
        (ctrl[0][0]-55, ctrl[0][1]+70),
        (ctrl[1][0], ctrl[1][1]+78),
        (ctrl[2][0]+12, ctrl[2][1]+92),
        (ctrl[3][0]+65, ctrl[3][1]-86),
        (ctrl[4][0]+78, ctrl[4][1]-78),
        (ctrl[5][0]+82, ctrl[5][1]-70),
    ]
    for i, (p, label_pos) in enumerate(zip(ctrl, label_positions)):
        draw.ellipse((p[0] - 12, p[1] - 12, p[0] + 12, p[1] + 12), fill="#D55E00", outline="#111111", width=3)
        draw.line((p, label_pos), fill="#777777", width=2)
        box = draw.textbbox(label_pos, f"P{i}", font=label, anchor="mm")
        draw.rectangle((box[0]-8, box[1]-4, box[2]+8, box[3]+4), fill="white")
        draw.text(label_pos, f"P{i}", font=label, fill="#111111", anchor="mm")
    vv = cv(V)
    draw.ellipse((vv[0] - 9, vv[1] - 9, vv[0] + 9, vv[1] + 9), fill="#222222")
    v_label = (vv[0]+45, vv[1]+62)
    draw.line((vv, v_label), fill="#777777", width=2)
    box = draw.textbbox(v_label, "V", font=label, anchor="mm")
    draw.rectangle((box[0]-8, box[1]-4, box[2]+8, box[3]+4), fill="white")
    draw.text(v_label, "V", font=label, fill="#111111", anchor="mm")
    # The explanatory legend occupies a separate band below the geometry.
    legend_y = 820
    draw.line((150, legend_y, 260, legend_y), fill="#8C8C8C", width=4)
    draw.text((285, legend_y), "Đa giác điều khiển", font=label, fill="#111111", anchor="lm")
    draw.line((650, legend_y, 760, legend_y), fill="#0072B2", width=11)
    draw.text((785, legend_y), "Đường Bézier bậc năm", font=label, fill="#111111", anchor="lm")
    draw.ellipse((1240, legend_y-12, 1264, legend_y+12), fill="#D55E00", outline="#111111", width=3)
    draw.text((1285, legend_y), "Điểm điều khiển", font=label, fill="#111111", anchor="lm")
    canvas.save(path, dpi=(300, 300))


def make_results_figure(path: Path, metrics):
    canvas = Image.new("RGB", (2400, 920), "white")
    draw = ImageDraw.Draw(canvas)
    title_font = ImageFont.truetype(FONT_BOLD, 35)
    label_font = ImageFont.truetype(FONT_REGULAR, 27)
    small_font = ImageFont.truetype(FONT_REGULAR, 23)
    bold_small = ImageFont.truetype(FONT_BOLD, 23)
    methods = ["raw", "simple", "savitzky_golay", "constrained", "pstmo"]
    labels = {"raw": "Raw", "simple": "Simple", "savitzky_golay": "S–G", "constrained": "Constrained", "pstmo": "PSTMO"}
    colors = {"raw": "#777777", "simple": "#009E73", "savitzky_golay": "#E69F00", "constrained": "#CC79A7", "pstmo": "#0072B2"}
    envs = ["open_arena", "narrow_aisles", "warehouse_cross_aisles"]
    env_labels = ["Không gian mở", "Lối đi hẹp", "Kho giao cắt"]

    draw.text((1200, 25), "Đánh đổi mức uốn hình học và thời gian xử lý", font=title_font, fill="#111111", anchor="ma")
    legend_y = 83
    legend_x = 475
    for method in methods:
        color = colors[method]
        draw.rectangle((legend_x, legend_y - 13, legend_x + 32, legend_y + 13), fill=color, outline="#111111", width=2)
        draw.text((legend_x + 44, legend_y), labels[method], font=label_font, fill="#111111", anchor="lm")
        legend_x += 285 if method != "savitzky_golay" else 245

    panels = [
        (95, 155, 1120, 690, "Eκ trung bình (thang logarit)", "Eκ (m⁻¹)", "E"),
        (1280, 155, 1025, 690, "Thời gian xử lý trung bình", "T (ms)", "T"),
    ]
    for panel_index, (left, top, width, height, panel_title, y_label, key) in enumerate(panels):
        plot_left = left + 105
        plot_top = top + 70
        plot_right = left + width - 35
        plot_bottom = top + height - 95
        draw.text((left + width / 2, top + 15), f"({chr(97 + panel_index)}) {panel_title}", font=label_font, fill="#111111", anchor="ma")
        draw.line((plot_left, plot_top, plot_left, plot_bottom), fill="#222222", width=3)
        draw.line((plot_left, plot_bottom, plot_right, plot_bottom), fill="#222222", width=3)
        if key == "E":
            ticks = [1, 2, 5, 10, 20, 50, 100, 200, 400]
            ymin, ymax = 1.0, 400.0

            def ypos(value):
                return plot_bottom - (math.log10(max(value, ymin)) - math.log10(ymin)) / (math.log10(ymax) - math.log10(ymin)) * (plot_bottom - plot_top)
        else:
            ticks = [0, 50, 100, 150, 200, 250]
            ymin, ymax = 0.0, 250.0

            def ypos(value):
                return plot_bottom - (value - ymin) / (ymax - ymin) * (plot_bottom - plot_top)
        for tick in ticks:
            py = ypos(tick)
            draw.line((plot_left, py, plot_right, py), fill="#DDDDDD", width=2)
            draw.text((plot_left - 13, py), str(tick), font=small_font, fill="#111111", anchor="rm")
        draw.text((left + 21, (plot_top + plot_bottom) / 2), y_label, font=small_font, fill="#111111", anchor="mm")
        group_width = (plot_right - plot_left) / len(envs)
        bar_width = group_width * 0.135
        for env_index, (env, env_label) in enumerate(zip(envs, env_labels)):
            center = plot_left + group_width * (env_index + 0.5)
            for method_index, method in enumerate(methods):
                value = metrics[env][method][key]
                x0 = center + (method_index - 2.5) * bar_width
                x1 = x0 + bar_width * 0.82
                y0 = ypos(value)
                draw.rectangle((x0, y0, x1, plot_bottom), fill=colors[method], outline="#111111", width=2)
            draw.text((center, plot_bottom + 18), env_label, font=small_font, fill="#111111", anchor="ma")
        draw.rectangle((plot_left, plot_top, plot_right, plot_bottom), outline="#222222", width=3)
    canvas.save(path, dpi=(300, 300))


# The following publication figures deliberately contain one analytical view
# each.  They replace the earlier slide-like multi-panel compositions.
def make_corner_figure_v2(path: Path):
    canvas = Image.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(canvas)
    label = ImageFont.truetype(FONT_REGULAR, 32)
    italic = ImageFont.truetype(FONT_ITALIC, 31)

    A0, V, B0 = (130, 610), (850, 610), (1290, 145)
    draw_dashed(draw, [A0, V, B0], "#666666", 7, dash=24, gap=15)
    A, B = (520, 610), (1010, 440)
    ctrl = [A, (660, 610), (760, 610), (900, 555), (955, 495), B]
    curve = []
    for j in range(301):
        t = j / 300
        x = y = 0.0
        for i, p in enumerate(ctrl):
            c = math.comb(5, i) * (1-t)**(5-i) * t**i
            x += c*p[0]; y += c*p[1]
        curve.append((x, y))
    draw.line((A0, A), fill="#0072B2", width=10)
    draw.line(curve, fill="#0072B2", width=12, joint="curve")
    draw.line((B, B0), fill="#0072B2", width=10)
    for p, name, dx, dy in [(A, "A", -18, 28), (V, "V", 10, 28), (B, "B", 14, -42)]:
        draw.ellipse((p[0]-10, p[1]-10, p[0]+10, p[1]+10), fill="#D55E00", outline="#111111", width=2)
        draw.text((p[0]+dx, p[1]+dy), name, font=italic, fill="#111111")
    draw.arc((735, 420, 930, 625), 276, 327, fill="#7A3DF0", width=6)
    draw.text((800, 452), "Δψ", font=italic, fill="#7A3DF0")
    draw_dashed(draw, [(330, 700), (450, 700)], "#666666", 7, dash=24, gap=15)
    draw.text((475, 700), "đường thô", font=label, fill="#111111", anchor="lm")
    draw.line((770, 700, 890, 700), fill="#0072B2", width=10)
    draw.text((915, 700), "đường sau làm mượt", font=label, fill="#111111", anchor="lm")
    canvas.save(path, dpi=(300, 300))


def make_drive_figure_v2(path: Path):
    canvas = Image.new("RGB", (1400, 1010), "white")
    draw = ImageDraw.Draw(canvas)
    label = ImageFont.truetype(FONT_REGULAR, 36)
    italic = ImageFont.truetype(FONT_ITALIC, 40)
    draw.rounded_rectangle((485, 260, 915, 650), radius=45, fill="#D9EAF4", outline="#222222", width=6)
    draw.rectangle((435, 325, 490, 485), fill="#444444")
    draw.rectangle((910, 325, 965, 485), fill="#444444")
    # Wheel separation dimension above the robot.
    draw.line((462, 215, 938, 215), fill="#333333", width=5)
    draw.line((462, 195, 462, 235), fill="#333333", width=5)
    draw.line((938, 195, 938, 235), fill="#333333", width=5)
    draw_arrow(draw, (700, 465), (700, 78), "#0072B2", 9, 30)
    draw.text((738, 92), "v", font=italic, fill="#0072B2")
    draw_arrow(draw, (462, 405), (462, 85), "#D55E00", 8, 28)
    draw.text((420, 92), "v_L", font=italic, fill="#D55E00", anchor="ra")
    draw_arrow(draw, (938, 405), (938, 85), "#009E73", 8, 28)
    draw.text((978, 92), "v_R", font=italic, fill="#009E73", anchor="la")
    draw.arc((600, 365, 850, 620), 285, 85, fill="#7A3DF0", width=8)
    draw.text((855, 500), "ω", font=italic, fill="#7A3DF0")
    # Definitions are placed below the drawing, never on top of the robot.
    draw.line((120, 725, 200, 725), fill="#333333", width=5)
    draw.text((225, 725), "b: khoảng cách tâm hai bánh", font=label, fill="#111111", anchor="lm")
    draw.line((120, 790, 200, 790), fill="#D55E00", width=9)
    draw.text((225, 790), "v_L: vận tốc bánh trái", font=label, fill="#111111", anchor="lm")
    draw.line((755, 790, 835, 790), fill="#009E73", width=9)
    draw.text((860, 790), "v_R: vận tốc bánh phải", font=label, fill="#111111", anchor="lm")
    draw.text((700, 910), "v_L = v(1 − bκ/2)        v_R = v(1 + bκ/2)", font=label, fill="#111111", anchor="ma")
    canvas.save(path, dpi=(300, 300))


def make_swept_figure_v2(path: Path):
    canvas = Image.new("RGBA", (1400, 1020), "white")
    draw = ImageDraw.Draw(canvas, "RGBA")
    label = ImageFont.truetype(FONT_REGULAR, 32)
    italic = ImageFont.truetype(FONT_ITALIC, 31)
    curve = []
    for j in range(301):
        t = j/300
        curve.append((150+1000*t, 650-440*(3*t*t-2*t*t*t)))
    draw.line(curve, fill="#0072B2", width=10)

    def footprint(center, theta, alpha):
        length, width = 190, 135
        pts=[]
        for dx,dy in [(-length/2,-width/2),(length/2,-width/2),(length/2,width/2),(-length/2,width/2)]:
            pts.append((center[0]+dx*math.cos(theta)-dy*math.sin(theta),
                        center[1]+dx*math.sin(theta)+dy*math.cos(theta)))
        draw.polygon(pts, fill=(0,114,178,alpha), outline=(0,72,115,230))

    for t in (0.05,0.22,0.39,0.56,0.73,0.90):
        center=(150+1000*t,650-440*(3*t*t-2*t*t*t))
        theta=math.atan2(-440*(6*t-6*t*t),1000)
        footprint(center,theta,55)
    draw.rectangle((1240, 80, 1370, 330), fill="#4D004B", outline="#111111", width=5)
    # Separate legend band: no descriptive text is drawn on the geometry.
    y1, y2, y3 = 790, 865, 940
    draw.rectangle((115, y1-24, 195, y1+24), fill=(0,114,178,70), outline=(0,72,115,230), width=3)
    draw.text((225, y1), "Các tư thế hình bao robot được lấy mẫu", font=label, fill="#111111", anchor="lm")
    draw.line((115, y2, 235, y2), fill="#0072B2", width=10)
    draw.text((260, y2), "Đường tâm robot r(s)", font=label, fill="#111111", anchor="lm")
    draw.rectangle((115, y3-24, 195, y3+24), fill="#4D004B", outline="#111111", width=3)
    draw.text((225, y3), "Ô vật cản", font=label, fill="#111111", anchor="lm")
    canvas.convert("RGB").save(path, dpi=(300, 300))


def make_shared_edge_figure_v2(path: Path):
    canvas = Image.new("RGB", (1500, 760), "white")
    draw = ImageDraw.Draw(canvas)
    label = ImageFont.truetype(FONT_REGULAR, 32)
    italic = ImageFont.truetype(FONT_ITALIC, 33)
    left_v, right_v = (250, 430), (1250, 430)
    draw.line([(90, 640), left_v, right_v, (1410, 170)], fill="#666666", width=8, joint="curve")
    # Two distinct transition portions terminating on the shared edge.
    draw.arc((150, 320, 570, 650), 105, 183, fill="#0072B2", width=12)
    draw.arc((930, 250, 1350, 585), 274, 342, fill="#0072B2", width=12)
    for p,name in [(left_v,"Vᵢ"),(right_v,"Vᵢ₊₁")]:
        draw.ellipse((p[0]-11,p[1]-11,p[0]+11,p[1]+11),fill="#111111")
        draw.text((p[0],p[1]-40),name,font=italic,fill="#111111",anchor="ms")
    left_end, right_start = 560, 930
    y1, y2 = 570, 665
    # Full shared-edge length.
    draw.line((left_v[0],y2,right_v[0],y2),fill="#333333",width=4)
    for x in (left_v[0],right_v[0]): draw.line((x,y2-14,x,y2+14),fill="#333333",width=4)
    draw.text(((left_v[0]+right_v[0])/2,y2+22),"Lᵢ",font=italic,fill="#111111",anchor="mt")
    # Occupied portions and positive margin.
    draw.line((left_v[0],y1,left_end,y1),fill="#D55E00",width=9)
    draw.line((left_end,y1,right_start,y1),fill="#7A3DF0",width=9)
    draw.line((right_start,y1,right_v[0],y1),fill="#009E73",width=9)
    for x in (left_v[0],left_end,right_start,right_v[0]): draw.line((x,y1-13,x,y1+13),fill="#333333",width=3)
    draw.text(((left_v[0]+left_end)/2,y1-24),"dᵢ",font=italic,fill="#D55E00",anchor="ms")
    draw.text(((left_end+right_start)/2,y1-24),"m",font=italic,fill="#7A3DF0",anchor="ms")
    draw.text(((right_start+right_v[0])/2,y1-24),"dᵢ₊₁",font=italic,fill="#009E73",anchor="ms")
    canvas.save(path,dpi=(300,300))


def make_environment_roi_figure(path: Path, env: str, scenario: str):
    pgm = Image.open(MAP_DIR / f"{env}.pgm").convert("L")
    meta = yaml.safe_load((MAP_DIR / f"{env}.yaml").read_text(encoding="utf-8"))
    res = float(meta["resolution"])
    xmin, ymin = map(float, meta["origin"][:2])
    xmax, ymax = xmin + pgm.width * res, ymin + pgm.height * res
    data = json.loads((RVIZ_DIR / f"{env}__{scenario}__ThetaStar.json").read_text(encoding="utf-8"))
    raw = data["paths"]["raw"]["poses"]
    smooth = data["paths"]["pstmo"]["poses"]

    canvas = Image.new("RGB", (1200, 1000), "white")
    draw = ImageDraw.Draw(canvas)
    label = ImageFont.truetype(FONT_REGULAR, 32)
    small = ImageFont.truetype(FONT_REGULAR, 30)
    left, top, map_w, map_h = 150, 45, 900, 600
    scale_x, scale_y = map_w / (xmax - xmin), map_h / (ymax - ymin)
    # The map is rendered with equal metric scales, so occupied cells remain square.
    map_image = pgm.convert("RGB").resize((map_w, map_h), Image.Resampling.NEAREST)
    canvas.paste(map_image, (left, top))

    def cv(p):
        return (left + (p["x"] - xmin) * scale_x,
                top + map_h - (p["y"] - ymin) * scale_y)

    # One-metre grid, thin enough to preserve the occupancy map underneath.
    gx = math.ceil(xmin)
    while gx <= math.floor(xmax):
        x = left + (gx - xmin) * scale_x
        draw.line((x, top, x, top + map_h), fill="#D8D8D8", width=1)
        gx += 1
    gy = math.ceil(ymin)
    while gy <= math.floor(ymax):
        y = top + map_h - (gy - ymin) * scale_y
        draw.line((left, y, left + map_w, y), fill="#D8D8D8", width=1)
        gy += 1

    # Sampled 0.44 m x 0.34 m robot footprints show motion and orientation.
    overlay = Image.new("RGBA", canvas.size, (255, 255, 255, 0))
    odraw = ImageDraw.Draw(overlay, "RGBA")
    sample_count = 5 if len(smooth) < 300 else 7
    sampled_indices = sorted({round(i * (len(smooth) - 1) / (sample_count - 1)) for i in range(sample_count)})
    for idx in sampled_indices:
        pose = smooth[idx]
        cx, cy = cv(pose)
        yaw = 2.0 * math.atan2(float(pose.get("qz", 0.0)), float(pose.get("qw", 1.0)))
        corners = []
        for dx, dy in ((-0.22, -0.17), (0.22, -0.17), (0.22, 0.17), (-0.22, 0.17)):
            wx = pose["x"] + dx * math.cos(yaw) - dy * math.sin(yaw)
            wy = pose["y"] + dx * math.sin(yaw) + dy * math.cos(yaw)
            corners.append((left + (wx - xmin) * scale_x,
                            top + map_h - (wy - ymin) * scale_y))
        odraw.polygon(corners, fill=(86, 180, 233, 72), outline=(0, 95, 145, 230))
        hx = cx + 0.27 * scale_x * math.cos(yaw)
        hy = cy - 0.27 * scale_y * math.sin(yaw)
        odraw.line((cx, cy, hx, hy), fill=(0, 95, 145, 230), width=4)
    canvas = Image.alpha_composite(canvas.convert("RGBA"), overlay).convert("RGB")
    draw = ImageDraw.Draw(canvas)

    draw.line([cv(p) for p in smooth], fill="#0072B2", width=7, joint="curve")
    draw_dashed(draw, [cv(p) for p in raw], "#D55E00", 6, dash=17, gap=10)

    start_xy = cv({"x": data["start"][0], "y": data["start"][1]})
    goal_xy = cv({"x": data["goal"][0], "y": data["goal"][1]})
    draw.ellipse((start_xy[0]-10, start_xy[1]-10, start_xy[0]+10, start_xy[1]+10),
                 fill="#009E73", outline="#111111", width=2)
    gx, gy = goal_xy
    draw.polygon(((gx, gy-13), (gx+13, gy), (gx, gy+13), (gx-13, gy)),
                 fill="#CC79A7", outline="#111111")
    draw.rectangle((left, top, left + map_w, top + map_h), outline="#222222", width=4)

    for value in (-6, -3, 0, 3, 6):
        x = left + (value - xmin) * scale_x
        draw.line((x, top + map_h, x, top + map_h + 10), fill="#222222", width=3)
        draw.text((x, top + map_h + 14), str(value), font=small, fill="#111111", anchor="ma")
    for value in (-4, -2, 0, 2, 4):
        y = top + map_h - (value - ymin) * scale_y
        draw.line((left - 10, y, left, y), fill="#222222", width=3)
        draw.text((left - 17, y), str(value), font=small, fill="#111111", anchor="rm")
    draw.text((left + map_w / 2, top + map_h + 52), "Tọa độ x (m)", font=label, fill="#111111", anchor="ma")
    y_label = Image.new("RGBA", (260, 50), (255, 255, 255, 0))
    ImageDraw.Draw(y_label).text((130, 25), "Tọa độ y (m)", font=label, fill="#111111", anchor="mm")
    y_label = y_label.rotate(90, expand=True)
    canvas.paste(y_label, (33, top + (map_h - y_label.height) // 2), y_label)
    draw = ImageDraw.Draw(canvas)

    # One item per row; no two text labels share the same horizontal band.
    y1, y2, y3, y4, y5 = 750, 805, 860, 915, 970
    draw_dashed(draw, [(115, y1), (235, y1)], "#D55E00", 6, dash=17, gap=10)
    draw.text((260, y1), "Đường thô từ bộ lập kế hoạch toàn cục", font=small, fill="#111111", anchor="lm")
    draw.line((115, y2, 235, y2), fill="#0072B2", width=7)
    draw.text((260, y2), "Đường sau làm mượt bằng phương pháp đề xuất", font=small, fill="#111111", anchor="lm")
    draw.rectangle((115, y3-12, 150, y3+12), fill="#B8DBEF", outline="#005F91", width=2)
    draw.text((165, y3), "Các tư thế hình bao robot", font=small, fill="#111111", anchor="lm")
    draw.ellipse((115, y4-10, 135, y4+10), fill="#009E73", outline="#111111", width=2)
    draw.text((165, y4), "Điểm bắt đầu", font=small, fill="#111111", anchor="lm")
    draw.polygon(((125, y5-12), (137, y5), (125, y5+12), (113, y5)), fill="#CC79A7", outline="#111111")
    draw.text((165, y5), "Điểm đích", font=small, fill="#111111", anchor="lm")
    canvas.save(path, dpi=(300, 300))


def make_metric_figure_v2(path: Path, metrics, key: str):
    canvas=Image.new("RGB",(1200,1100),"white"); draw=ImageDraw.Draw(canvas)
    title_font=ImageFont.truetype(FONT_BOLD,32)
    label=ImageFont.truetype(FONT_REGULAR,30)
    tick_font=ImageFont.truetype(FONT_REGULAR,26)
    legend_font=ImageFont.truetype(FONT_REGULAR,30)
    methods=["raw","simple","savitzky_golay","constrained","pstmo"]
    names={
        "raw":"Đường thô từ bộ lập kế hoạch toàn cục",
        "simple":"Navigation2 Simple Smoother",
        "savitzky_golay":"Navigation2 Savitzky–Golay Smoother",
        "constrained":"Navigation2 Constrained Smoother",
        "pstmo":"Path Smoothing and Turning-Maneuver Optimization (PSTMO)",
    }
    colors={"raw":"#777777","simple":"#009E73","savitzky_golay":"#E69F00","constrained":"#CC79A7","pstmo":"#0072B2"}
    envs=["open_arena","narrow_aisles","warehouse_cross_aisles"]
    enames=[("Không gian","mở"),("Lối đi","hẹp"),("Kho có lối","giao cắt")]
    left,top,right,bottom=115,95,1165,590
    draw.line((left,top,left,bottom),fill="#222222",width=4); draw.line((left,bottom,right,bottom),fill="#222222",width=4)
    if key=="E":
        ticks=[1,2,5,10,20,50,100,200,400]; ymin,ymax=1,400
        ypos=lambda v: bottom-(math.log10(max(v,ymin))-math.log10(ymin))/(math.log10(ymax)-math.log10(ymin))*(bottom-top)
        chart_title="Tích phân bình phương độ cong trung bình Eκ (m⁻¹)"
    else:
        ticks=[0,50,100,150,200]; ymin,ymax=0,200
        ypos=lambda v: bottom-(v-ymin)/(ymax-ymin)*(bottom-top)
        chart_title="Thời gian xử lý thuật toán trung bình T (ms)"
    draw.text(((left+right)/2,42),chart_title,font=title_font,fill="#111111",anchor="mm")
    for tick in ticks:
        y=ypos(tick); draw.line((left,y,right,y),fill="#DDDDDD",width=2); draw.text((left-18,y),str(tick),font=tick_font,fill="#111111",anchor="rm")
    gw=(right-left)/3; bw=gw*0.14
    for ei,(env,ename) in enumerate(zip(envs,enames)):
        center=left+gw*(ei+0.5)
        for mi,m in enumerate(methods):
            val=metrics[env][m][key]; x0=center+(mi-2.5)*bw; x1=x0+bw*0.82
            draw.rectangle((x0,ypos(val),x1,bottom),fill=colors[m],outline="#111111",width=2)
        draw.text((center,bottom+20),ename[0],font=label,fill="#111111",anchor="ma")
        draw.text((center,bottom+53),ename[1],font=label,fill="#111111",anchor="ma")
    draw.rectangle((left,top,right,bottom),outline="#222222",width=4)
    # One full method name per line, outside the data rectangle.
    lx, ly = 125, 735
    for row, m in enumerate(methods):
        y = ly + row * 66
        draw.rectangle((lx,y-14,lx+32,y+14),fill=colors[m],outline="#111111",width=2)
        draw.text((lx+48,y),names[m],font=legend_font,fill="#111111",anchor="lm")
    canvas.save(path,dpi=(300,300))


def clear_paragraph_content(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def build_document(template_path: Path, problem_path: Path, drive_path: Path,
                   swept_path: Path, bezier_path: Path, edge_path: Path,
                   environment_paths: dict[str, Path], results_e_path: Path,
                   results_t_path: Path, metrics):
    doc = Document(str(template_path))
    ensure_custom_styles(doc)
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in footer.paragraphs:
                clear_paragraph_content(paragraph)

    # Preserve the exact floating author layout from abstract final 2.docx.
    title = next(p for p in doc.paragraphs if p.text.startswith("Path Smoothing via"))
    clear_paragraph_content(title)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("Làm mượt đường đi bằng tối ưu chuyển tiếp tại góc có xét hình bao cho robot di động vi sai")
    set_run_font(run, 24)

    abstract_text = (
        "Đường do bộ lập kế hoạch toàn cục sinh ra thường là chuỗi điểm gấp khúc; dù không va chạm trên bản đồ chi phí, "
        "đường đó chưa bảo đảm chuyển hướng liên tục và khả thi đối với robot hai bánh vi sai. Bài báo đề xuất PSTMO, một "
        "bộ hậu xử lý thay lân cận góc bằng đoạn Bézier bậc năm liên tục hình học bậc hai (G²). Các ứng viên được sàng lọc theo "
        "độ cong, giới hạn vận tốc thân và bánh, vùng quét hình bao robot; sau đó quy hoạch động chọn tổ hợp không chồng lấn "
        "trên toàn đường. Phương pháp được đánh giá trong ROS 2 Navigation2 trên ba bản đồ—không gian mở, lối đi hẹp và kho "
        "có lối giao cắt—với năm bộ lập kế hoạch toàn cục, tạo 15 nhóm đầu vào ghép cặp. So với Simple và Constrained, "
        "PSTMO giảm trung bình tích phân bình phương độ cong lần lượt 75,42% và 90,75%, đồng thời giảm chiều dài 0,98% và "
        "2,01%. Giá trị trung bình của PSTMO là 8,645 m, 2,887 m⁻¹ và 96,0 ms cho chiều dài, mức uốn tích lũy và thời gian "
        "xử lý. Kết quả cho thấy khả năng cải thiện hình học rõ rệt, đổi lại chi phí tính toán cao hơn các bộ làm mượt đối chứng."
    )
    abstract = next(p for p in doc.paragraphs if p.text.startswith("Abstract—"))
    clear_paragraph_content(abstract)
    abstract.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abstract.paragraph_format.first_line_indent = Inches(0)
    abstract.paragraph_format.space_after = Pt(3)
    r = abstract.add_run("Tóm tắt—")
    set_run_font(r, 9, bold=True, italic=True)
    r = abstract.add_run(abstract_text)
    set_run_font(r, 9, bold=True)

    keywords = next(p for p in doc.paragraphs if p.text.startswith("Keywords—"))
    clear_paragraph_content(keywords)
    keywords.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    keywords.paragraph_format.first_line_indent = Inches(0)
    keywords.paragraph_format.space_after = Pt(4)
    r = keywords.add_run("Từ khóa—")
    set_run_font(r, 9, bold=True, italic=True)
    r = keywords.add_run("robot vi sai, làm mượt đường đi, Bézier bậc năm, hình bao robot, Navigation2, PSTMO.")
    set_run_font(r, 9, italic=True)

    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    configure_section(body_section, columns=2)
    # The IEEE A4 template starts its two-column flow at Abstract, not after
    # Keywords.  Move the continuous section break created above immediately
    # before the abstract while leaving the title/author area unchanged.
    section_break_paragraph = doc.paragraphs[-1]
    abstract._p.addprevious(section_break_paragraph._p)

    add_heading1(doc, "I. GIỚI THIỆU")
    add_body(doc, (
        "Robot tự hành dẫn đường và robot di động tự hành ngày càng được dùng trong kho vận, sản xuất và vận chuyển nội bộ. "
        "Trong một ngăn xếp điều hướng mô-đun như ROS 2 Navigation2 (Nav2), bộ lập kế hoạch toàn cục tìm đường tránh vật cản, "
        "còn bộ điều khiển chuyển đường đó thành lệnh vận tốc [1]. Vì vậy, chất lượng hình học của đường ảnh hưởng trực tiếp "
        "đến khả năng bám đường, mức giảm tốc khi rẽ và tính ổn định của chuyển động."), first=True)
    add_body(doc, (
        "Các bộ tìm kiếm trên lưới thường trả về một đa tuyến. Tại một đỉnh gãy, hướng tiếp tuyến đổi đột ngột; nếu robot đi "
        "qua với vận tốc khác không thì vận tốc góc hoặc vận tốc bánh yêu cầu có thể tăng mạnh. Bộ điều khiển khi đó phải giảm "
        "tốc, lệch khỏi đường, hoặc dừng để quay tại chỗ. Làm mượt sau lập kế hoạch đã được nghiên cứu bằng dải đàn hồi [2], "
        "các đường cong hình học [3], [4] và tối ưu số. Tuy nhiên, mỗi nhóm phải đánh đổi giữa tính liên tục, chi phí tính toán, "
        "khả năng xét động học và mô hình va chạm của toàn thân robot."))
    add_figure(doc, problem_path, 3.20,
               "Đỉnh gãy của đường thô và chuyển tiếp cục bộ giữ nguyên tuyến nhưng phân bố biến thiên hướng trên một đoạn hữu hạn.", 1)
    add_body(doc, (
        "Bài báo đề xuất Path Smoothing and Turning-Maneuver Optimization (PSTMO), đặt giữa bộ lập kế hoạch toàn cục và bộ "
        "điều khiển Nav2. PSTMO chỉ thay cục bộ vùng quanh góc nhưng quyết định trên toàn đường. Ba đóng góp chính là: (1) cấu "
        "trúc Bézier bậc năm tạo mối nối thẳng–cong có tiếp tuyến và độ cong liên tục; (2) sàng lọc cứng theo động học robot "
        "vi sai và vùng quét hình bao, kết hợp quy hoạch động để ngăn các đoạn chuyển tiếp ở góc kề nhau chồng lấn; và (3) đánh "
        "giá ghép cặp trên 15 tổ hợp bản đồ–bộ lập kế hoạch bằng chiều dài, tích phân bình phương độ cong và thời gian xử lý."))
    add_body(doc, (
        "Thử nghiệm mô phỏng cho thấy PSTMO đạt mức uốn tích lũy thấp nhất trong cả ba môi trường được chọn, trong khi chiều "
        "dài không tăng và thời gian xử lý vẫn ở thang dưới một phần tư giây cho từng ca. Phần II tổng quan nghiên cứu liên quan; "
        "Phần III xây dựng mô hình và bài toán; Phần IV trình bày PSTMO; Phần V mô tả thử nghiệm và thảo luận kết quả; cuối cùng, "
        "Phần VI kết luận bài báo."))

    add_heading1(doc, "II. CÁC NGHIÊN CỨU LIÊN QUAN")
    add_heading2(doc, "A. Các đường cong và phần tử hình học")
    add_body(doc, (
        "Ý tưởng gốc của nhóm này là thay đỉnh đa tuyến bằng một phần tử có quy luật độ cong biết trước. Fleury và cộng sự dùng "
        "clothoid cùng anticlothoid để nối các đoạn chuyển động và duy trì vận tốc bánh liên tục, kể cả khi đổi chiều [3]. Brezak "
        "và Petrović phát triển cách làm mượt trực tuyến bằng clothoid cho robot vi sai, cho phép độ cong ban đầu khác không [4]. "
        "Clothoid có ưu điểm độ cong biến thiên tuyến tính theo chiều dài cung, nhưng liên quan tích phân Fresnel và thường cần "
        "xấp xỉ, bảng tra hoặc giải số; thông tin khoảng hở vật cản cũng phải được cung cấp cho đoạn cần thay."), first=True)
    add_body(doc, (
        "Bézier cung cấp biểu thức đóng và tính chất bao lồi thuận lợi. Yang và Sukkarieh xây dựng thuật toán giải tích bằng Bézier "
        "bậc ba với độ cong liên tục và giới hạn độ cong [5]; Bu và cộng sự phát triển các lượt rẽ Bézier bậc ba liên tục độ cong "
        "cho xe dạng ô tô [6]. Xu và cộng sự tối ưu một đường Bézier liên tục bằng bầy đàn để giảm chiều dài, điểm không khả thi "
        "và độ cong lớn nhất [7]. Parque và Miyashita dùng Differential Evolution để khớp Bézier với quỹ đạo robot quan sát [8]. "
        "Các công trình này thể hiện tốt tính trơn hình học, nhưng nhiều đánh giá dùng robot điểm hoặc môi trường lưới đơn giản; "
        "khớp đường từ dữ liệu cũng không tự bảo đảm vùng quét hình bao không va chạm."))
    add_body(doc, (
        "Hai họ clothoid và Bézier vì thế có đánh đổi khác nhau. Clothoid đưa quy luật biến thiên độ cong vào chính định nghĩa "
        "đường nhưng phép tính trực tuyến phức tạp hơn; Bézier dễ lấy đạo hàm, dễ điều chỉnh bằng điểm điều khiển nhưng phải thiết "
        "kế cấu trúc điểm một cách có chủ ý mới bảo đảm độ cong ở mối nối. Với đường Nav2 gồm nhiều đoạn thẳng ngắn, một đường "
        "cong toàn cục còn có thể dịch xa hành lang ban đầu. PSTMO chọn các chuyển tiếp Bézier cục bộ để giữ lại phần lớn đường "
        "thô, đồng thời dùng điều kiện khả thi riêng để bù cho hạn chế của tính chất bao lồi."))

    add_heading2(doc, "B. Biến dạng đường và tối ưu quỹ đạo")
    add_body(doc, (
        "Elastic Bands là nền tảng của nhóm biến dạng: một đường không va chạm được biểu diễn bởi chuỗi vùng tự do chồng lấn và "
        "được co bởi lực nội tại, đồng thời bị vật cản đẩy ra [2]. Cách này thích hợp cập nhật cục bộ theo thời gian thực, nhưng là "
        "tối ưu cục bộ và không tự đổi lớp đồng luân của đường. Dolgov và cộng sự kết hợp Hybrid A* với trường Voronoi và tối ưu "
        "gradient để tạo đường khả thi cho xe tự hành [9]. CHOMP đưa ra gradient hiệp biến trên hàm chi phí trơn và vật cản [10]; "
        "hiệu quả phụ thuộc đường khởi tạo và có thể dừng ở cực tiểu cục bộ."), first=True)
    add_body(doc, (
        "Timed Elastic Band (TEB) bổ sung khoảng thời gian giữa các trạng thái để tối ưu đồng thời hình học, vận tốc, gia tốc và "
        "vật cản trong một bài toán thưa [11]. Các mở rộng sau đó xây dựng điều khiển dự đoán tối ưu thời gian [12] và tối ưu "
        "động học–động lực học cho xe dạng ô tô [13]. Nhóm này xử lý trực tiếp quỹ đạo theo thời gian và môi trường động, nhưng "
        "cần điều chỉnh nhiều trọng số, lời giải vẫn cục bộ và chi phí cao hơn một phép thay góc có cấu trúc."))
    add_body(doc, (
        "Điểm mạnh của tối ưu toàn đường là các đỉnh có thể dịch chuyển phối hợp, nhưng điều đó cũng làm tăng số biến và khiến "
        "lời giải nhạy với đường khởi tạo, trọng số vật cản và cách rời rạc hóa. Ngược lại, một phép biến đổi độc lập tại từng góc "
        "nhẹ hơn nhưng dễ sinh hai cung chồng lên cùng một cạnh. Do đó, bài toán của nghiên cứu này không phải thay thế TEB hoặc "
        "CHOMP, mà là tìm mức trung gian: ứng viên có cấu trúc được kiểm tra cục bộ, còn tính tương thích được quyết định ở cấp "
        "toàn đường."))

    add_heading2(doc, "C. Phương pháp lai và khoảng trống nghiên cứu")
    add_body(doc, (
        "GRIPS kết hợp biến dạng gradient theo trường khoảng cách với phép nối tắt có kiểm tra, qua đó giảm chiều dài và độ cong "
        "trên nhiều bộ lập kế hoạch và hàm lái [14]. Andreasson và cộng sự tạo các miền ràng buộc lồi không va chạm quanh đường "
        "lattice rồi tối ưu liên tục, đạt sai số tư thế cuối nhỏ trên AGV công nghiệp [15]. Đổi lại, các cách này cần trường khoảng "
        "cách hoặc miền ràng buộc được chuẩn bị, và kiến trúc thường dành cho xe dạng ô tô hay toàn bộ trạng thái liên tục."), first=True)
    add_body(doc, (
        "Nhìn chung, tài liệu đã giải quyết tốt từng khía cạnh: liên tục độ cong, tối ưu theo vật cản, hoặc ràng buộc theo thời gian. "
        "Khoảng trống còn lại trong bối cảnh Nav2 của robot vi sai là phối hợp đồng thời mối nối thẳng–cong có độ cong đầu cuối "
        "bằng không, giới hạn hai bánh, kiểm tra vùng quét hình bao và tương tác giữa các góc kề nhau, nhưng vẫn giữ cấu trúc hậu "
        "xử lý nhẹ và không thay tuyến đường toàn cục. PSTMO được xây dựng cho khoảng trống này."))
    add_body(doc, (
        "Từ tổng quan trên, bốn yêu cầu được dùng để định hướng phương pháp: đoạn thay phải nối được với cạnh thẳng mà không tạo "
        "bước nhảy độ cong; giới hạn của hai bánh phải được suy ra trực tiếp từ hình học; va chạm phải được đánh giá trên toàn hình "
        "bao thay vì chỉ tại tâm; và lựa chọn ở một góc không được làm mất chiều dài cần thiết của góc kế tiếp. Việc đánh giá sau "
        "đó phải tách rõ chất lượng hình học khỏi chi phí xử lý, tránh diễn giải Eκ như năng lượng tiêu thụ."))

    add_heading1(doc, "III. MÔ HÌNH TOÁN HỌC VÀ PHÁT BIỂU BÀI TOÁN")
    add_heading2(doc, "A. Mô hình robot và biểu diễn đường")
    add_body(doc, (
        "Trạng thái robot là z=(x,y,θ), với (x,y) là vị trí tâm hình học và θ là góc hướng trong hệ quy chiếu bản đồ. Với vận "
        "tốc tuyến tính v và vận tốc góc ω, mô hình động học hai bánh vi sai được viết như (1)."), first=True)
    add_equation(doc, "ẋ = v cos θ,   ẏ = v sin θ,   θ̇ = ω.", 1)
    add_body(doc, (
        "Đường thô từ bộ lập kế hoạch là P={p₀,…,pₙ}, với pᵢ=(rᵢ,ψᵢ) và rᵢ=(xᵢ,yᵢ). Chiều dài chỉ được tính trên "
        "tọa độ phẳng theo (2), không cộng góc hướng vào chuẩn khoảng cách. Một điểm pᵢ là góc khi "
        "hai vectơ đơn vị của cạnh vào u và cạnh ra v tạo góc có dấu φᵢ=atan2(u×v,u·v) lớn hơn ngưỡng phát hiện."))
    add_equation(doc, "L(P) = Σⁿ⁻¹ᵢ₌₀ ‖rᵢ₊₁ − rᵢ‖₂.", 2)

    add_heading2(doc, "B. Độ cong, tính liên tục và các ràng buộc")
    add_body(doc, (
        "Với đường tham số r(t)=(x(t),y(t)), độ cong có dấu được xác định bởi (3). Hai đại lượng mô tả hình học là độ cong "
        "lớn nhất κmax và mức uốn tích lũy Eκ trong (4). Eκ có đơn vị m⁻¹ và chỉ là chỉ số hình học, không phải năng lượng "
        "điện tiêu thụ."), first=True)
    add_equation(doc, "κ(t) = (x′y″ − y′x″)/(x′² + y′²)³ᐟ².", 3)
    add_equation(doc, "κmax = maxₛ |κ(s)|,   Eκ = ∫₀ᴸ κ²(s) ds.", 4)
    add_body(doc, (
        "Liên tục G¹ bảo đảm hướng tiếp tuyến trùng nhau, còn G² yêu cầu thêm độ cong liên tục. Vì đoạn thẳng có κ=0, đoạn "
        "chuyển tiếp nối với nó phải có κ bằng không tại hai đầu. Với khoảng cách hai bánh b, vận tốc hai bánh và quan hệ theo "
        "độ cong là (5); ứng viên phải thỏa giới hạn vận tốc thân, vận tốc góc, vận tốc từng bánh và gia tốc ngang."))
    add_equation(doc, "vL=v(1−bκ/2),   vR=v(1+bκ/2),   ω=vκ.", 5)
    add_figure(doc, drive_path, 3.05,
               "Mô hình động học robot hai bánh vi sai và quy ước vận tốc bánh.", 2)
    add_body(doc, (
        "Gọi Fbody là đa giác hình bao trong hệ thân xe. Tại tư thế z, hình bao trên bản đồ là Fmap(z)=[x,y]ᵀ+R(θ)Fbody. "
        "Hợp các hình bao tại những tư thế nội suy dọc đường tạo vùng quét; vùng này không được giao ô vật cản. Nếu hai góc "
        "liên tiếp cùng dùng cạnh dài L_i, tổng khoảng cắt của chúng còn phải chừa một biên dương để tránh chồng lấn."))
    add_figure(doc, swept_path, 3.05,
               "Vùng quét được tạo từ hợp các hình bao lấy mẫu dọc đường tâm; va chạm được kiểm tra trên toàn vùng này.", 3)

    add_heading2(doc, "C. Phát biểu bài toán")
    add_body(doc, (
        "Đầu vào gồm đường P, bản đồ chi phí, Fbody và các giới hạn chuyển động. Cần tìm đường sau làm mượt P̃ giữ nguyên p₀ và "
        "pₙ, không đổi thứ tự các hành lang của P, liên tục G² tại mối nối được thay, thỏa (5), không va chạm theo vùng quét và "
        "không chồng lấn giữa các góc. Trong miền khả thi đó, bài toán đa mục tiêu ưu tiên giảm Eκ, tránh tăng L không cần thiết "
        "và giữ thời gian xử lý phù hợp cho bước hậu xử lý trực tuyến. Nếu không có chuyển tiếp khả thi, phương án giữ nguyên hoặc "
        "quay tại chỗ được bảo toàn thay vì tạo một đường cong không an toàn."), first=True)
    add_equation(doc, "tìm P* ∈ F(P, M, Fbody, Llim) để giảm [Eκ(P*), L(P*), Talg].", 6)
    add_body(doc, (
        "Trong (6), M là costmap, Llim là tập giới hạn chuyển động và F là miền các đường thỏa đồng thời điều kiện đầu–đích, G², "
        "động học, vùng quét và không chồng lấn. Đây là phát biểu đa mục tiêu, không giả định ba đại lượng có cùng đơn vị hoặc "
        "một bộ trọng số phổ quát. Triển khai PSTMO trước hết loại mọi phương án ngoài F, rồi mới so sánh các ứng viên còn lại. "
        "Cách tách khả thi khỏi chất lượng ngăn một đường cong có Eκ nhỏ nhưng va chạm được bù bởi điểm số hình học tốt."))

    add_heading1(doc, "IV. PHƯƠNG PHÁP PSTMO ĐỀ XUẤT")
    add_heading2(doc, "A. Tổng quan phương pháp")
    add_body(doc, (
        "PSTMO nhận đường Nav2 và trả về đường cùng hệ quy chiếu. Trước hết, thuật toán loại nhiễu răng cưa nhỏ mà không vượt "
        "hành lang an toàn, rồi phát hiện các góc có ý nghĩa. Tại mỗi góc, một tập nhỏ phương án giữ nguyên, quay tại chỗ hoặc "
        "chuyển tiếp được tạo ra. Mỗi phương án được kiểm tra độc lập; sau đó một phép tối ưu theo chuỗi chọn các phương án tương "
        "thích giữa các góc. Cuối cùng, các đoạn thẳng còn lại được ghép với chuyển tiếp đã chọn và toàn đường được hậu kiểm."), first=True)
    add_body(doc, (
        "Ba trạng thái xử lý được giữ trong cùng một biểu diễn. Góc rất nhỏ có thể đi qua mà không chèn cung; góc có đủ không "
        "gian nhận một chuyển tiếp; còn quay tại chỗ là phương án dự phòng khi chuyển tiếp tịnh tiến không khả thi. Sự hiện diện "
        "của phương án dự phòng giúp thuật toán không buộc mọi góc phải được làm tròn, một giả định có thể nguy hiểm trong lối "
        "hẹp. Quy trình trên vẫn trả về nav_msgs/Path; thông tin thời gian chỉ phục vụ kiểm tra và lựa chọn, không biến đầu ra "
        "thành một quỹ đạo đã đóng dấu thời gian."))

    add_heading2(doc, "B. Xây dựng đoạn chuyển tiếp và lựa chọn tham số")
    add_body(doc, (
        "Xét đỉnh V, vectơ đơn vị u trên cạnh vào và v trên cạnh ra. Hai điểm tiếp giáp A=V−du và B=V+dv được xác định bởi "
        "khoảng cắt d. Đặt q=αd, sáu điểm điều khiển được sắp như (7) và đường Bézier bậc năm được cho bởi (8)."), first=True)
    add_equation(doc, ["P₀=A, P₁=A+qu, P₂=A+2qu;", "P₃=B−2qv, P₄=B−qv, P₅=B."], 7)
    add_equation(doc, "B(t)=Σ⁵ᵢ₌₀ C(5,i)(1−t)⁵⁻ⁱtⁱPᵢ,   0≤t≤1.", 8)
    add_figure(doc, bezier_path, 3.12, "Cấu trúc sáu điểm điều khiển. Ba điểm đầu và ba điểm cuối thẳng hàng, cách đều nên độ cong tại hai đầu bằng không.", 4)
    add_body(doc, (
        "Cấu trúc này cho B′ cùng hướng với hai cạnh và B″ bằng không tại hai đầu; vì vậy đoạn chuyển tiếp nối với đoạn thẳng "
        "theo G². Tăng d mở rộng vùng chuyển hướng, còn α điều chỉnh hình dạng và phân bố độ cong. PSTMO biến thiên hai tham số "
        "để tạo một số ứng viên đại diện; chi tiết lưới tìm kiếm và ngưỡng dừng không được trình bày trong khuôn khổ bài báo này."))
    add_body(doc, (
        "Tác động của hai tham số có thể tách về mặt hình học. d quyết định phần chiều dài cạnh được dành cho thao tác rẽ và do "
        "đó bị chặn bởi vật cản cũng như góc lân cận. Với cùng d, α nhỏ đặt các điểm điều khiển gần đầu mút hơn và thay đổi phân "
        "bố độ cong; α lớn làm tay đòn tiếp tuyến dài hơn. PSTMO không giả định một α phù hợp cho mọi góc mà chỉ giữ các cặp "
        "(d,α) vượt qua toàn bộ phép kiểm tra bắt buộc."))

    add_heading2(doc, "C. Đánh giá và lựa chọn các đoạn chuyển tiếp")
    add_body(doc, (
        "Ứng viên trước tiên phải có đạo hàm không suy biến và độ cong hữu hạn. Từ κ, hệ thống suy ra vận tốc góc và vận tốc "
        "hai bánh theo (5), đồng thời áp các giới hạn thân xe, bánh xe và gia tốc ngang. Tiếp theo, đa giác hình bao được lấy mẫu "
        "theo vị trí và góc hướng dọc chuyển tiếp; bất kỳ giao cắt nào với ô cấm đều làm ứng viên bị loại. Việc dùng toàn bộ hình "
        "bao tránh trường hợp tâm robot còn ở vùng tự do nhưng góc thân xe quét vào kệ."), first=True)
    add_body(doc, (
        "Các góc không độc lập vì hai chuyển tiếp kề nhau có thể cùng chiếm một cạnh. Với biên m và chiều dài cạnh chung L_i, "
        "hai lựa chọn liên tiếp chỉ tương thích khi thỏa (9). PSTMO xem mỗi phương án tại góc như một trạng thái và dùng quy "
        "hoạch động để chọn chuỗi khả thi có chi phí tích lũy thấp. Nếu N là số góc và mỗi góc có không quá K phương án, bước "
        "này có độ phức tạp O(NK²)."))
    add_equation(doc, "d_i + d_(i+1) + m ≤ L_i.", 9)
    add_figure(doc, edge_path, 3.12,
               "Ngân sách chiều dài trên cạnh chung của hai góc kề nhau; khoảng m ngăn hai chuyển tiếp chồng lấn.", 5)
    add_body(doc, (
        "Chi phí của trạng thái chỉ được so sánh sau khi ứng viên đã khả thi. Quy hoạch động lưu chi phí tốt nhất kết thúc ở từng "
        "phương án của góc hiện tại và con trỏ về phương án tương thích trước đó. Nhờ truy vết từ góc cuối, thuật toán thu được "
        "một chuỗi thống nhất thay vì chọn tham lam ứng viên tốt nhất riêng lẻ rồi sửa xung đột về sau."))

    add_heading2(doc, "D. Tạo và kiểm tra đường đầu ra")
    add_body(doc, (
        "Đường đầu ra được ghép theo thứ tự từ đoạn thẳng còn lại và các chuyển tiếp đã chọn, sau đó lấy mẫu lại với khoảng cách "
        "đều. Tư thế đầu và đích được giữ nguyên. Một hậu kiểm độc lập xác nhận tính hữu hạn, thứ tự điểm, giới hạn chuyển động, "
        "vùng quét hình bao và mối nối; nếu bất biến cuối không thỏa, kết quả không được công bố cho bộ điều khiển."), first=True)
    add_body(doc, (
        "Hậu kiểm được thực hiện trên đường đã ghép chứ không chỉ trên từng ứng viên, vì phép lấy mẫu lại và mối nối giữa các "
        "thành phần có thể tạo tình huống không xuất hiện khi kiểm tra cục bộ. Đường chỉ được chấp nhận khi số liệu hữu hạn, thứ "
        "tự tiến dọc tuyến được giữ, tư thế đầu–đích không đổi và mọi mẫu nội suy của hình bao đều hợp lệ trên cùng costmap đầu "
        "vào. Cơ chế này ưu tiên giữ một đường khả thi hơn việc luôn trả về một đường đã làm tròn."))

    add_heading1(doc, "V. THỬ NGHIỆM VÀ ĐÁNH GIÁ KẾT QUẢ")
    add_heading2(doc, "A. Thiết lập thử nghiệm")
    add_body(doc, (
        "Thử nghiệm dùng Ubuntu 24.04, ROS 2 Jazzy, Nav2 và Gazebo Harmonic trên máy Intel Core i5-12450HX. Robot có hình bao "
        "chữ nhật 0,44×0,34 m, vệt bánh vật lý b=0,2548 m và bản đồ độ phân giải 0,05 m. Mô hình động học và bộ làm mượt dùng "
        "giá trị vật lý này; plugin DiffDrive của Gazebo dùng vệt bánh hiệu dụng 0,2834 m sau hiệu chuẩn tiếp xúc–odometry. "
        "Giới hạn chung là v≤0,30 m/s, "
        "|ω|≤0,80 rad/s, |vL,R|≤0,36 m/s và gia tốc ngang không quá 0,18 m/s². Khoảng mẫu kiểm tra Bézier là 0,02 m; "
        "khoảng điểm đầu ra và biên m đều là 0,05 m."), first=True)
    add_table_caption(doc, "I", "Các phương án được dùng trong phép so sánh")
    baseline_rows = [
        ("Raw", "Đường nguyên bản của bộ lập kế hoạch; không làm mượt.", "Mốc đầu vào"),
        ("Simple", "Lặp hiệu chỉnh điểm theo sai lệch dữ liệu và sai phân bậc hai; tối đa 1000 vòng, hai lần tinh chỉnh.", "Đối chứng nhẹ"),
        ("Savitzky–Golay", "Lọc đa thức cửa sổ 7 điểm với hệ số [−2, 3, 6, 7, 6, 3, −2]/21; hai lần tinh chỉnh.", "Đối chứng lọc cục bộ"),
        ("Constrained", "Tối ưu lặp có thành phần độ trơn và costmap; wsmooth=200000, wcost=0,015, wcurve=wdist=0.", "Đối chứng có ràng buộc"),
        ("PSTMO", "Chuyển tiếp Bézier G², giới hạn hai bánh, vùng quét hình bao và quy hoạch động.", "Phương pháp đề xuất"),
    ]
    add_ieee_table(doc, ["Phương án", "Nguyên lý/cấu hình dùng trong thử nghiệm", "Vai trò"], baseline_rows,
                   widths=[0.78, 1.70, 0.77], font_size=6.6)
    add_body(doc, (
        "Bảng I làm rõ rằng Raw không phải một bộ làm mượt. Simple là phép khuếch tán hình học lặp; Savitzky–Golay là bộ lọc "
        "cục bộ cố định; còn Constrained giải một bài toán tối ưu. Trong cấu hình thực nghiệm, trọng số độ cong của Constrained "
        "bằng không, vì vậy phương pháp này không trực tiếp tối thiểu hóa Eκ. Các cấu hình trên được giữ nguyên cho cả ba bản đồ."), first=True)
    add_body(doc, (
        "Năm nguồn đường gồm NavFn A*, NavFn Dijkstra, ThetaStar, Smac2D và SmacHybrid. Trên từng đường thô, bốn bộ làm "
        "mượt Simple, Savitzky–Golay, Constrained và PSTMO được chạy độc lập; Raw là mốc không làm mượt. Các phương án dùng "
        "cùng dữ liệu costmap và cùng giới hạn. Mỗi tổ hợp được chạy một lần, nên các giá trị sau là thống kê mô tả trên nhóm "
        "đầu vào ghép cặp, không phải ước lượng độ bất định."))
    add_table_caption(doc, "II", "Thông số chính của hệ thống thử nghiệm")
    setup_rows = [
        ("Nền tảng", "Ubuntu 24.04; ROS 2 Jazzy; Gazebo Harmonic"),
        ("CPU", "Intel Core i5-12450HX"),
        ("Costmap", "0,05 m/ô"),
        ("Hình bao robot", "Hình chữ nhật 0,44×0,34 m"),
        ("Vệt bánh", "Vật lý/mô hình: 0,2548 m; hiệu dụng Gazebo: 0,2834 m"),
        ("Giới hạn vận tốc", "vmax=0,30 m/s; |ω|max=0,80 rad/s; |vbánh|max=0,36 m/s"),
        ("Giới hạn gia tốc ngang", "0,18 m/s²"),
        ("Khoảng mẫu", "Kiểm tra 0,02 m; đầu ra 0,05 m"),
    ]
    setup_table = add_ieee_table(doc, ["Thành phần", "Giá trị"], setup_rows, widths=[1.22, 2.03], font_size=7.0)
    for row in setup_table.rows[:-1]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True

    add_heading2(doc, "B. Kịch bản thử nghiệm")
    add_body(doc, (
        "Ba bản đồ được chọn đúng theo phạm vi bài báo: (i) không gian mở với một khối cản trung tâm, từ (−2,20; −0,60) m "
        "đến (1,20; −0,60) m; (ii) lối đi hẹp theo đường chéo tây nam–đông bắc, từ (−5,00; −3,00) m đến (5,00; 3,00) m; "
        "và (iii) kho có lối giao cắt, từ (−2,00; −2,80) m đến (2,00; 2,80) m. Năm bộ lập kế hoạch tạo tổng cộng 15 nhóm "
        "đầu vào; tất cả năm phương án đều hoàn tất trong 15 nhóm này."), first=True)

    add_figure(doc, environment_paths["open_arena"], 3.18,
               "Toàn cảnh ca thử nghiệm dùng bộ lập kế hoạch Theta* trong không gian mở. Bản đồ lưới và đường đi được ghi từ RViz2; các hình chữ nhật xanh nhạt biểu diễn những tư thế hình bao robot lấy mẫu dọc đường.", 6)
    add_figure(doc, environment_paths["narrow_aisles"], 3.18,
               "Toàn cảnh ca thử nghiệm dùng bộ lập kế hoạch Theta* trong lối đi hẹp. Path Smoothing and Turning-Maneuver Optimization phân bố lại chuyển hướng nhưng vẫn giữ đường trong cùng hành lang.", 7)
    add_figure(doc, environment_paths["warehouse_cross_aisles"], 3.18,
               "Toàn cảnh ca thử nghiệm dùng bộ lập kế hoạch Theta* trong kho có lối giao cắt. Đường sau làm mượt và toàn bộ các tư thế hình bao được hậu kiểm trên cùng bản đồ chi phí.", 8)

    add_heading2(doc, "C. Chỉ số đánh giá")
    add_body(doc, (
        "Ba chỉ số duy nhất dùng để so sánh là chiều dài L trong (2), mức uốn tích lũy Eκ trong (4), và thời gian xử lý thuật "
        "toán tính từ lúc nhận đến khi trả đường. Thời gian Raw phản ánh bước chuyển tiếp/sao chép trong cùng bộ đo, không gồm "
        "thời gian lập kế hoạch toàn cục. Bộ đo có lượng tử khoảng 3 ms, vì vậy các trung bình rất nhỏ chỉ biểu thị dưới độ phân "
        "giải đáng tin cậy chứ không phải chi phí bằng không."), first=True)
    add_body(doc, (
        "L được cộng từ khoảng cách Euclid giữa các mẫu liên tiếp. Để tính Eκ thống nhất giữa các phương pháp có mật độ điểm "
        "khác nhau, mỗi đường được đánh giá theo cùng quy trình lấy mẫu theo chiều dài cung; độ cong rời rạc sau đó được bình "
        "phương và tích phân số. Trung bình toàn bộ được lấy trên 15 cặp bản đồ–bộ lập kế hoạch, tức mỗi bộ lập kế hoạch đóng góp "
        "một đường cho từng môi trường. Cách ghép cặp này bảo đảm mọi bộ làm mượt trong một nhóm nhận cùng đường Raw."))

    add_heading2(doc, "D. Kết quả và thảo luận")
    add_table_caption(doc, "III", "Trung bình trên 15 nhóm bản đồ–bộ lập kế hoạch")
    labels = {"raw": "Raw", "simple": "Simple", "savitzky_golay": "Savitzky–Golay", "constrained": "Constrained", "pstmo": "PSTMO"}
    overall_rows = []
    for method in ["raw", "simple", "savitzky_golay", "constrained", "pstmo"]:
        m = metrics["all"][method]
        overall_rows.append((labels[method], f"{m['L']:.3f}".replace(".", ","), f"{m['E']:.3f}".replace(".", ","), f"{m['T']:.1f}".replace(".", ",")))
    add_ieee_table(doc, ["Phương pháp", "L (m)", "Eκ (m⁻¹)", "T (ms)"], overall_rows, widths=[1.27, 0.62, 0.78, 0.58], font_size=7.4)
    add_body(doc, (
        "Bảng III cho thấy PSTMO đạt Eκ=2,887 m⁻¹, thấp nhất trong năm phương án. So với Raw, Simple, Savitzky–Golay và "
        "Constrained, mức giảm tương ứng là 98,26%, 75,42%, 91,06% và 90,75%. Đồng thời, chiều dài trung bình giảm lần lượt "
        "1,97%, 0,98%, 1,62% và 2,01%. Do đó, cải thiện mức uốn không đạt được bằng cách kéo dài hoặc đi vòng hơn trong bộ ca "
        "này. Tuy nhiên, PSTMO cần 96,0 ms, cao hơn Simple (1,0 ms), Savitzky–Golay (0,2 ms) và Constrained (19,0 ms)."), first=True)
    add_body(doc, (
        "Raw có Eκ rất lớn vì độ cong được tập trung tại các đổi hướng rời rạc. Simple giảm mạnh mức này nhờ dịch chuyển lặp "
        "các điểm, trong khi Savitzky–Golay vẫn để lại biến thiên cục bộ nên Eκ trung bình cao hơn Simple dù thời gian thấp. "
        "Constrained ưu tiên độ trơn và chi phí costmap trong đúng cấu hình đã thử, nhưng wcurve=0 làm Eκ không phải thành phần "
        "được tối ưu trực tiếp. PSTMO tác động đúng vào lân cận góc và ép κ về không ở hai đầu chuyển tiếp, phù hợp với mức giảm "
        "Eκ quan sát được."))

    env_labels = {
        "open_arena": "Không gian mở",
        "narrow_aisles": "Lối đi hẹp",
        "warehouse_cross_aisles": "Kho giao cắt",
    }
    for suffix, env in zip(("a", "b", "c"), ("open_arena", "narrow_aisles", "warehouse_cross_aisles")):
        env_rows = []
        for method in ["raw", "simple", "savitzky_golay", "constrained", "pstmo"]:
            value = metrics[env][method]
            env_rows.append((labels[method], f"{value['L']:.3f}".replace(".", ","),
                             f"{value['E']:.3f}".replace(".", ","), f"{value['T']:.1f}".replace(".", ",")))
        add_table_caption(doc, f"IV({suffix})", f"Kết quả trung bình — {env_labels[env]}")
        env_table = add_ieee_table(doc, ["Phương pháp", "L (m)", "Eκ (m⁻¹)", "T (ms)"], env_rows,
                                   widths=[1.35, 0.60, 0.75, 0.55], font_size=7.4)
        for row in env_table.rows:
            for cell in row.cells:
                set_cell_margins(cell, top=28, start=45, bottom=28, end=45)
        for row in env_table.rows[:-1]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    add_body(doc, (
        "Theo Bảng IV(a)–(c), mức giảm Eκ của PSTMO so với Simple là 84,36% ở không gian mở, 51,03% ở lối đi hẹp và 82,89% ở "
        "kho giao cắt. Lối đi hẹp có nhiều góc liên tiếp và ít chiều dài cạnh dùng chung, nên điều kiện (9) thu hẹp không gian "
        "lựa chọn; đây là môi trường PSTMO cải thiện ít nhất và xử lý lâu nhất (154,2 ms). Ngược lại, vùng mở cho phép chọn đoạn "
        "chuyển tiếp rộng hơn và đạt Eκ=1,954 m⁻¹."), first=True)
    add_body(doc, (
        "Ở không gian mở, cả bốn bộ làm mượt đều duy trì chiều dài quanh 4,2–4,4 m, nhưng PSTMO tách biệt rõ về Eκ. Trong lối "
        "đi hẹp, đường dài khoảng 14 m và có nhiều đoạn dùng chung giữa các góc; thời gian PSTMO tăng gần gấp ba so với vùng mở "
        "do phải đánh giá nhiều phương án và quan hệ tương thích hơn. Tại kho giao cắt, các dãy kệ tạo nhiều lựa chọn đổi hướng "
        "nhưng hành lang vẫn đủ rộng để PSTMO đạt Eκ=2,380 m⁻¹ và L=7,541 m."))
    add_body(doc, (
        "Chiều dài của PSTMO thấp hơn mọi phương án khi lấy trung bình trên cả ba môi trường, nhưng mức chênh chỉ khoảng 1–2%. "
        "Vì vậy, kết luận chính không phải PSTMO tìm một tuyến toàn cục ngắn hơn; tuyến vẫn do bộ lập kế hoạch quyết định. Lợi ích "
        "chủ yếu là phân bố lại chuyển hướng trong cùng hành lang để giảm mức uốn tích lũy. Đổi lại, thời gian xử lý tăng theo "
        "số góc và số ứng viên khả thi; đây là đánh đổi cần được xét khi tần suất lập kế hoạch lại cao."))

    add_figure(doc, results_e_path, 3.18,
               "Tích phân bình phương độ cong trung bình trên ba môi trường; trục tung dùng thang logarit và mỗi cột là trung bình của năm bộ lập kế hoạch toàn cục.", 9)
    add_figure(doc, results_t_path, 3.18,
               "Thời gian xử lý thuật toán trung bình trên ba môi trường; mỗi cột là trung bình của năm bộ lập kế hoạch toàn cục.", 10)
    add_body(doc, (
        "Hình 9 và Hình 10 xác nhận đánh đổi chính: PSTMO luôn nằm thấp nhất ở biểu đồ Eκ nhưng cao nhất về thời gian xử lý. Thứ tự này "
        "ổn định trên cả ba môi trường, nên kết quả tổng hợp không bị tạo bởi riêng một bản đồ. Đồng thời, khoảng cách thời gian "
        "giữa lối đi hẹp và hai môi trường còn lại phản ánh ảnh hưởng của số góc và các ràng buộc cạnh chung, thay vì chỉ phụ "
        "thuộc chiều dài đường."), first=True)
    add_body(doc, (
        "Kết quả chỉ chứng minh chất lượng hình học trên ba bản đồ tĩnh và một điểm đầu–đích cho mỗi bản đồ. Một lượt chạy không "
        "cho phép lập khoảng tin cậy; lượng tử thời gian làm các giá trị dưới vài mili giây kém phân giải. Ngoài ra, kiểm tra "
        "không va chạm là nhị phân trên costmap đã dùng, không đồng nghĩa với mọi biên an toàn ngoài thực địa. PSTMO tìm trên "
        "tập tham số rời rạc nên không bảo đảm tối ưu liên tục toàn cục; liên tục G² theo hình học cũng chưa bảo đảm gia tốc giật "
        "theo thời gian bằng không."))

    add_heading1(doc, "VI. KẾT LUẬN")
    add_body(doc, (
        "Bài báo đã giải quyết bước chuyển từ đa tuyến có góc gãy của bộ lập kế hoạch Nav2 sang đường phù hợp hơn với robot "
        "vi sai. PSTMO dùng chuyển tiếp Bézier bậc năm G², kiểm tra động học và vùng quét hình bao, rồi phối hợp các góc bằng "
        "quy hoạch động. Trên 15 nhóm ghép cặp thuộc ba môi trường, PSTMO đạt L=8,645 m và Eκ=2,887 m⁻¹; Eκ giảm 75,42% "
        "so với Simple và 90,75% so với Constrained, đổi lại thời gian xử lý trung bình tăng lên 96,0 ms."), first=True)
    add_body(doc, (
        "Các bước tiếp theo gồm lặp lại nhiều lần trên nhiều điểm đầu–đích, báo cáo sai số bám đường, khảo sát nhiễu, trượt bánh, "
        "tải và vật cản động, đồng thời thử nghiệm robot thật. Về thuật toán, cần tối ưu tham số liên tục, đưa khoảng hở vào ràng "
        "buộc thay vì chỉ kiểm tra va chạm, và gắn tham số hóa thời gian để đánh giá gia tốc giật cũng như năng lượng thực."))

    references_heading = add_heading1(doc, "TÀI LIỆU THAM KHẢO")
    references_heading.paragraph_format.page_break_before = True
    refs = [
        "S. Macenski, F. Martín, R. White, and J. Ginés Clavero, “The Marathon 2: A Navigation System,” in Proc. IEEE/RSJ Int. Conf. Intelligent Robots and Systems (IROS), 2020, pp. 2718–2725, doi: 10.1109/IROS45743.2020.9341207.",
        "S. Quinlan and O. Khatib, “Elastic Bands: Connecting Path Planning and Control,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 1993, pp. 802–807.",
        "S. Fleury, P. Souères, J.-P. Laumond, and R. Chatila, “Primitives for Smoothing Mobile Robot Trajectories,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 1993, pp. 832–839.",
        "M. Brezak and I. Petrović, “Path Smoothing Using Clothoids for Differential Drive Mobile Robots,” IFAC Proc. Volumes, vol. 44, no. 1, pp. 1133–1138, 2011, doi: 10.3182/20110828-6-IT-1002.02944.",
        "K. Yang and S. Sukkarieh, “An Analytical Continuous-Curvature Path-Smoothing Algorithm,” IEEE Trans. Robotics, vol. 26, no. 3, pp. 561–568, Jun. 2010, doi: 10.1109/TRO.2010.2042990.",
        "X. Bu, H. Su, W. Zou, and P. Wang, “Curvature Continuous Path Smoothing Based on Cubic Bezier Curves for Car-Like Vehicles,” in Proc. IEEE Int. Conf. Robotics and Biomimetics (ROBIO), 2015, pp. 1453–1458, doi: 10.1109/ROBIO.2015.7418975.",
        "L. Xu, D. Wang, B. Song, and M. Cao, “Global Smooth Path Planning for Mobile Robots Based on Continuous Bezier Curve,” in Proc. Chinese Automation Congress (CAC), 2017, pp. 2081–2085, doi: 10.1109/CAC.2017.8243114.",
        "V. Parque and T. Miyashita, “Smooth Curve Fitting of Mobile Robot Trajectories Using Differential Evolution,” IEEE Access, vol. 8, pp. 82855–82866, 2020, doi: 10.1109/ACCESS.2020.2991003.",
        "D. Dolgov, S. Thrun, M. Montemerlo, and J. Diebel, “Practical Search Techniques in Path Planning for Autonomous Driving,” in Proc. 1st Int. Symp. Search Techniques in Artificial Intelligence and Robotics, 2008.",
        "N. Ratliff, M. Zucker, J. A. Bagnell, and S. Srinivasa, “CHOMP: Gradient Optimization Techniques for Efficient Motion Planning,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 2009, pp. 489–494.",
        "C. Rösmann, W. Feiten, T. Wösch, F. Hoffmann, and T. Bertram, “Trajectory Modification Considering Dynamic Constraints of Autonomous Robots,” in ROBOTIK 2012, 2012, pp. 74–79.",
        "C. Rösmann, F. Hoffmann, and T. Bertram, “Timed-Elastic-Bands for Time-Optimal Point-to-Point Nonlinear Model Predictive Control,” in Proc. European Control Conf. (ECC), 2015, pp. 3352–3357, doi: 10.1109/ECC.2015.7331052.",
        "C. Rösmann, F. Hoffmann, and T. Bertram, “Kinodynamic Trajectory Optimization and Control for Car-Like Robots,” in Proc. IEEE/RSJ Int. Conf. Intelligent Robots and Systems (IROS), 2017, pp. 5681–5686, doi: 10.1109/IROS.2017.8206458.",
        "E. Heiden, L. Palmieri, S. Koenig, K. O. Arras, and G. S. Sukhatme, “Gradient-Informed Path Smoothing for Wheeled Mobile Robots,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 2018, pp. 1710–1717, doi: 10.1109/ICRA.2018.8460818.",
        "H. Andreasson, J. Saarinen, M. Cirillo, T. Stoyanov, and A. J. Lilienthal, “Fast, Continuous State Path Smoothing to Improve Navigation Accuracy,” in Proc. IEEE Int. Conf. Robotics and Automation (ICRA), 2015, pp. 662–669, doi: 10.1109/ICRA.2015.7139250.",
    ]
    for i, ref in enumerate(refs, 1):
        add_reference(doc, i, ref)
        if i == 8:
            column_break = doc.add_paragraph()
            column_break.paragraph_format.space_after = Pt(0)
            column_break.add_run().add_break(WD_BREAK.COLUMN)

    # Keep document properties clean and conference-appropriate.
    doc.core_properties.title = "Làm mượt đường đi bằng tối ưu chuyển tiếp tại góc có xét hình bao cho robot di động vi sai"
    doc.core_properties.subject = "Bản thảo tiếng Việt theo mẫu IEEE A4 cho ICEEIS"
    doc.core_properties.author = "Hai Linh Pham; Tien Cuong Nguyen; Viet Bao Nguyen; Thi Phuong Thao Nguyen; Thi Ly Pham"
    doc.core_properties.keywords = "PSTMO; path smoothing; differential drive; quintic Bezier; footprint; Nav2"
    doc.core_properties.comments = "Số liệu lấy từ benchmark_hinh_hoc_175_luot.csv; hình bản đồ lấy từ PGM và bản ghi RViz2 JSON."

    doc.save(str(OUTPUT_DOCX))


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    _, metrics = load_results()
    with tempfile.TemporaryDirectory(prefix="iceeis_paper_") as temp_name:
        temp = Path(temp_name)
        converted = temp / ABSTRACT_TEMPLATE.name
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx", "--outdir", str(temp), str(ABSTRACT_TEMPLATE)],
            check=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
        )
        if not converted.exists():
            raise RuntimeError("LibreOffice did not create a Transitional OOXML template copy")
        problem = temp / "corner_problem.png"
        drive = temp / "differential_drive.png"
        swept = temp / "footprint_sweep.png"
        bezier = temp / "bezier_transition.png"
        edge = temp / "shared_edge.png"
        results_e = temp / "curvature_results.png"
        results_t = temp / "runtime_results.png"
        environment_paths = {
            "open_arena": temp / "open_arena_rviz.png",
            "narrow_aisles": temp / "narrow_aisles_rviz.png",
            "warehouse_cross_aisles": temp / "warehouse_cross_aisles_rviz.png",
        }
        make_corner_figure_v2(problem)
        make_drive_figure_v2(drive)
        make_swept_figure_v2(swept)
        make_bezier_figure(bezier)
        make_shared_edge_figure_v2(edge)
        make_environment_roi_figure(environment_paths["open_arena"], "open_arena", "center_block_detour")
        make_environment_roi_figure(environment_paths["narrow_aisles"], "narrow_aisles", "southwest_northeast_weave")
        make_environment_roi_figure(environment_paths["warehouse_cross_aisles"], "warehouse_cross_aisles", "cross_aisle_transfer")
        make_metric_figure_v2(results_e, metrics, "E")
        make_metric_figure_v2(results_t, metrics, "T")
        build_document(converted, problem, drive, swept, bezier, edge,
                       environment_paths, results_e, results_t, metrics)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
