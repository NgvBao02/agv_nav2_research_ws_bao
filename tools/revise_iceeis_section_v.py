#!/usr/bin/env python3
"""Revise the wording and argument flow of Section V in the ICEEIS paper."""

from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "final_bao_ICEEIS/ICEEIS_2026_PSTMO_final.docx"
OUTPUT = ROOT / "final_bao_ICEEIS/ICEEIS_2026_PSTMO_final_section_V_revised.docx"


REPLACEMENTS = {
    "Thử nghiệm dùng Ubuntu 24.04, ROS 2 Jazzy, Nav2 và Gazebo Harmonic trên máy Intel Core i5-12450HX.": (
        "Năm phương án xử lý đường được đưa vào phép so sánh gồm Raw, Simple, Savitzky–Golay, Constrained và PSTMO. "
        "Trong từng nhóm thử nghiệm, Raw giữ nguyên đường do bộ lập kế hoạch toàn cục tạo và được dùng làm mốc; bốn bộ "
        "làm mượt còn lại nhận đúng đường Raw đó làm đầu vào. Nguyên lý, cấu hình và vai trò của từng phương án được tóm "
        "tắt trong Bảng I."
    ),
    "Bảng I làm rõ rằng Raw không phải một bộ làm mượt.": (
        "Bảng I phân biệt rõ vai trò của các phương án đối chứng. Raw không phải bộ làm mượt; Simple và Savitzky–Golay "
        "đại diện cho hai cách xử lý cục bộ có chi phí thấp, còn Constrained là phương án tối ưu lặp có sử dụng bản đồ chi "
        "phí. Trong cấu hình thử nghiệm, trọng số độ cong của Constrained bằng không nên Eκ không phải thành phần được tối "
        "thiểu hóa trực tiếp. Mọi cấu hình được giữ nguyên trên cả ba bản đồ và năm nguồn đường; do đó, kết quả chỉ mô tả "
        "các cấu hình đã thử, không được hiểu là hiệu năng cực đại có thể đạt được sau khi hiệu chỉnh riêng từng phương pháp."
    ),
    "Năm nguồn đường gồm NavFn A*, NavFn Dijkstra, ThetaStar, Smac2D và SmacHybrid.": (
        "Thử nghiệm được thực hiện trên Ubuntu 24.04, ROS 2 Jazzy, Nav2 và Gazebo Harmonic, với CPU Intel Core i5-12450HX. "
        "Robot có hình bao chữ nhật 0,44×0,34 m. Khoảng cách hai bánh b=0,2548 m được dùng trong mô hình và các phép kiểm tra "
        "động học của bộ làm mượt; riêng plugin DiffDrive trong Gazebo dùng giá trị hiệu dụng 0,2834 m sau hiệu chỉnh tiếp xúc "
        "bánh–mặt sàn và odometry. Các giới hạn chuyển động, độ phân giải bản đồ và khoảng lấy mẫu được liệt kê trong Bảng II."
    ),
    "Ba bản đồ được chọn đúng theo phạm vi bài báo:": (
        "Ba bản đồ được chọn để đại diện cho ba cấu trúc hình học khác nhau: (i) không gian mở với một khối cản trung tâm, "
        "từ (−2,20; −0,60) m đến (1,20; −0,60) m; (ii) lối đi hẹp theo đường chéo tây nam–đông bắc, từ "
        "(−5,00; −3,00) m đến (5,00; 3,00) m; và (iii) kho có lối giao cắt, từ (−2,00; −2,80) m đến "
        "(2,00; 2,80) m. Cả năm phương án đều tạo được đầu ra trong toàn bộ 15 nhóm ghép cặp. Hình 6–8 minh họa một nhóm "
        "đại diện trong mỗi môi trường với đường Raw của Theta* và đầu ra PSTMO."
    ),
    "Ba chỉ số duy nhất dùng để so sánh": (
        "Ba chỉ số được dùng để so sánh là chiều dài L trong (2), tích phân bình phương độ cong Eκ trong (4) và thời gian "
        "xử lý thuật toán T. T được đo từ khi phương án nhận đường đến khi trả kết quả và không bao gồm thời gian của bộ lập "
        "kế hoạch toàn cục. Đối với Raw, T chỉ gồm thao tác tiếp nhận, chuẩn hóa và sao chép đường trong cùng giao diện đo; "
        "vì vậy, giá trị này phản ánh chi phí nền chứ không phải thời gian làm mượt. Bộ đo có độ phân giải xấp xỉ 3 ms, nên "
        "các giá trị trung bình nhỏ hơn mức này không thể được phân biệt tin cậy với giá trị bằng không."
    ),
    "L được cộng từ khoảng cách Euclid giữa các mẫu liên tiếp.": (
        "Chiều dài L được tính bằng tổng khoảng cách Euclid giữa các mẫu liên tiếp. Để giảm ảnh hưởng của mật độ điểm khác "
        "nhau, mọi đường đầu ra được lấy mẫu lại theo cùng khoảng cách chiều dài cung trước khi tính độ cong rời rạc và tích "
        "phân số Eκ. Trung bình trong Bảng III được lấy trên 15 nhóm ghép cặp; trung bình trong mỗi phần của Bảng IV được lấy "
        "trên năm đường do năm bộ lập kế hoạch tạo trong cùng môi trường. Thiết kế ghép cặp bảo đảm các phương án trong một "
        "nhóm nhận cùng đường Raw, nhưng các trung bình này vẫn chỉ là thống kê mô tả vì mỗi tổ hợp được chạy một lần."
    ),
    "Bảng III cho thấy PSTMO đạt Eκ=2,887 m⁻¹": (
        "Bảng III tổng hợp kết quả trên 15 nhóm ghép cặp. PSTMO đạt Eκ=2,887 m⁻¹, thấp hơn Raw, Simple, Savitzky–Golay và "
        "Constrained lần lượt 98,26%, 75,42%, 91,06% và 90,75%. Chiều dài trung bình của PSTMO là 8,645 m, thấp hơn bốn "
        "phương án còn lại từ 0,98% đến 2,01%; vì vậy, trong tập thử nghiệm này, mức giảm Eκ không đi kèm với việc kéo dài "
        "đường. Đổi lại, PSTMO cần 96,0 ms, so với 1,0 ms của Simple, 0,2 ms của Savitzky–Golay và 19,0 ms của Constrained."
    ),
    "Raw có Eκ rất lớn": (
        "Giá trị Eκ lớn của Raw xuất phát từ các đổi hướng rời rạc của đường đa tuyến sau khi được đánh giá bằng cùng quy "
        "trình lấy mẫu; vì vậy, đại lượng này nên được hiểu là chỉ số mức uốn theo giao thức đánh giá, không phải độ cong liên "
        "tục tại một đỉnh lý tưởng. Simple làm giảm mạnh chỉ số này bằng cách dịch chuyển lặp các điểm, trong khi "
        "Savitzky–Golay vẫn giữ lại nhiều biến thiên cục bộ. Constrained không trực tiếp tối thiểu hóa Eκ trong cấu hình đã "
        "thử. PSTMO tác động tại lân cận góc và buộc độ cong về không ở hai đầu đoạn chuyển tiếp; kết quả quan sát được phù "
        "hợp với mục tiêu hình học đó. Tuy nhiên, do các phương án không được hiệu chỉnh để tối ưu cùng một hàm mục tiêu, phép "
        "so sánh không đại diện cho mọi cấu hình có thể có của từng phương pháp."
    ),
    "Theo Bảng IV(a)-(c), mức giảm Eκ": (
        "Bảng IV(a)–(c) cho thấy PSTMO đạt Eκ trung bình thấp nhất trong cả ba môi trường: 1,954 m⁻¹ ở không gian mở, "
        "4,326 m⁻¹ ở lối đi hẹp và 2,380 m⁻¹ ở kho có lối giao cắt. So với Simple, mức giảm tương ứng là 84,36%, 51,03% "
        "và 82,89%. Lối đi hẹp có mức cải thiện tương đối nhỏ nhất, đồng thời có thời gian PSTMO lớn nhất là 154,2 ms. Xu "
        "hướng này phù hợp với việc các góc liên tiếp phải chia sẻ ngân sách cạnh theo (9), nhưng thí nghiệm hiện tại chưa "
        "tách riêng ảnh hưởng của số góc, chiều dài cạnh và khoảng hở vật cản để thiết lập quan hệ nhân quả."
    ),
    "Ở không gian mở, cả bốn bộ làm mượt": (
        "Thời gian xử lý của PSTMO thay đổi rõ theo môi trường: 54,0 ms ở không gian mở, 154,2 ms ở lối đi hẹp và 79,8 ms "
        "ở kho có lối giao cắt. Thứ tự này nhất quán với số lượng phương án tại góc và quan hệ tương thích mà thuật toán phải "
        "đánh giá, thay vì chỉ phụ thuộc vào chiều dài đường. Tuy nhiên, mỗi bản đồ chỉ có một cặp điểm đầu–đích, nên các số "
        "liệu trên chưa đủ để xây dựng một mô hình tổng quát về thời gian xử lý theo cấu trúc môi trường."
    ),
    "Chiều dài của PSTMO thấp hơn mọi phương án": (
        "Chiều dài của PSTMO thấp hơn các phương án còn lại trong trung bình tổng hợp, nhưng mức chênh chỉ khoảng 1–2%. "
        "Kết quả này không có nghĩa PSTMO tìm được một tuyến toàn cục ngắn hơn, vì tuyến và thứ tự hành lang vẫn do bộ lập "
        "kế hoạch quyết định. Lợi ích chính là phân bố lại chuyển hướng trong cùng hành lang để giảm Eκ mà không làm tăng "
        "chiều dài; đánh đổi tương ứng là chi phí xử lý lớn hơn khi số góc và số phương án cần kiểm tra tăng."
    ),
    "Hình 9 và Hình 10 xác nhận đánh đổi chính": (
        "Hình 9 và Hình 10 trực quan hóa các trung bình theo môi trường trong Bảng IV. PSTMO có Eκ trung bình thấp nhất và "
        "thời gian xử lý trung bình cao nhất ở cả ba môi trường; do đó, đánh đổi quan sát trong Bảng III không bị chi phối bởi "
        "riêng một bản đồ. Tuy vậy, mỗi cột chỉ là trung bình của năm đường đầu vào và không kèm độ bất định, nên các biểu đồ "
        "được dùng để mô tả xu hướng chứ không hỗ trợ suy luận thống kê."
    ),
    "Kết quả chỉ chứng minh chất lượng hình học": (
        "Phạm vi đánh giá còn giới hạn ở ba bản đồ tĩnh, một cặp điểm đầu–đích cho mỗi bản đồ, năm nguồn đường và một lần chạy "
        "cho mỗi tổ hợp. Vì vậy, chưa thể ước lượng độ bất định hoặc khái quát thời gian xử lý sang các tuyến khác. Độ phân giải "
        "của bộ đo cũng làm các giá trị dưới vài mili giây kém tin cậy. Kiểm tra không phát hiện va chạm là kết quả nhị phân "
        "trên bản đồ chi phí đã dùng, không thay thế đánh giá khoảng hở và sai số mô hình ngoài thực địa. Ngoài ra, Eκ chỉ mô "
        "tả hình học, không trực tiếp phản ánh sai số bám đường, năng lượng hoặc tải động lực học. PSTMO tìm kiếm trên tập tham "
        "số rời rạc nên không bảo đảm tối ưu liên tục toàn cục; liên tục G² theo hình học cũng không bảo đảm gia tốc giật theo "
        "thời gian bằng không."
    ),
}


SPECIAL_REPLACEMENTS = {
    "Hình 6. Toàn cảnh ca thử nghiệm": (
        "Hình 6. ",
        "Ca thử nghiệm đại diện với Theta* trong không gian mở. Đường Raw, đường PSTMO và các tư thế hình bao robot được "
        "hiển thị trên cùng bản đồ lưới ghi từ RViz2.",
    ),
    "Hình 7. Toàn cảnh ca thử nghiệm": (
        "Hình 7. ",
        "Ca thử nghiệm đại diện với Theta* trong lối đi hẹp. PSTMO phân bố lại chuyển hướng trong khi vẫn giữ đường trong "
        "hành lang ban đầu.",
    ),
    "Hình 8. Toàn cảnh ca thử nghiệm": (
        "Hình 8. ",
        "Ca thử nghiệm đại diện với Theta* trong kho có lối giao cắt. Đường PSTMO và các tư thế hình bao robot được kiểm tra "
        "cuối trên cùng bản đồ chi phí.",
    ),
    "BẢNG III\nTRUNG BÌNH TRÊN 15 NHÓM BẢN ĐỒ-BỘ LẬP KẾ HOẠCH": (
        "BẢNG III\n",
        "TRUNG BÌNH TRÊN 15 NHÓM GHÉP CẶP",
    ),
    "BẢNG IV(c)\nKẾT QUẢ TRUNG BÌNH — KHO GIAO CẮT": (
        "BẢNG IV(c)\n",
        "KẾT QUẢ TRUNG BÌNH — KHO CÓ LỐI GIAO CẮT",
    ),
}


PROTOCOL_PARAGRAPH = (
    "Trong mỗi môi trường, năm bộ lập kế hoạch NavFn A*, NavFn Dijkstra, Theta*, Smac2D và SmacHybrid tạo năm đường Raw. "
    "Simple, Savitzky–Golay, Constrained và PSTMO sau đó được chạy độc lập trên từng đường, tạo 15 nhóm ghép cặp trên ba "
    "môi trường. Các phương án trong cùng nhóm dùng chung bản đồ chi phí và giới hạn chuyển động. Mỗi tổ hợp chỉ được chạy một "
    "lần; do đó, các kết quả dưới đây là thống kê mô tả và không được dùng để lập khoảng tin cậy."
)


def first_run_properties(paragraph):
    for run in paragraph.runs:
        if run._r.rPr is not None:
            return deepcopy(run._r.rPr)
    return None


def replace_text(paragraph, text: str, template=None) -> None:
    template = template or paragraph
    run_properties = first_run_properties(template)
    if template is not paragraph and template._p.pPr is not None:
        old_ppr = paragraph._p.pPr
        if old_ppr is not None:
            paragraph._p.remove(old_ppr)
        paragraph._p.insert(0, deepcopy(template._p.pPr))
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def replace_two_runs(paragraph, first_text: str, second_text: str) -> None:
    if len(paragraph.runs) < 2:
        raise RuntimeError(f"Expected two formatted runs in paragraph: {paragraph.text!r}")
    first_properties = deepcopy(paragraph.runs[0]._r.rPr) if paragraph.runs[0]._r.rPr is not None else None
    second_properties = deepcopy(paragraph.runs[1]._r.rPr) if paragraph.runs[1]._r.rPr is not None else None
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    first_run = paragraph.add_run(first_text)
    second_run = paragraph.add_run(second_text)
    if first_properties is not None:
        first_run._r.insert(0, first_properties)
    if second_properties is not None:
        second_run._r.insert(0, second_properties)


def revise() -> None:
    document = Document(SOURCE)
    replaced = set()
    special_replaced = set()
    for paragraph in document.paragraphs:
        original = paragraph.text.strip()
        matched_special = False
        for prefix, (first_text, second_text) in SPECIAL_REPLACEMENTS.items():
            if original.startswith(prefix):
                replace_two_runs(paragraph, first_text, second_text)
                special_replaced.add(prefix)
                matched_special = True
                break
        if matched_special:
            continue
        for prefix, revised in REPLACEMENTS.items():
            if original.startswith(prefix):
                replace_text(paragraph, revised)
                replaced.add(prefix)
                break

    missing = set(REPLACEMENTS) - replaced
    if missing:
        raise RuntimeError(f"Could not locate paragraphs: {sorted(missing)}")
    missing_special = set(SPECIAL_REPLACEMENTS) - special_replaced
    if missing_special:
        raise RuntimeError(f"Could not locate specially formatted paragraphs: {sorted(missing_special)}")

    scenario_heading_index = next(
        index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == "B. Kịch bản thử nghiệm"
    )
    protocol_paragraph = document.paragraphs[scenario_heading_index - 1]
    if protocol_paragraph.text.strip():
        raise RuntimeError("Expected a blank paragraph immediately before subsection B")
    body_template = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith("Thử nghiệm được thực hiện trên Ubuntu 24.04")
    )
    replace_text(protocol_paragraph, PROTOCOL_PARAGRAPH, template=body_template)

    document.core_properties.comments = (
        "Phần V đã được biên tập lại để làm rõ thiết kế ghép cặp, giao thức đo, phạm vi diễn giải và các giới hạn của kết quả; "
        "toàn bộ số liệu, bảng và hình được giữ nguyên."
    )
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    revise()
